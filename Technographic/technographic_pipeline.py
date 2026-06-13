from __future__ import annotations

import argparse
import csv
import json
import os
import re
import sys
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
INPUT_FILE = BASE_DIR / "input" / "compnys.txt"
RAW_ROOT = BASE_DIR / "raw"
API_LOG_ROOT = BASE_DIR / "apilogs"
REPORTS_DIR = BASE_DIR / "reports"

NOT_AVAILABLE = "Not Available"
NOT_VERIFIED = "Not Verified"

TARGET_FIELDS = [
    "detected technologies (full list)",
    "hardware / devices / endpoint signals",
    "print / MPS / Collaboration (UC) stack",
    "IT spend estimate",
    "last seen/freshness date per technology",
]

API_COMPARISON_FIELDS = [
    "API Name",
    "Endpoint Used",
    "Status (Success/Fail)",
    "Fields Returned",
    "Free-Tier Limitations",
    "Paid-Tier Benefits",
    "Notes",
]

API_TRACING_FIELDS = [
    "Tool Name",
    "Category",
    "API Available (Y/N)",
    "Authentication Type",
    "Free Credits / Tokens",
    "Credits / Tokens Used",
    "Rate Limit",
    "Companies Processed",
    "Coverage (%)",
    "Success Rate (%)",
    "Error Rate (%)",
    "Average Latency",
    "Gated Fields",
    "Free Tier Limitation",
    "Paid Plan Cost",
    "Paid Tier Benefits",
    "Ease of Integration",
    "API Documentation Quality",
    "Evidence Link",
    "Overall API Score",
    "Status",
    "Remarks",
    "Data Completeness (%)",
    "Records Retrieved",
    "Estimated Cost per 100 Companies",
    "Raw Export Saved (Y/N)",
]

CALL_LOG_FIELDS = [
    "timestamp",
    "api_name",
    "company_name",
    "domain",
    "endpoint_used",
    "http_status",
    "status",
    "latency_ms",
    "credits_before",
    "credits_after",
    "credits_used",
    "credit_evidence",
    "rate_limit",
    "records_retrieved",
    "error_message",
    "raw_response_path",
    "api_log_path",
]

COMPANY_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "detected_technologies_full_list",
    "hardware_devices_endpoint_signals",
    "print_mps_collaboration_uc_stack",
    "it_spend_estimate",
    "technology_count",
    "freshness_latest_seen",
    "freshness_oldest_seen",
    "data_completeness_percent",
    "raw_response_path",
]

TECH_DETAIL_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "technology_name",
    "category",
    "vendor",
    "confidence",
    "first_seen",
    "last_seen",
    "freshness_date",
    "source_field",
    "signal_type",
    "raw_response_path",
]

MISSING_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "field",
    "status",
    "reason",
    "raw_response_path",
]

HARDWARE_TERMS = {
    "endpoint", "device", "hardware", "laptop", "desktop", "server", "printer", "scanner",
    "thin client", "pos", "mobile device", "iphone", "ipad", "android", "windows", "macos",
    "linux", "chromebook", "intune", "jamf", "workspace one", "sccm", "endpoint management",
}
PRINT_UC_TERMS = {
    "print", "printer", "mps", "managed print", "xerox", "ricoh", "canon", "konica", "hp printer",
    "collaboration", "unified communications", " uc ", "zoom", "webex", "teams", "slack",
    "google meet", "sharepoint", "poly", "voip", "telephony", "contact center",
}
SPEND_TERMS = {"it_spend", "technology_spend", "spend", "budget", "estimated_it_spend", "ict_spend"}
DATE_KEYS = {
    "last_seen", "last_detected", "detected_at", "updated_at", "last_updated", "last_check_date",
    "freshness_date", "first_verified_at", "last_verified_at", "first_date_found", "last_date_found",
}


@dataclass
class Company:
    source_rank: str
    company_name: str
    domain: str
    linkedin_url: str
    source_basis: str


@dataclass
class ApiAdapter:
    name: str
    display_name: str
    env_key: str
    base_url_env: str
    default_base_url: str
    endpoint_label: str
    auth_type: str
    category: str
    evidence_link: str
    free_tier: str
    paid_benefits: str
    paid_plan_cost: str
    rate_limit: str
    api_doc_quality: str
    estimated_credits_per_success: float | None = None

    def request(self, company: Company, api_key: str, base_url: str) -> dict[str, Any]:
        raise NotImplementedError

    def credit_balance_request(self, api_key: str, base_url: str) -> dict[str, Any] | None:
        return None

    def extract_credit_balance(self, payload: Any) -> str:
        return ""


class TheirStackAdapter(ApiAdapter):
    def request(self, company: Company, api_key: str, base_url: str) -> dict[str, Any]:
        return {
            "method": "POST",
            "url": f"{base_url.rstrip('/')}/v1/companies/technologies",
            "headers": {
                "Accept": "application/json",
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            "json": {
                "company_domain": company.domain,
                "company_name": company.company_name,
                "company_linkedin_url": company.linkedin_url,
            },
        }

    def credit_balance_request(self, api_key: str, base_url: str) -> dict[str, Any] | None:
        return {
            "method": "GET",
            "url": f"{base_url.rstrip('/')}/v0/billing/credit-balance",
            "headers": {"Accept": "application/json", "Authorization": f"Bearer {api_key}"},
        }

    def extract_credit_balance(self, payload: Any) -> str:
        if not isinstance(payload, dict):
            return ""
        if isinstance(payload.get("api_credits"), (int, float)) and isinstance(payload.get("used_api_credits"), (int, float)):
            return str(payload["api_credits"] - payload["used_api_credits"])
        for key in ("remaining_credits", "credits_remaining", "credit_balance", "credits", "available_credits", "balance"):
            value = payload.get(key)
            if isinstance(value, (int, float, str)):
                return str(value)
        for value in payload.values():
            if isinstance(value, dict):
                nested = self.extract_credit_balance(value)
                if nested:
                    return nested
        return ""


class CoresignalAdapter(ApiAdapter):
    def request(self, company: Company, api_key: str, base_url: str) -> dict[str, Any]:
        return {
            "method": "GET",
            "url": f"{base_url.rstrip('/')}/company_multi_source/enrich",
            "headers": {"accept": "application/json", "apikey": api_key},
            "params": {"website": f"https://{company.domain}"},
        }


ADAPTERS: list[ApiAdapter] = [
    TheirStackAdapter(
        name="theirstack",
        display_name="TheirStack Technographics API",
        env_key="THEIRSTACK_API_KEY",
        base_url_env="THEIRSTACK_BASE_URL",
        default_base_url="https://api.theirstack.com",
        endpoint_label="POST /v1/companies/technologies",
        auth_type="Bearer API key",
        category="Technographic company enrichment",
        evidence_link="https://theirstack.com/en/docs/api-reference/companies/technographics_v1",
        free_tier="Credit balance available through GET /v0/billing/credit-balance when account permits.",
        paid_benefits="More credits, higher volume, and deeper technographic/job coverage depending on subscription.",
        paid_plan_cost="Vendor/account specific",
        rate_limit="Vendor/account specific; no rate-limit header required for the run.",
        api_doc_quality="Good - endpoint, authentication, credit endpoint, and response purpose documented.",
        estimated_credits_per_success=3,
    ),
    CoresignalAdapter(
        name="coresignal",
        display_name="Coresignal Multi-source Company Enrichment API",
        env_key="CORESIGNAL_API_KEY",
        base_url_env="CORESIGNAL_BASE_URL",
        default_base_url="https://api.coresignal.com/cdapi/v2",
        endpoint_label="GET /company_multi_source/enrich?website={URL}",
        auth_type="API key header",
        category="Company enrichment and technographic signals",
        evidence_link="https://docs.coresignal.com/company-api/multi-source-company-api",
        free_tier="Free trial/subscription credits are account specific; enrich calls consume collect credits.",
        paid_benefits="Larger collect-credit allowance, broader enrichment volume, bulk collect, and field selection.",
        paid_plan_cost="Vendor/account specific",
        rate_limit="18 enrichment requests per second documented for Multi-source Company API.",
        api_doc_quality="Good - endpoint, rate limit, and credit cost are documented.",
        estimated_credits_per_success=2,
    ),
]


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def load_dotenv() -> None:
    env_path = BASE_DIR / ".env"
    if not env_path.exists():
        return
    for raw in env_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def ensure_dirs() -> None:
    REPORTS_DIR.mkdir(exist_ok=True)
    for adapter in ADAPTERS:
        (RAW_ROOT / adapter.name / "data").mkdir(parents=True, exist_ok=True)
        (API_LOG_ROOT / adapter.name).mkdir(parents=True, exist_ok=True)


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_") or "company"


def normalize_domain(domain: str) -> str:
    domain = re.sub(r"^https?://", "", (domain or "").strip(), flags=re.IGNORECASE)
    domain = domain.split("/")[0]
    return domain[4:] if domain.lower().startswith("www.") else domain


def read_companies(path: Path, limit: int | None) -> list[Company]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = []
        for row in reader:
            if not (row.get("company_name") or "").strip():
                continue
            rows.append(
                Company(
                    source_rank=(row.get("source_rank") or "").strip(),
                    company_name=(row.get("company_name") or "").strip(),
                    domain=normalize_domain(row.get("domain") or ""),
                    linkedin_url=(row.get("linkedin_url") or "").strip(),
                    source_basis=(row.get("source_basis") or "").strip(),
                )
            )
            if limit and len(rows) >= limit:
                break
        return rows


def safe_rel(path: Path) -> str:
    try:
        return str(path.relative_to(BASE_DIR)).replace("\\", "/")
    except ValueError:
        return str(path)


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def append_jsonl(path: Path, event: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(event, ensure_ascii=False) + "\n")


def redact_request(request_spec: dict[str, Any]) -> dict[str, Any]:
    redacted = json.loads(json.dumps(request_spec, default=str))
    headers = redacted.get("headers") or {}
    for key in list(headers):
        if key.lower() in {"authorization", "apikey", "api_key", "x-api-key"}:
            headers[key] = "[REDACTED]"
    return redacted


def clean_headers(headers: requests.structures.CaseInsensitiveDict[str]) -> dict[str, str]:
    keep: dict[str, str] = {}
    for key, value in headers.items():
        lowered = key.lower()
        if lowered in {"date", "content-type"} or "rate" in lowered or "retry-after" in lowered or "credit" in lowered:
            keep[key] = re.sub(r"pk=[A-Za-z0-9]+", "pk=[REDACTED]", value)
    return keep


def http_request(spec: dict[str, Any], timeout: int) -> dict[str, Any]:
    started_at = now_iso()
    started = time.perf_counter()
    try:
        response = requests.request(timeout=timeout, **spec)
        latency = round((time.perf_counter() - started) * 1000, 2)
        try:
            payload = response.json()
        except ValueError:
            payload = {"text": response.text}
        return {
            "ok": 200 <= response.status_code < 300,
            "status_code": response.status_code,
            "latency_ms": latency,
            "payload": payload,
            "headers": clean_headers(response.headers),
            "started_at": started_at,
            "finished_at": now_iso(),
            "error": "",
        }
    except requests.RequestException as exc:
        return {
            "ok": False,
            "status_code": "",
            "latency_ms": round((time.perf_counter() - started) * 1000, 2),
            "payload": {"error": str(exc)},
            "headers": {},
            "started_at": started_at,
            "finished_at": now_iso(),
            "error": str(exc),
        }


def walk(data: Any, path: str = ""):
    yield path, data
    if isinstance(data, dict):
        for key, value in data.items():
            next_path = f"{path}.{key}" if path else str(key)
            yield from walk(value, next_path)
    elif isinstance(data, list):
        for index, value in enumerate(data):
            next_path = f"{path}.{index}" if path else str(index)
            yield from walk(value, next_path)


def scalar(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False)


def date_like(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    return value if re.search(r"\d{4}[-/]\d{1,2}[-/]\d{1,2}|\d{4}-\d{2}-\d{2}T", value) else ""


def first_by_keys(obj: dict[str, Any], keys: list[str]) -> str:
    for key in keys:
        value = obj.get(key)
        if value not in (None, ""):
            return scalar(value)
    return ""


def find_spend(payload: Any) -> str:
    for path, value in walk(payload):
        if any(term in path.lower() for term in SPEND_TERMS) and isinstance(value, (str, int, float)):
            return scalar(value)
    return NOT_AVAILABLE


def extract_theirstack(payload: Any) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    items = payload.get("data", []) if isinstance(payload, dict) else []
    for index, item in enumerate(items if isinstance(items, list) else []):
        if not isinstance(item, dict):
            continue
        tech = item.get("technology") if isinstance(item.get("technology"), dict) else item
        if not isinstance(tech, dict):
            continue
        name = first_by_keys(tech, ["name", "technology", "keyword_name"])
        if not name:
            continue
        first_seen = first_by_keys(item, ["first_date_found", "first_seen"]) or first_by_keys(tech, ["first_date_found", "first_seen"])
        last_seen = first_by_keys(item, ["last_date_found", "last_seen"]) or first_by_keys(tech, ["last_date_found", "last_seen"])
        rows.append(
            {
                "technology_name": name,
                "category": first_by_keys(tech, ["category", "parent_category"]),
                "vendor": first_by_keys(tech, ["vendor", "company", "provider"]),
                "confidence": first_by_keys(item, ["confidence", "score"]),
                "first_seen": first_seen,
                "last_seen": last_seen,
                "freshness_date": last_seen or first_seen or NOT_AVAILABLE,
                "source_field": f"data.{index}",
                "signal_type": "technology",
            }
        )
    return rows


def extract_coresignal(payload: Any) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    items = payload.get("technologies_used", []) if isinstance(payload, dict) else []
    for index, item in enumerate(items if isinstance(items, list) else []):
        if not isinstance(item, dict):
            continue
        name = first_by_keys(item, ["technology", "name", "technology_name"])
        if not name:
            continue
        first_seen = first_by_keys(item, ["first_verified_at", "first_seen"])
        last_seen = first_by_keys(item, ["last_verified_at", "last_seen", "updated_at"])
        rows.append(
            {
                "technology_name": name,
                "category": first_by_keys(item, ["category", "technology_category"]),
                "vendor": first_by_keys(item, ["vendor", "provider"]),
                "confidence": first_by_keys(item, ["confidence", "score"]),
                "first_seen": first_seen,
                "last_seen": last_seen,
                "freshness_date": last_seen or first_seen or NOT_AVAILABLE,
                "source_field": f"technologies_used.{index}",
                "signal_type": "technology",
            }
        )
    return rows


def extract_generic(payload: Any) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for path, value in walk(payload):
        path_lower = path.lower()
        if "technolog" not in path_lower and ".tech" not in path_lower:
            continue
        if isinstance(value, dict):
            name = first_by_keys(value, ["name", "technology", "technology_name", "product", "software_name"])
            if not name or (path, name) in seen:
                continue
            seen.add((path, name))
            freshness = ""
            for key in DATE_KEYS:
                freshness = date_like(value.get(key))
                if freshness:
                    break
            rows.append(
                {
                    "technology_name": name,
                    "category": first_by_keys(value, ["category", "type", "group"]),
                    "vendor": first_by_keys(value, ["vendor", "provider", "manufacturer"]),
                    "confidence": first_by_keys(value, ["confidence", "score"]),
                    "first_seen": first_by_keys(value, ["first_seen", "first_verified_at", "first_date_found"]),
                    "last_seen": first_by_keys(value, ["last_seen", "last_verified_at", "last_date_found", "updated_at"]),
                    "freshness_date": freshness or NOT_AVAILABLE,
                    "source_field": path,
                    "signal_type": "technology",
                }
            )
    return rows


def extract_details(api_name: str, payload: Any) -> list[dict[str, str]]:
    if api_name == "theirstack":
        details = extract_theirstack(payload)
    elif api_name == "coresignal":
        details = extract_coresignal(payload)
    else:
        details = []
    return details or extract_generic(payload)


def classify(details: list[dict[str, str]], terms: set[str]) -> list[str]:
    found: list[str] = []
    for row in details:
        haystack = f" {row.get('technology_name', '')} {row.get('category', '')} {row.get('vendor', '')} {row.get('source_field', '')} ".lower()
        if any(term in haystack for term in terms):
            name = row.get("technology_name", "")
            if name and name not in found:
                found.append(name)
    return found


def summarise_company(company: Company, adapter: ApiAdapter, payload: Any, raw_path: str) -> tuple[dict[str, str], list[dict[str, str]], list[dict[str, str]]]:
    details = extract_details(adapter.name, payload)
    names = sorted({row["technology_name"] for row in details if row.get("technology_name")})
    hardware = classify(details, HARDWARE_TERMS)
    print_uc = classify(details, PRINT_UC_TERMS)
    freshness = sorted({row["freshness_date"] for row in details if row.get("freshness_date") not in ("", NOT_AVAILABLE)})
    spend = find_spend(payload)

    populated = [
        bool(names),
        bool(hardware),
        bool(print_uc),
        spend != NOT_AVAILABLE,
        bool(freshness),
    ]
    completeness = round(sum(populated) / len(TARGET_FIELDS) * 100, 2)
    summary = {
        "source_rank": company.source_rank,
        "company_name": company.company_name,
        "domain": company.domain,
        "api_name": adapter.display_name,
        "detected_technologies_full_list": "; ".join(names) if names else NOT_AVAILABLE,
        "hardware_devices_endpoint_signals": "; ".join(hardware) if hardware else NOT_AVAILABLE,
        "print_mps_collaboration_uc_stack": "; ".join(print_uc) if print_uc else NOT_AVAILABLE,
        "it_spend_estimate": spend,
        "technology_count": str(len(names)),
        "freshness_latest_seen": freshness[-1] if freshness else NOT_AVAILABLE,
        "freshness_oldest_seen": freshness[0] if freshness else NOT_AVAILABLE,
        "data_completeness_percent": f"{completeness:.2f}",
        "raw_response_path": raw_path,
    }

    detail_rows = []
    for detail in details:
        detail_rows.append(
            {
                "source_rank": company.source_rank,
                "company_name": company.company_name,
                "domain": company.domain,
                "api_name": adapter.display_name,
                "raw_response_path": raw_path,
                **detail,
            }
        )

    reasons = [
        (TARGET_FIELDS[0], bool(names), "No technology names were parsed from the response."),
        (TARGET_FIELDS[1], bool(hardware), "No hardware, device, or endpoint terms were detected."),
        (TARGET_FIELDS[2], bool(print_uc), "No print, MPS, UC, or collaboration terms were detected."),
        (TARGET_FIELDS[3], spend != NOT_AVAILABLE, "No spend-like field was present in the raw response."),
        (TARGET_FIELDS[4], bool(freshness), "No first/last seen or freshness date was parsed for technologies."),
    ]
    missing = [
        {
            "source_rank": company.source_rank,
            "company_name": company.company_name,
            "domain": company.domain,
            "api_name": adapter.display_name,
            "field": field,
            "status": NOT_AVAILABLE,
            "reason": reason,
            "raw_response_path": raw_path,
        }
        for field, ok, reason in reasons
        if not ok
    ]
    return summary, detail_rows, missing


def to_float(value: str) -> float | None:
    try:
        return float(str(value).replace(",", ""))
    except (TypeError, ValueError):
        return None


def credit_delta(before: str, after: str) -> str:
    before_num = to_float(before)
    after_num = to_float(after)
    if before_num is None or after_num is None:
        return ""
    return f"{max(before_num - after_num, 0):.2f}".rstrip("0").rstrip(".")


def check_credit_balance(adapter: ApiAdapter, api_key: str, base_url: str, timeout: int) -> tuple[str, str]:
    spec = adapter.credit_balance_request(api_key, base_url)
    if not spec:
        return "", "No provider credit-balance endpoint configured."
    result = http_request(spec, timeout)
    balance = adapter.extract_credit_balance(result["payload"])
    evidence = f"HTTP {result['status_code']} from credit-balance endpoint"
    if not balance:
        evidence += "; balance field was not parsed"
    return balance, evidence


def write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def build_comparison_rows(adapters: list[ApiAdapter], call_rows: list[dict[str, str]], detail_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    rows = []
    for adapter in adapters:
        calls = [row for row in call_rows if row["api_name"] == adapter.display_name]
        successes = [row for row in calls if row["status"] == "Success"]
        returned = sorted(
            field
            for field in {"technology_name", "category", "vendor", "confidence", "first_seen", "last_seen", "freshness_date"}
            if any(row.get(field) for row in detail_rows if row["api_name"] == adapter.display_name)
        )
        notes = "; ".join(sorted({row["error_message"] for row in calls if row["error_message"]})) or f"{len(successes)} successful calls out of {len(calls)}."
        rows.append(
            {
                "API Name": adapter.display_name,
                "Endpoint Used": adapter.endpoint_label,
                "Status (Success/Fail)": "Success" if successes else "Fail",
                "Fields Returned": "; ".join(returned) if returned else NOT_AVAILABLE,
                "Free-Tier Limitations": adapter.free_tier,
                "Paid-Tier Benefits": adapter.paid_benefits,
                "Notes": notes,
            }
        )
    return rows


def score_api(success_rate: float, completeness: float, records: int) -> str:
    score = 1.0
    score += min(success_rate / 100 * 1.5, 1.5)
    score += min(completeness / 100 * 1.5, 1.5)
    score += 1.0 if records else 0.0
    return f"{min(score, 5.0):.2f}/5"


def build_tracing_rows(adapters: list[ApiAdapter], call_rows: list[dict[str, str]], company_rows: list[dict[str, str]], detail_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    rows = []
    for adapter in adapters:
        calls = [row for row in call_rows if row["api_name"] == adapter.display_name]
        company_summaries = [row for row in company_rows if row["api_name"] == adapter.display_name]
        successes = [row for row in calls if row["status"] == "Success"]
        failures = [row for row in calls if row["status"] != "Success"]
        latencies = [float(row["latency_ms"]) for row in calls if row["latency_ms"]]
        records = sum(int(row["records_retrieved"]) for row in calls if row["records_retrieved"].isdigit())
        completeness_values = [float(row["data_completeness_percent"]) for row in company_summaries if row["data_completeness_percent"]]
        avg_completeness = sum(completeness_values) / len(completeness_values) if completeness_values else 0
        used_values = [row["credits_used"] for row in calls if row["credits_used"]]
        real_values = [value for value in used_values if not value.startswith("Estimated")]
        coverage = len(successes) / len(calls) * 100 if calls else 0
        success_rate = coverage
        error_rate = len(failures) / len(calls) * 100 if calls else 0
        cost_per_100 = ""
        if adapter.estimated_credits_per_success is not None:
            cost_per_100 = f"{adapter.estimated_credits_per_success * 100:g} credits per 100 successful companies"
        gated = "; ".join(sorted({row["field"] for row in build_missing_for_api(adapter.display_name, company_rows)})) or "None identified in parsed output"
        rows.append(
            {
                "Tool Name": adapter.display_name,
                "Category": adapter.category,
                "API Available (Y/N)": "Y" if calls else "N",
                "Authentication Type": adapter.auth_type,
                "Free Credits / Tokens": next((row["credits_before"] for row in calls if row["credits_before"]), NOT_VERIFIED),
                "Credits / Tokens Used": "; ".join(real_values or used_values) if used_values else NOT_VERIFIED,
                "Rate Limit": adapter.rate_limit,
                "Companies Processed": str(len(calls)),
                "Coverage (%)": f"{coverage:.2f}",
                "Success Rate (%)": f"{success_rate:.2f}",
                "Error Rate (%)": f"{error_rate:.2f}",
                "Average Latency": f"{(sum(latencies) / len(latencies)):.2f} ms" if latencies else NOT_AVAILABLE,
                "Gated Fields": gated,
                "Free Tier Limitation": adapter.free_tier,
                "Paid Plan Cost": adapter.paid_plan_cost,
                "Paid Tier Benefits": adapter.paid_benefits,
                "Ease of Integration": "Easy - standard REST call; response reshaping needed for HPI target fields.",
                "API Documentation Quality": adapter.api_doc_quality,
                "Evidence Link": adapter.evidence_link,
                "Overall API Score": score_api(success_rate, avg_completeness, records),
                "Status": "Success" if successes else "Fail",
                "Remarks": "; ".join(sorted({row["error_message"] for row in failures if row["error_message"]})) or "Raw responses and API logs saved.",
                "Data Completeness (%)": f"{avg_completeness:.2f}",
                "Records Retrieved": str(records),
                "Estimated Cost per 100 Companies": cost_per_100 or NOT_VERIFIED,
                "Raw Export Saved (Y/N)": "Y" if any(row["raw_response_path"] for row in calls) else "N",
            }
        )
    return rows


def build_missing_for_api(api_name: str, company_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    missing: list[dict[str, str]] = []
    for row in company_rows:
        if row["api_name"] != api_name:
            continue
        checks = [
            (TARGET_FIELDS[0], row["detected_technologies_full_list"] != NOT_AVAILABLE),
            (TARGET_FIELDS[1], row["hardware_devices_endpoint_signals"] != NOT_AVAILABLE),
            (TARGET_FIELDS[2], row["print_mps_collaboration_uc_stack"] != NOT_AVAILABLE),
            (TARGET_FIELDS[3], row["it_spend_estimate"] != NOT_AVAILABLE),
            (TARGET_FIELDS[4], row["freshness_latest_seen"] != NOT_AVAILABLE),
        ]
        missing.extend({"field": field} for field, ok in checks if not ok)
    return missing


def write_workbook(path: Path, sheets: dict[str, tuple[list[str], list[dict[str, Any]]]]) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill(fill_type="solid", fgColor="1F4E79")
    header_font = Font(bold=True, color="FFFFFF")
    for title, (fields, rows) in sheets.items():
        ws = wb.create_sheet(title[:31])
        ws.append(fields)
        for row in rows:
            ws.append([row.get(field, "") for field in fields])
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for row_cells in ws.iter_rows(min_row=2):
            for cell in row_cells:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        for column in ws.columns:
            width = min(max(max(len(str(cell.value or "")) for cell in column) + 2, 12), 60)
            ws.column_dimensions[get_column_letter(column[0].column)].width = width
        ws.freeze_panes = "A2"
    wb.save(path)


def add_doc_table(document: Document, fields: list[str], rows: list[dict[str, Any]], max_rows: int = 25) -> None:
    table = document.add_table(rows=1, cols=len(fields))
    table.style = "Table Grid"
    for index, field in enumerate(fields):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        run = paragraph.add_run(field)
        run.bold = True
        run.font.size = Pt(8)
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for index, field in enumerate(fields):
            cells[index].text = str(row.get(field, ""))
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)


def write_docx(path: Path, comparison: list[dict[str, str]], tracing: list[dict[str, str]], companies: list[dict[str, str]], missing: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HPI Technographic API Comparison Report")
    run.bold = True
    run.font.size = Pt(17)
    document.add_paragraph(f"Generated: {now_iso()}")
    document.add_paragraph("Scope: 10-company technographic extraction using TheirStack and Coresignal. Raw API responses are saved under raw/ and per-call JSONL traces are saved under apilogs/.")
    document.add_heading("API Comparison", level=1)
    add_doc_table(document, API_COMPARISON_FIELDS, comparison)
    document.add_heading("API Tracing", level=1)
    add_doc_table(document, API_TRACING_FIELDS, tracing)
    document.add_heading("Company Technographics", level=1)
    add_doc_table(document, COMPANY_FIELDS, companies, max_rows=20)
    document.add_heading("Missing / Gated Fields", level=1)
    add_doc_table(document, MISSING_FIELDS, missing or [{"field": "None", "status": "Complete", "reason": "No missing fields found in parsed output."}], max_rows=40)
    document.save(path)


def selected_adapters(names: str | None) -> list[ApiAdapter]:
    if not names:
        return ADAPTERS
    wanted = {name.strip().lower() for name in names.split(",") if name.strip()}
    adapters = [adapter for adapter in ADAPTERS if adapter.name in wanted]
    missing = wanted - {adapter.name for adapter in adapters}
    if missing:
        raise ValueError(f"Unknown API name(s): {', '.join(sorted(missing))}")
    return adapters


def raw_has_technology_records(adapter: ApiAdapter, raw_path: Path) -> bool:
    if not raw_path.exists():
        return False
    try:
        return bool(extract_details(adapter.name, load_json(raw_path)))
    except (OSError, json.JSONDecodeError):
        return False


def run(args: argparse.Namespace) -> int:
    load_dotenv()
    ensure_dirs()
    adapters = selected_adapters(args.apis)
    companies = read_companies(Path(args.input) if args.input else INPUT_FILE, args.limit)
    if not companies:
        raise ValueError("No companies found in input file.")

    missing_keys = [adapter.env_key for adapter in adapters if not os.getenv(adapter.env_key) and not args.reuse_raw]
    if missing_keys:
        raise ValueError("Missing API keys in Technographic/.env: " + ", ".join(missing_keys))

    company_rows: list[dict[str, str]] = []
    detail_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []

    for adapter in adapters:
        api_key = os.getenv(adapter.env_key, "")
        base_url = os.getenv(adapter.base_url_env, adapter.default_base_url) or adapter.default_base_url
        for company in companies:
            raw_path = RAW_ROOT / adapter.name / "data" / f"{slugify(company.company_name)}.json"
            log_path = API_LOG_ROOT / adapter.name / f"{slugify(company.company_name)}.jsonl"
            credits_before = ""
            credits_after = ""
            credit_evidence = ""
            should_reuse_raw = raw_path.exists() and (
                args.reuse_raw
                or (args.refresh_empty_raw and raw_has_technology_records(adapter, raw_path))
            )
            if should_reuse_raw:
                result = {
                    "ok": True,
                    "status_code": 200,
                    "latency_ms": 0,
                    "payload": load_json(raw_path),
                    "headers": {},
                    "started_at": now_iso(),
                    "finished_at": now_iso(),
                    "error": "",
                }
                credit_evidence = "Reused saved raw response; no live credits consumed."
            else:
                credits_before, credit_evidence = check_credit_balance(adapter, api_key, base_url, args.timeout)
                spec = adapter.request(company, api_key, base_url)
                result = http_request(spec, args.timeout)
                save_json(raw_path, result["payload"])
                if adapter.credit_balance_request(api_key, base_url):
                    credits_after, after_evidence = check_credit_balance(adapter, api_key, base_url, args.timeout)
                    credit_evidence = f"{credit_evidence}; after call: {after_evidence}"
                append_jsonl(
                    log_path,
                    {
                        "logged_at": now_iso(),
                        "api_name": adapter.display_name,
                        "company": company.__dict__,
                        "request": redact_request(spec),
                        "response": {
                            "status_code": result["status_code"],
                            "ok": result["ok"],
                            "latency_ms": result["latency_ms"],
                            "headers": result["headers"],
                            "error": result["error"],
                        },
                        "raw_response_path": safe_rel(raw_path),
                    },
                )

            raw_rel = safe_rel(raw_path)
            log_rel = safe_rel(log_path)
            summary, details, missing = summarise_company(company, adapter, result["payload"], raw_rel)
            company_rows.append(summary)
            detail_rows.extend(details)
            missing_rows.extend(missing)
            records = len(details)
            status = "Success" if result["ok"] and records else "Fail"
            used = credit_delta(credits_before, credits_after)
            if not used and result["ok"] and records and adapter.estimated_credits_per_success is not None:
                used = f"Estimated {adapter.estimated_credits_per_success:g}"
            if args.reuse_raw:
                used = "0 (reuse-raw)"
            elif should_reuse_raw:
                used = "0 (existing raw with parsed records)"
            error = result["error"]
            if result["ok"] and not records:
                error = "HTTP success but no technology records parsed."
            elif not result["ok"] and not error:
                error = json.dumps(result["payload"], ensure_ascii=False)[:500]
            call_rows.append(
                {
                    "timestamp": now_iso(),
                    "api_name": adapter.display_name,
                    "company_name": company.company_name,
                    "domain": company.domain,
                    "endpoint_used": adapter.endpoint_label,
                    "http_status": str(result["status_code"]),
                    "status": status,
                    "latency_ms": str(result["latency_ms"]),
                    "credits_before": credits_before,
                    "credits_after": credits_after,
                    "credits_used": used,
                    "credit_evidence": credit_evidence,
                    "rate_limit": "; ".join(f"{k}: {v}" for k, v in result["headers"].items() if "rate" in k.lower()) or adapter.rate_limit,
                    "records_retrieved": str(records),
                    "error_message": error,
                    "raw_response_path": raw_rel,
                    "api_log_path": log_rel,
                }
            )
            print(f"{adapter.name}: {company.company_name} | {status} | HTTP {result['status_code']} | records={records} | latency={result['latency_ms']} ms")

    comparison_rows = build_comparison_rows(adapters, call_rows, detail_rows)
    tracing_rows = build_tracing_rows(adapters, call_rows, company_rows, detail_rows)

    write_csv(REPORTS_DIR / "api_comparison_report.csv", comparison_rows, API_COMPARISON_FIELDS)
    write_csv(REPORTS_DIR / "api_tracing_report.csv", tracing_rows, API_TRACING_FIELDS)
    write_csv(REPORTS_DIR / "api_call_log.csv", call_rows, CALL_LOG_FIELDS)
    write_csv(REPORTS_DIR / "company_technographics.csv", company_rows, COMPANY_FIELDS)
    write_csv(REPORTS_DIR / "technology_detail.csv", detail_rows, TECH_DETAIL_FIELDS)
    write_csv(REPORTS_DIR / "missing_fields_report.csv", missing_rows, MISSING_FIELDS)
    save_json(REPORTS_DIR / "company_technographics.json", company_rows)
    save_json(
        REPORTS_DIR / "run_manifest.json",
        {
            "generated_at": now_iso(),
            "input": safe_rel(Path(args.input) if args.input else INPUT_FILE),
            "companies_processed": len(companies),
            "apis": [adapter.name for adapter in adapters],
            "raw_root": safe_rel(RAW_ROOT),
            "api_log_root": safe_rel(API_LOG_ROOT),
            "reports": [
                "reports/api_comparison_report.csv",
                "reports/api_tracing_report.csv",
                "reports/api_call_log.csv",
                "reports/company_technographics.csv",
                "reports/technology_detail.csv",
                "reports/missing_fields_report.csv",
            ],
        },
    )
    xlsx_path = REPORTS_DIR / f"hpi_technographic_api_comparison_{stamp()}.xlsx"
    docx_path = REPORTS_DIR / f"hpi_technographic_api_comparison_{stamp()}.docx"
    write_workbook(
        xlsx_path,
        {
            "API Comparison": (API_COMPARISON_FIELDS, comparison_rows),
            "API Tracing": (API_TRACING_FIELDS, tracing_rows),
            "API Call Log": (CALL_LOG_FIELDS, call_rows),
            "Company Technographics": (COMPANY_FIELDS, company_rows),
            "Technology Detail": (TECH_DETAIL_FIELDS, detail_rows),
            "Missing Fields": (MISSING_FIELDS, missing_rows),
        },
    )
    write_docx(docx_path, comparison_rows, tracing_rows, company_rows, missing_rows)
    print(f"Saved reports to {safe_rel(REPORTS_DIR)}")
    print(f"Workbook: {safe_rel(xlsx_path)}")
    print(f"DOCX: {safe_rel(docx_path)}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Run HPI technographic extraction and API audit for TheirStack and Coresignal.")
    parser.add_argument("--input", help="Input CSV path. Defaults to input/compnys.txt.")
    parser.add_argument("--limit", type=int, default=10, help="Number of companies to process. Defaults to 10.")
    parser.add_argument("--apis", default="theirstack,coresignal", help="Comma-separated APIs: theirstack,coresignal.")
    parser.add_argument("--reuse-raw", action="store_true", help="Build reports from existing raw responses without live API calls.")
    parser.add_argument("--refresh-empty-raw", action="store_true", help="Reuse raw files that already parse technology records; call the API only for missing/empty raw files.")
    parser.add_argument("--timeout", type=int, default=60, help="HTTP timeout in seconds.")
    args = parser.parse_args()
    try:
        return run(args)
    except Exception as exc:
        print(f"Technographic pipeline failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
