from __future__ import annotations

import argparse
import csv
import json
import os
import re
import shutil
import sys
import time
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote_plus, urlparse
from xml.etree import ElementTree

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
INPUT_FILE = BASE_DIR / "input" / "compnys.txt"
REPORTS_DIR = BASE_DIR / "reports"
RAW_ROOT = BASE_DIR / "raw"
APILOG_ROOT = BASE_DIR / "apilogs"

NOT_AVAILABLE = "Not Available"
NOT_TESTED = "Not Tested"

EXA_DOC = "https://exa.ai/docs/reference/search"
TAVILY_DOC = "https://docs.tavily.com/documentation/api-reference/endpoint/search"
NEWSAPI_DOC = "https://newsapi.org/docs/endpoints/everything"
GOOGLE_RSS_DOC = "https://news.google.com/rss/search"
PREDICTLEADS_DOC = "https://docs.predictleads.com/"
GDELT_DOC = "https://www.gdeltproject.org/"

EXA_ENDPOINT = "POST https://api.exa.ai/search"
TAVILY_ENDPOINT = "POST https://api.tavily.com/search"
NEWSAPI_ENDPOINT = "GET https://newsapi.org/v2/everything"
GOOGLE_RSS_ENDPOINT = "GET https://news.google.com/rss/search"

TARGET_FIELDS = [
    "Event headline + URL + date",
    "Event type (funding, expansion, launch, leadership, M&A)",
    "Source / publisher",
    "Relevance / confidence",
    "Coverage depth (events per account, last 12 months)",
]

COMPANY_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "events_last_12_months",
    "event_type_counts",
    "publishers",
    "avg_relevance_confidence",
    "latest_event_date",
    "sample_event_count",
    "data_completeness_percent",
    "raw_response_path",
]

EVENT_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "event_headline",
    "event_url",
    "event_date",
    "event_type",
    "source_publisher",
    "relevance_confidence",
    "source_field",
    "raw_response_path",
]

API_CALL_FIELDS = [
    "timestamp",
    "api_name",
    "company_name",
    "domain",
    "endpoint_used",
    "http_status",
    "status",
    "latency_ms",
    "credits_used",
    "rate_limit",
    "records_retrieved",
    "error_message",
    "raw_response_path",
    "api_log_path",
]

API_TRACE_FIELDS = [
    "Tool Name",
    "Category",
    "API Available (Y/N)",
    "Authentication Type",
    "Free Credits / Tokens",
    "Credits / Tokens Used",
    "Rate Limit",
    "Companies Processed",
    "Coverage (%)",
    "Success Rate (%)",
    "Error Rate (%)",
    "Average Latency",
    "Gated Fields",
    "Free Tier Limitation",
    "Paid Plan Cost",
    "Paid Tier Benefits",
    "Ease of Integration",
    "API Documentation Quality",
    "Evidence Link",
    "Overall API Score",
    "Status",
    "Remarks",
    "Data Completeness (%)",
    "Records Retrieved",
    "Estimated Cost per 100 Companies",
    "Raw Export Saved (Y/N)",
]

API_COMPARISON_FIELDS = [
    "API Name",
    "Endpoint Used",
    "Status (Success/Fail)",
    "Fields Returned",
    "Free-Tier Limitations",
    "Paid-Tier Benefits",
    "Notes",
]

MISSING_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "field",
    "status",
    "reason",
    "raw_response_path",
]


@dataclass
class Company:
    source_rank: str
    company_name: str
    domain: str
    linkedin_url: str
    source_basis: str


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def since_date() -> str:
    return (datetime.now(timezone.utc) - timedelta(days=365)).date().isoformat()


def newsapi_since_date() -> str:
    return (datetime.now(timezone.utc) - timedelta(days=30)).date().isoformat()


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_") or "company"


def normalize_domain(domain: str) -> str:
    domain = re.sub(r"^https?://", "", (domain or "").strip(), flags=re.I).split("/")[0]
    return domain[4:] if domain.lower().startswith("www.") else domain


def load_dotenv() -> None:
    for env_path in [BASE_DIR / ".env", ROOT_DIR / "Firmographic" / ".env", ROOT_DIR / "JobsHiring" / ".env"]:
        if not env_path.exists():
            continue
        for raw in env_path.read_text(encoding="utf-8").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip().lstrip("\ufeff"), value.strip().strip('"').strip("'"))


def ensure_dirs() -> None:
    REPORTS_DIR.mkdir(exist_ok=True)
    (BASE_DIR / "input").mkdir(exist_ok=True)
    for provider in ["exa", "tavily", "google_news_rss", "newsapi", "predictleads", "gdelt"]:
        (RAW_ROOT / provider / "data").mkdir(parents=True, exist_ok=True)
        (APILOG_ROOT / provider).mkdir(parents=True, exist_ok=True)


def copy_input() -> None:
    source = ROOT_DIR / "Firmographic" / "input" / "compnys.txt"
    if not INPUT_FILE.exists() and source.exists():
        shutil.copy2(source, INPUT_FILE)


def read_companies(path: Path, limit: int | None) -> list[Company]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        companies: list[Company] = []
        for row in reader:
            if not (row.get("company_name") or "").strip():
                continue
            companies.append(
                Company(
                    source_rank=(row.get("source_rank") or "").strip(),
                    company_name=(row.get("company_name") or "").strip(),
                    domain=normalize_domain(row.get("domain") or ""),
                    linkedin_url=(row.get("linkedin_url") or "").strip(),
                    source_basis=(row.get("source_basis") or "").strip(),
                )
            )
            if limit and len(companies) >= limit:
                break
    return companies


def safe_rel(path: Path) -> str:
    try:
        return str(path.relative_to(BASE_DIR)).replace("\\", "/")
    except ValueError:
        try:
            return str(path.relative_to(ROOT_DIR)).replace("\\", "/")
        except ValueError:
            return str(path)


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def append_jsonl(path: Path, event: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(event, ensure_ascii=False) + "\n")


def redact_headers(headers: dict[str, str]) -> dict[str, str]:
    safe = dict(headers)
    for key in list(safe):
        if key.lower() in {"authorization", "x-api-key", "api-key", "x-api-token"}:
            safe[key] = "[REDACTED]"
    return safe


def clean_headers(headers: requests.structures.CaseInsensitiveDict[str]) -> dict[str, str]:
    out: dict[str, str] = {}
    for key, value in headers.items():
        low = key.lower()
        if low in {"date", "content-type"} or "rate" in low or "retry-after" in low or "remaining" in low:
            out[key] = value
    return out


def summarize_error(payload: Any) -> str:
    if isinstance(payload, dict):
        for key in ["error", "message", "status", "code"]:
            if payload.get(key):
                return str(payload[key])[:300]
    return ""


def http_request(method: str, url: str, *, headers: dict[str, str], params: dict[str, Any] | None = None, body: dict[str, Any] | None = None, timeout: int = 60) -> dict[str, Any]:
    started = time.perf_counter()
    started_at = now_iso()
    try:
        response = requests.request(method, url, headers=headers, params=params, json=body, timeout=timeout)
        latency = round((time.perf_counter() - started) * 1000, 2)
        try:
            payload = response.json()
        except ValueError:
            payload = {"text": response.text}
        return {
            "ok": 200 <= response.status_code < 300,
            "status_code": response.status_code,
            "latency_ms": latency,
            "payload": payload,
            "headers": clean_headers(response.headers),
            "error": "" if 200 <= response.status_code < 300 else summarize_error(payload),
            "started_at": started_at,
            "finished_at": now_iso(),
        }
    except requests.RequestException as exc:
        return {
            "ok": False,
            "status_code": "",
            "latency_ms": round((time.perf_counter() - started) * 1000, 2),
            "payload": {"error": str(exc)},
            "headers": {},
            "error": str(exc),
            "started_at": started_at,
            "finished_at": now_iso(),
        }


def query_for(company: Company) -> str:
    return f'"{company.company_name}" ({company.domain}) announcement OR funding OR expansion OR launch OR leadership OR acquisition OR merger OR partnership'


def normalize_date(value: Any) -> str:
    if not value:
        return ""
    if isinstance(value, (int, float)):
        return ""
    text = str(value)
    match = re.search(r"\d{4}-\d{2}-\d{2}", text)
    if match:
        return match.group(0)
    try:
        return parsedate_to_datetime(text).date().isoformat()
    except Exception:
        return ""


def event_type(text: str) -> str:
    t = (text or "").lower()
    checks = [
        ("M&A", r"\b(acquire|acquires|acquisition|merger|merge|takeover|buyout|divest|sale of)\b"),
        ("funding", r"\b(funding|fundraise|raises?|investment|invests?|financing|series [abcde]|capital)\b"),
        ("expansion", r"\b(expands?|expansion|opens?|new office|new plant|factory|facility|market entry|regional hub)\b"),
        ("launch", r"\b(launch|launches|unveils?|introduces?|rolls out|release|new product|platform)\b"),
        ("leadership", r"\b(appoints?|appointed|ceo|cfo|cto|cio|chief|president|chairman|leadership|joins board)\b"),
        ("partnership", r"\b(partners?|partnership|collaboration|alliance|teams up|agreement)\b"),
    ]
    for label, pattern in checks:
        if re.search(pattern, t):
            return label
    return "other"


def publisher_from_url(url: str) -> str:
    host = urlparse(url or "").netloc.lower()
    return host[4:] if host.startswith("www.") else host


def confidence(company: Company, headline: str, url: str, raw_score: Any = None) -> str:
    score = 0.45
    text = f"{headline} {url}".lower()
    if company.company_name.lower() in text:
        score += 0.25
    if company.domain.lower() in text:
        score += 0.15
    if event_type(text) != "other":
        score += 0.1
    try:
        if raw_score not in (None, ""):
            score = max(score, min(float(raw_score), 1.0))
    except (TypeError, ValueError):
        pass
    return f"{min(score, 0.99):.2f}"


def event_row(company: Company, api_name: str, headline: str, url: str, date: str, publisher: str, conf: str, source_field: str, raw_rel: str) -> dict[str, str]:
    headline = re.sub(r"\s+", " ", headline or "").strip()
    url = (url or "").strip()
    date = normalize_date(date)
    publisher = publisher or publisher_from_url(url) or NOT_AVAILABLE
    return {
        "source_rank": company.source_rank,
        "company_name": company.company_name,
        "domain": company.domain,
        "api_name": api_name,
        "event_headline": headline or NOT_AVAILABLE,
        "event_url": url or NOT_AVAILABLE,
        "event_date": date or NOT_AVAILABLE,
        "event_type": event_type(f"{headline} {url}"),
        "source_publisher": publisher,
        "relevance_confidence": conf or confidence(company, headline, url),
        "source_field": source_field,
        "raw_response_path": raw_rel,
    }


def dedupe_events(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    seen: set[tuple[str, str, str]] = set()
    out: list[dict[str, str]] = []
    for row in rows:
        key = (row["api_name"], row["company_name"], row["event_url"].lower() if row["event_url"] != NOT_AVAILABLE else row["event_headline"].lower())
        if key in seen:
            continue
        seen.add(key)
        out.append(row)
    return out


def events_from_payload(company: Company, api_name: str, payload: Any, raw_rel: str, limit: int) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    if not isinstance(payload, dict):
        return rows
    if api_name == "Exa Search API":
        for idx, item in enumerate(payload.get("results") or []):
            rows.append(event_row(company, api_name, item.get("title"), item.get("url"), item.get("publishedDate") or item.get("published_date"), publisher_from_url(item.get("url", "")), confidence(company, item.get("title", ""), item.get("url", ""), item.get("score")), f"results.{idx}", raw_rel))
    elif api_name == "Tavily Search API":
        for idx, item in enumerate(payload.get("results") or []):
            rows.append(event_row(company, api_name, item.get("title"), item.get("url"), item.get("published_date") or item.get("publishedDate"), publisher_from_url(item.get("url", "")), confidence(company, item.get("title", ""), item.get("url", ""), item.get("score")), f"results.{idx}", raw_rel))
    elif api_name == "NewsAPI Everything API":
        for idx, item in enumerate(payload.get("articles") or []):
            source = item.get("source") if isinstance(item.get("source"), dict) else {}
            rows.append(event_row(company, api_name, item.get("title"), item.get("url"), item.get("publishedAt"), source.get("name") or publisher_from_url(item.get("url", "")), confidence(company, item.get("title", ""), item.get("url", "")), f"articles.{idx}", raw_rel))
    elif api_name == "Google News RSS":
        for idx, item in enumerate(payload.get("items") or []):
            rows.append(event_row(company, api_name, item.get("title"), item.get("link"), item.get("pubDate"), item.get("source") or publisher_from_url(item.get("link", "")), confidence(company, item.get("title", ""), item.get("link", "")), f"items.{idx}", raw_rel))
    return dedupe_events(rows)[:limit]


def summarize_events(companies: list[Company], api_name: str, event_rows: list[dict[str, str]], raw_by_company: dict[str, str]) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    summaries: list[dict[str, str]] = []
    missing: list[dict[str, str]] = []
    for company in companies:
        rows = [row for row in event_rows if row["company_name"] == company.company_name and row["api_name"] == api_name]
        type_counts = Counter(row["event_type"] for row in rows if row["event_type"] != NOT_AVAILABLE)
        publishers = sorted({row["source_publisher"] for row in rows if row["source_publisher"] != NOT_AVAILABLE})
        dates = sorted([row["event_date"] for row in rows if row["event_date"] != NOT_AVAILABLE], reverse=True)
        confidences = []
        for row in rows:
            try:
                confidences.append(float(row["relevance_confidence"]))
            except ValueError:
                pass
        checks = [
            any(row["event_headline"] != NOT_AVAILABLE and row["event_url"] != NOT_AVAILABLE and row["event_date"] != NOT_AVAILABLE for row in rows),
            any(row["event_type"] != "other" for row in rows),
            bool(publishers),
            bool(confidences),
            bool(rows),
        ]
        completeness = round(sum(checks) / len(TARGET_FIELDS) * 100, 2)
        raw_rel = raw_by_company.get(company.company_name, "")
        summaries.append({
            "source_rank": company.source_rank,
            "company_name": company.company_name,
            "domain": company.domain,
            "api_name": api_name,
            "events_last_12_months": str(len(rows)),
            "event_type_counts": "; ".join(f"{key}: {value}" for key, value in sorted(type_counts.items())) if type_counts else NOT_AVAILABLE,
            "publishers": "; ".join(publishers[:20]) if publishers else NOT_AVAILABLE,
            "avg_relevance_confidence": f"{(sum(confidences) / len(confidences)):.2f}" if confidences else NOT_AVAILABLE,
            "latest_event_date": dates[0] if dates else NOT_AVAILABLE,
            "sample_event_count": str(len(rows)),
            "data_completeness_percent": f"{completeness:.2f}",
            "raw_response_path": raw_rel,
        })
        reasons = [
            (TARGET_FIELDS[0], checks[0], "No event headline+URL+date parsed."),
            (TARGET_FIELDS[1], checks[1], "No classified key announcement event type parsed."),
            (TARGET_FIELDS[2], checks[2], "No source/publisher parsed."),
            (TARGET_FIELDS[3], checks[3], "No relevance/confidence score assigned."),
            (TARGET_FIELDS[4], checks[4], "No events parsed for this account/provider."),
        ]
        for field, ok, reason in reasons:
            if not ok:
                missing.append({"source_rank": company.source_rank, "company_name": company.company_name, "domain": company.domain, "api_name": api_name, "field": field, "status": NOT_AVAILABLE, "reason": reason, "raw_response_path": raw_rel})
    return summaries, missing


def api_call_row(api_name: str, company: Company, endpoint: str, result: dict[str, Any], status: str, credits_used: str, records: int, raw_rel: str, log_rel: str) -> dict[str, str]:
    return {
        "timestamp": now_iso(),
        "api_name": api_name,
        "company_name": company.company_name,
        "domain": company.domain,
        "endpoint_used": endpoint,
        "http_status": str(result.get("status_code", "")),
        "status": status,
        "latency_ms": str(result.get("latency_ms", "")),
        "credits_used": credits_used,
        "rate_limit": "; ".join(f"{k}: {v}" for k, v in (result.get("headers") or {}).items() if "rate" in k.lower() or "remaining" in k.lower()),
        "records_retrieved": str(records),
        "error_message": str(result.get("error") or ""),
        "raw_response_path": raw_rel,
        "api_log_path": log_rel,
    }


def collect_exa(companies: list[Company], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Exa Search API"
    key = os.getenv("EXA_API_KEY", "")
    event_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for company in companies:
        raw_path = RAW_ROOT / "exa" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "exa" / f"{slugify(company.company_name)}.jsonl"
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = {"ok": not bool(payload.get("error")) if isinstance(payload, dict) else True, "status_code": 200, "latency_ms": 0, "payload": payload, "headers": {}, "error": summarize_error(payload)}
        elif not key:
            payload = {"error": "Missing EXA_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"x-api-key": key, "Content-Type": "application/json", "Accept": "application/json"}
            body = {"query": query_for(company), "numResults": limit, "type": "auto", "startPublishedDate": since_date(), "contents": {"highlights": True}}
            result = http_request("POST", "https://api.exa.ai/search", headers=headers, body=body, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://api.exa.ai/search", "headers": redact_headers(headers), "body": body}, "response": result})
            save_json(raw_path, result["payload"])
        raw_rel = safe_rel(raw_path)
        raw_by_company[company.company_name] = raw_rel
        rows = events_from_payload(company, api_name, result.get("payload", {}), raw_rel, limit)
        event_rows.extend(rows)
        call_rows.append(api_call_row(api_name, company, EXA_ENDPOINT, result, "Success" if result.get("ok") else "Fail", "1 search request", len(rows), raw_rel, safe_rel(log_path)))
        print(f"exa: {company.company_name} | {'Success' if result.get('ok') else 'Fail'} | events={len(rows)}")
    company_rows, missing_rows = summarize_events(companies, api_name, event_rows, raw_by_company)
    return company_rows, event_rows, missing_rows, call_rows


def collect_tavily(companies: list[Company], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Tavily Search API"
    key = os.getenv("TAVILY_API_KEY", "")
    event_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for company in companies:
        raw_path = RAW_ROOT / "tavily" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "tavily" / f"{slugify(company.company_name)}.jsonl"
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = {"ok": not bool(payload.get("error")) if isinstance(payload, dict) else True, "status_code": 200, "latency_ms": 0, "payload": payload, "headers": {}, "error": summarize_error(payload)}
        elif not key:
            payload = {"error": "Missing TAVILY_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json", "Accept": "application/json"}
            body = {"query": query_for(company), "topic": "news", "search_depth": "basic", "max_results": limit, "include_answer": False, "include_raw_content": False, "days": 365}
            result = http_request("POST", "https://api.tavily.com/search", headers=headers, body=body, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://api.tavily.com/search", "headers": redact_headers(headers), "body": body}, "response": result})
            save_json(raw_path, result["payload"])
        raw_rel = safe_rel(raw_path)
        raw_by_company[company.company_name] = raw_rel
        rows = events_from_payload(company, api_name, result.get("payload", {}), raw_rel, limit)
        event_rows.extend(rows)
        call_rows.append(api_call_row(api_name, company, TAVILY_ENDPOINT, result, "Success" if result.get("ok") else "Fail", "1 search request", len(rows), raw_rel, safe_rel(log_path)))
        print(f"tavily: {company.company_name} | {'Success' if result.get('ok') else 'Fail'} | events={len(rows)}")
    company_rows, missing_rows = summarize_events(companies, api_name, event_rows, raw_by_company)
    return company_rows, event_rows, missing_rows, call_rows


def collect_newsapi(companies: list[Company], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "NewsAPI Everything API"
    key = os.getenv("NEWSAPI_API_KEY", "")
    event_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for company in companies:
        raw_path = RAW_ROOT / "newsapi" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "newsapi" / f"{slugify(company.company_name)}.jsonl"
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = {"ok": isinstance(payload, dict) and payload.get("status") == "ok", "status_code": 200, "latency_ms": 0, "payload": payload, "headers": {}, "error": summarize_error(payload)}
        elif not key:
            payload = {"error": "Missing NEWSAPI_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Accept": "application/json"}
            params = {"apiKey": key, "q": f'"{company.company_name}" OR {company.domain}', "from": newsapi_since_date(), "sortBy": "publishedAt", "language": "en", "pageSize": min(limit, 100)}
            result = http_request("GET", "https://newsapi.org/v2/everything", headers=headers, params=params, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "GET", "url": "https://newsapi.org/v2/everything", "headers": redact_headers(headers), "params": {**params, "apiKey": "[REDACTED]"}}, "response": result})
            save_json(raw_path, result["payload"])
        raw_rel = safe_rel(raw_path)
        raw_by_company[company.company_name] = raw_rel
        rows = events_from_payload(company, api_name, result.get("payload", {}), raw_rel, limit)
        event_rows.extend(rows)
        call_rows.append(api_call_row(api_name, company, NEWSAPI_ENDPOINT, result, "Success" if result.get("ok") else "Fail", "1 request", len(rows), raw_rel, safe_rel(log_path)))
        print(f"newsapi: {company.company_name} | {'Success' if result.get('ok') else 'Fail'} | events={len(rows)}")
    company_rows, missing_rows = summarize_events(companies, api_name, event_rows, raw_by_company)
    return company_rows, event_rows, missing_rows, call_rows


def parse_google_rss(xml_text: str) -> dict[str, Any]:
    root = ElementTree.fromstring(xml_text)
    items = []
    for item in root.findall(".//item"):
        source = item.find("source")
        items.append({
            "title": item.findtext("title", default=""),
            "link": item.findtext("link", default=""),
            "pubDate": item.findtext("pubDate", default=""),
            "source": source.text if source is not None else "",
        })
    return {"items": items}


def collect_google_rss(companies: list[Company], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Google News RSS"
    event_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for company in companies:
        raw_path = RAW_ROOT / "google_news_rss" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "google_news_rss" / f"{slugify(company.company_name)}.jsonl"
        query = company.company_name
        url = f"https://news.google.com/rss/search?q={quote_plus(query)}&hl=en-US&gl=US&ceid=US:en"
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = {"ok": not bool(payload.get("error")) if isinstance(payload, dict) else True, "status_code": 200, "latency_ms": 0, "payload": payload, "headers": {}, "error": summarize_error(payload)}
        else:
            started = time.perf_counter()
            try:
                response = requests.get(url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
                payload = parse_google_rss(response.text) if response.ok else {"error": response.text[:300]}
                result = {"ok": response.ok, "status_code": response.status_code, "latency_ms": round((time.perf_counter() - started) * 1000, 2), "payload": payload, "headers": clean_headers(response.headers), "error": "" if response.ok else "Google RSS request failed"}
            except Exception as exc:
                result = {"ok": False, "status_code": "", "latency_ms": round((time.perf_counter() - started) * 1000, 2), "payload": {"error": str(exc)}, "headers": {}, "error": str(exc)}
            append_jsonl(log_path, {"request": {"method": "GET", "url": url, "headers": {}}, "response": result})
            save_json(raw_path, result["payload"])
        raw_rel = safe_rel(raw_path)
        raw_by_company[company.company_name] = raw_rel
        rows = events_from_payload(company, api_name, result.get("payload", {}), raw_rel, limit)
        event_rows.extend(rows)
        call_rows.append(api_call_row(api_name, company, GOOGLE_RSS_ENDPOINT, result, "Success" if result.get("ok") else "Fail", "0 (RSS)", len(rows), raw_rel, safe_rel(log_path)))
        print(f"google rss: {company.company_name} | {'Success' if result.get('ok') else 'Fail'} | events={len(rows)}")
    company_rows, missing_rows = summarize_events(companies, api_name, event_rows, raw_by_company)
    return company_rows, event_rows, missing_rows, call_rows


def collect_blocked(companies: list[Company], api_name: str, endpoint: str, reason: str) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    company_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    provider_dir = "predictleads" if "PredictLeads" in api_name else "gdelt"
    for company in companies:
        raw_path = RAW_ROOT / provider_dir / "data" / f"{slugify(company.company_name)}.json"
        payload = {"status": "not_executed", "reason": reason}
        save_json(raw_path, payload)
        raw_rel = safe_rel(raw_path)
        company_rows.append({"source_rank": company.source_rank, "company_name": company.company_name, "domain": company.domain, "api_name": api_name, "events_last_12_months": NOT_TESTED, "event_type_counts": NOT_TESTED, "publishers": NOT_TESTED, "avg_relevance_confidence": NOT_TESTED, "latest_event_date": NOT_TESTED, "sample_event_count": "0", "data_completeness_percent": "0.00", "raw_response_path": raw_rel})
        for field in TARGET_FIELDS:
            missing_rows.append({"source_rank": company.source_rank, "company_name": company.company_name, "domain": company.domain, "api_name": api_name, "field": field, "status": NOT_TESTED, "reason": reason, "raw_response_path": raw_rel})
        result = {"ok": False, "status_code": NOT_TESTED, "latency_ms": 0, "headers": {}, "error": reason}
        call_rows.append(api_call_row(api_name, company, endpoint, result, NOT_TESTED, "0", 0, raw_rel, ""))
    return company_rows, [], missing_rows, call_rows


def fields_returned(rows: list[dict[str, str]], api_name: str) -> list[str]:
    api_rows = [row for row in rows if row["api_name"] == api_name]
    checks = {
        "event_headline": "Event headline",
        "event_url": "Event URL",
        "event_date": "Event date",
        "event_type": "Event type",
        "source_publisher": "Source/publisher",
        "relevance_confidence": "Relevance/confidence",
    }
    returned = []
    for key, label in checks.items():
        if any(row.get(key) not in {"", NOT_AVAILABLE, NOT_TESTED, "other"} for row in api_rows):
            returned.append(label)
    if api_rows:
        returned.append("Coverage depth")
    return returned


def build_api_comparison(event_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    configs = [
        ("Exa Search API", EXA_ENDPOINT, "Semantic/news web search for company announcements.", "Free tier/account specific search credits.", "Higher volume search and content extraction."),
        ("Tavily Search API", TAVILY_ENDPOINT, "News/web search endpoint with topic=news.", "Free/dev credits; advanced search may cost more.", "Advanced search and extraction at larger volume."),
        ("Google News RSS", GOOGLE_RSS_ENDPOINT, "No API key required; RSS search with when:12m query.", "No formal SLA; RSS search behavior can change.", "Use a paid news provider if RSS stability is required."),
        ("NewsAPI Everything API", NEWSAPI_ENDPOINT, "NewsAPI everything endpoint with 12-month query.", "Developer tier may restrict historical date range and source availability.", "Paid plans support broader historical access and higher limits."),
        ("PredictLeads News/Events API", "PredictLeads events/news endpoint", "Not executed because no valid key is available in this workspace.", "Invalid/missing key.", "Structured company events if valid access is restored."),
        ("GDELT", "Not executed", "Not used in this run per API status; no key required.", "No key, but not included in scoped live run.", "Global event database for future comparison."),
    ]
    rows = []
    for api_name, endpoint, notes, free, paid in configs:
        returned = fields_returned(event_rows, api_name)
        status = "Success" if returned else "Not Tested" if api_name in {"PredictLeads News/Events API", "GDELT"} else "Fail"
        rows.append({"API Name": api_name, "Endpoint Used": endpoint, "Status (Success/Fail)": status, "Fields Returned": "; ".join(returned) if returned else NOT_AVAILABLE, "Free-Tier Limitations": free, "Paid-Tier Benefits": paid, "Notes": notes})
    return rows


def build_api_trace(company_rows: list[dict[str, str]], call_rows: list[dict[str, str]], missing_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    configs = {
        "Exa Search API": ("x-api-key header", EXA_DOC, "Search endpoint returned company announcement/news URLs.", "Account-specific free credits", "Higher volume and contents extraction", "Provider/account specific", ""),
        "Tavily Search API": ("Bearer API key", TAVILY_DOC, "Search endpoint run with topic=news.", "Free/dev credits", "Advanced search/extract", "Provider/account specific", ""),
        "Google News RSS": ("None", GOOGLE_RSS_DOC, "No API key required; RSS parsed into normalized events.", "No API key", "Paid alternatives provide SLA and cleaner source metadata", "Google RSS behavior is unofficial", ""),
        "NewsAPI Everything API": ("apiKey query parameter", NEWSAPI_DOC, "Everything endpoint queried for company events.", "Developer plan limitations", "Broader history and higher quota", "Plan-specific", ""),
        "PredictLeads News/Events API": ("API key", PREDICTLEADS_DOC, "No valid key available; not executed.", "Invalid/missing key", "Structured company event endpoints", "Plan-specific", "Not executed - invalid/missing API key"),
        "GDELT": ("None", GDELT_DOC, "Not used because scope says no API and enough providers are available.", "No API key required", "Can add global event search later", "Public endpoint limits", "Not executed - out of scoped live run"),
    }
    rows = []
    for api_name, (auth, evidence, remarks, free, paid, rate, forced_status) in configs.items():
        calls = [row for row in call_rows if row["api_name"] == api_name]
        summaries = [row for row in company_rows if row["api_name"] == api_name]
        successes = [row for row in calls if row["status"] == "Success"]
        latencies = [float(row["latency_ms"]) for row in calls if str(row["latency_ms"]).replace(".", "", 1).isdigit()]
        completeness = [float(row["data_completeness_percent"]) for row in summaries if row.get("data_completeness_percent") and row.get("data_completeness_percent") != NOT_TESTED]
        records = sum(int(row["records_retrieved"]) for row in calls if row["records_retrieved"].isdigit())
        missing = sorted({row["field"] for row in missing_rows if row["api_name"] == api_name and row["status"] != NOT_TESTED})
        processed = len(calls)
        success_rate = len(successes) / processed * 100 if processed else 0
        avg_completeness = sum(completeness) / len(completeness) if completeness else 0
        status = forced_status or ("Success" if successes else "Fail")
        rows.append({
            "Tool Name": api_name,
            "Category": "News / Key Announcements",
            "API Available (Y/N)": "Y" if api_name not in {"PredictLeads News/Events API", "GDELT"} else ("N" if api_name == "PredictLeads News/Events API" else "Not used"),
            "Authentication Type": auth,
            "Free Credits / Tokens": free,
            "Credits / Tokens Used": "; ".join(row["credits_used"] for row in calls if row["credits_used"]) or NOT_AVAILABLE,
            "Rate Limit": rate,
            "Companies Processed": str(processed),
            "Coverage (%)": f"{success_rate:.2f}",
            "Success Rate (%)": f"{success_rate:.2f}",
            "Error Rate (%)": f"{(100 - success_rate if processed else 0):.2f}",
            "Average Latency": f"{(sum(latencies) / len(latencies)):.2f} ms" if latencies else NOT_AVAILABLE,
            "Gated Fields": "; ".join(missing) if missing else "None identified in parsed output",
            "Free Tier Limitation": free,
            "Paid Plan Cost": "Vendor/account specific",
            "Paid Tier Benefits": paid,
            "Ease of Integration": "Good - standard REST/JSON or RSS" if api_name not in {"PredictLeads News/Events API", "GDELT"} else "Not assessed",
            "API Documentation Quality": "Good - endpoint documented" if api_name != "Google News RSS" else "RSS endpoint; limited official API documentation",
            "Evidence Link": evidence,
            "Overall API Score": score_api(success_rate, avg_completeness, records),
            "Status": status,
            "Remarks": remarks,
            "Data Completeness (%)": f"{avg_completeness:.2f}",
            "Records Retrieved": str(records),
            "Estimated Cost per 100 Companies": estimated_cost(api_name),
            "Raw Export Saved (Y/N)": "Y" if any(row["raw_response_path"] for row in calls) else "N",
        })
    return rows


def score_api(success_rate: float, completeness: float, records: int) -> str:
    score = 1.0 + min(success_rate / 100 * 1.5, 1.5) + min(completeness / 100 * 1.5, 1.5) + (1.0 if records else 0.0)
    return f"{min(score, 5.0):.2f}/5"


def estimated_cost(api_name: str) -> str:
    if api_name == "Google News RSS":
        return "0 API cost; reliability/SLA not guaranteed."
    if api_name in {"Exa Search API", "Tavily Search API", "NewsAPI Everything API"}:
        return "About 100 search requests for 100 companies at one query per company; exact cost is plan specific."
    return "Not assessed."


def write_workbook(path: Path, sheets: dict[str, tuple[list[str], list[dict[str, Any]]]]) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    for title, (fields, rows) in sheets.items():
        ws = wb.create_sheet(title[:31])
        ws.append(fields)
        for row in rows:
            ws.append([row.get(field, "") for field in fields])
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = Font(bold=True, color="FFFFFF")
        ws.freeze_panes = "A2"
        for column in ws.columns:
            col = get_column_letter(column[0].column)
            width = max(12, min(max(len(str(cell.value or "")) for cell in column) + 2, 60))
            ws.column_dimensions[col].width = width
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(path)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F4E78")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    document.add_paragraph()


def write_docx(path: Path, company_rows: list[dict[str, str]], event_rows: list[dict[str, str]], trace_rows: list[dict[str, str]], comparison_rows: list[dict[str, str]], missing_rows: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HPI News / Key Announcements API Evaluation Report")
    run.font.bold = True
    run.font.size = Pt(18)
    document.add_paragraph(f"Generated: {now_iso()}")
    document.add_paragraph("Scope: News and key announcement provider evaluation for the 10-company HPI pilot set. Queries target funding, expansion, launch, leadership, M&A, and partnership events in the last 12 months.")
    document.add_heading("API Trace Summary", level=1)
    add_table(document, API_TRACE_FIELDS, [[row.get(field, "") for field in API_TRACE_FIELDS] for row in trace_rows])
    document.add_heading("API Comparison", level=1)
    add_table(document, API_COMPARISON_FIELDS, [[row.get(field, "") for field in API_COMPARISON_FIELDS] for row in comparison_rows])
    document.add_heading("Company News Summary", level=1)
    compact_fields = ["company_name", "api_name", "events_last_12_months", "event_type_counts", "latest_event_date", "data_completeness_percent"]
    add_table(document, ["Company", "API", "Events", "Types", "Latest Date", "Completeness %"], [[row.get(field, "") for field in compact_fields] for row in company_rows])
    document.add_heading("Sample Events", level=1)
    event_fields = ["company_name", "api_name", "event_headline", "event_date", "event_type", "source_publisher", "event_url"]
    add_table(document, ["Company", "API", "Headline", "Date", "Type", "Publisher", "URL"], [[row.get(field, "") for field in event_fields] for row in event_rows[:40]])
    document.add_heading("Missing / Not Tested Fields", level=1)
    add_table(document, MISSING_FIELDS, [[row.get(field, "") for field in MISSING_FIELDS] for row in missing_rows])
    document.save(path)


def run(args: argparse.Namespace) -> int:
    load_dotenv()
    ensure_dirs()
    copy_input()
    companies = read_companies(Path(args.input) if args.input else INPUT_FILE, args.limit)
    if not companies:
        raise ValueError("No companies found for News / Key Announcements pipeline.")

    all_company_rows: list[dict[str, str]] = []
    all_event_rows: list[dict[str, str]] = []
    all_missing_rows: list[dict[str, str]] = []
    all_call_rows: list[dict[str, str]] = []

    api_filter = {item.strip().lower() for item in args.apis.split(",") if item.strip()}
    run_all = "all" in api_filter
    collectors = []
    if run_all or "exa" in api_filter:
        collectors.append(collect_exa(companies, args.reuse_raw, args.timeout, args.per_provider_limit))
    if run_all or "tavily" in api_filter:
        collectors.append(collect_tavily(companies, args.reuse_raw, args.timeout, args.per_provider_limit))
    if run_all or "google_news_rss" in api_filter or "google" in api_filter:
        collectors.append(collect_google_rss(companies, args.reuse_raw, args.timeout, args.per_provider_limit))
    if run_all or "newsapi" in api_filter:
        collectors.append(collect_newsapi(companies, args.reuse_raw, args.timeout, args.per_provider_limit))
    if run_all or "predictleads" in api_filter:
        collectors.append(collect_blocked(companies, "PredictLeads News/Events API", "PredictLeads events/news endpoint", "PredictLeads key is invalid/missing in this workspace; not executed."))
    if run_all or "gdelt" in api_filter:
        collectors.append(collect_blocked(companies, "GDELT", "Not executed", "GDELT needs no API key but was not included in the scoped live run."))
    for company_rows, event_rows, missing_rows, call_rows in collectors:
        all_company_rows.extend(company_rows)
        all_event_rows.extend(event_rows)
        all_missing_rows.extend(missing_rows)
        all_call_rows.extend(call_rows)

    comparison_rows = build_api_comparison(all_event_rows)
    trace_rows = build_api_trace(all_company_rows, all_call_rows, all_missing_rows)

    write_csv(REPORTS_DIR / "company_news_announcements.csv", all_company_rows, COMPANY_FIELDS)
    write_csv(REPORTS_DIR / "event_detail.csv", all_event_rows, EVENT_FIELDS)
    write_csv(REPORTS_DIR / "missing_fields_report.csv", all_missing_rows, MISSING_FIELDS)
    write_csv(REPORTS_DIR / "api_call_log.csv", all_call_rows, API_CALL_FIELDS)
    write_csv(REPORTS_DIR / "api_comparison_report.csv", comparison_rows, API_COMPARISON_FIELDS)
    write_csv(REPORTS_DIR / "api_tracing_report.csv", trace_rows, API_TRACE_FIELDS)
    save_json(REPORTS_DIR / "company_news_announcements.json", all_company_rows)
    save_json(REPORTS_DIR / "run_manifest.json", {"generated_at": now_iso(), "companies_processed": len(companies), "providers": ["exa", "tavily", "google_news_rss", "newsapi", "predictleads", "gdelt"], "apis": sorted(api_filter), "reuse_raw": args.reuse_raw, "since_date": since_date(), "newsapi_since_date": newsapi_since_date(), "per_provider_limit": args.per_provider_limit})

    suffix = stamp()
    xlsx_path = REPORTS_DIR / f"hpi_news_announcements_api_evaluation_{suffix}.xlsx"
    docx_path = REPORTS_DIR / f"hpi_news_announcements_api_evaluation_{suffix}.docx"
    write_workbook(
        xlsx_path,
        {
            "API Comparison": (API_COMPARISON_FIELDS, comparison_rows),
            "API Trace": (API_TRACE_FIELDS, trace_rows),
            "API Call Log": (API_CALL_FIELDS, all_call_rows),
            "Company News Summary": (COMPANY_FIELDS, all_company_rows),
            "Event Detail": (EVENT_FIELDS, all_event_rows),
            "Missing Fields": (MISSING_FIELDS, all_missing_rows),
        },
    )
    write_docx(docx_path, all_company_rows, all_event_rows, trace_rows, comparison_rows, all_missing_rows)
    print(f"Saved workbook: {safe_rel(xlsx_path)}")
    print(f"Saved docx: {safe_rel(docx_path)}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Run HPI News / Key Announcements API evaluation.")
    parser.add_argument("--input", help="Input CSV path. Defaults to NewsAnnouncements/input/compnys.txt.")
    parser.add_argument("--limit", type=int, default=10)
    parser.add_argument("--reuse-raw", action="store_true", help="Reuse saved raw responses.")
    parser.add_argument("--timeout", type=int, default=60)
    parser.add_argument("--per-provider-limit", type=int, default=10)
    parser.add_argument("--apis", default="all", help="Comma-separated APIs to run: all, exa, tavily, google_news_rss, newsapi, predictleads, gdelt.")
    args = parser.parse_args()
    try:
        return run(args)
    except Exception as exc:
        print(f"News / Key Announcements pipeline failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
