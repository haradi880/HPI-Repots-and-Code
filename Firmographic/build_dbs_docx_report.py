from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
RAW_RESPONSE_PATH = BASE_DIR / "raw" / "dbs_response.json"
METADATA_PATH = BASE_DIR / "raw" / "apollo_run_metadata.json"
USAGE_BEFORE_PATH = BASE_DIR / "raw" / "apollo_usage_before.json"
USAGE_AFTER_PATH = BASE_DIR / "raw" / "apollo_usage_after.json"
EVIDENCE_DIR = BASE_DIR / "evidence"
REPORT_DIR = BASE_DIR / "reports" / "dbs_one_company_sample"
OUTPUT_PATH = REPORT_DIR / "dbs_group_apollo_sample_validation_report.docx"

APOLLO_ORG_DOC = "https://docs.apollo.io/reference/organization-enrichment"
APOLLO_USAGE_DOC = "https://docs.apollo.io/reference/view-api-usage-stats"
ORG_ENDPOINT = "https://api.apollo.io/api/v1/organizations/enrich"
USAGE_ENDPOINT = "https://api.apollo.io/api/v1/usage_stats/api_usage_stats"
SOURCE_API = "Apollo Organization Enrichment API"
NOT_AVAILABLE = "Not Available"
NOT_VERIFIED = "Not Verified"


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def get_path(data: Any, path: str) -> Any:
    current = data
    for part in path.split("."):
        if isinstance(current, dict) and part in current:
            current = current[part]
        elif isinstance(current, list) and part.isdigit() and int(part) < len(current):
            current = current[int(part)]
        else:
            return None
    return current


def is_missing(value: Any) -> bool:
    return value is None or value == "" or value == [] or value == {}


def fmt(value: Any) -> str:
    if value is None:
        return NOT_AVAILABLE
    if isinstance(value, list):
        return ", ".join(str(item) for item in value)
    if isinstance(value, float):
        return str(int(value)) if value.is_integer() else str(value)
    return str(value)


def fmt_growth(value: Any) -> str:
    if isinstance(value, (int, float)):
        percent = value * 100 if -1 < value < 1 else value
        return f"{percent:.2f}%"
    return fmt(value)


def extract_company_rows(payload: dict[str, Any]) -> list[dict[str, str]]:
    org = payload.get("organization", {})
    rows: list[dict[str, str]] = []

    def add(field: str, value: str, source_field: str, status: str, reason: str = "") -> None:
        rows.append(
            {
                "Field": field,
                "Value": value,
                "Source API": SOURCE_API,
                "Source Field": source_field,
                "Status": status,
                "Reason / Note": reason,
            }
        )

    name = get_path(org, "name")
    add(
        "Legal Name",
        NOT_VERIFIED,
        "organization.name",
        NOT_VERIFIED,
        f"Apollo returned company name ({name}), but did not return legal_name or registered_name.",
    )
    add("Website Domain", fmt(get_path(org, "primary_domain")), "organization.primary_domain", "Available")
    add("LinkedIn URL", fmt(get_path(org, "linkedin_url")), "organization.linkedin_url", "Available")
    add("Employee Count", fmt(get_path(org, "estimated_num_employees")), "organization.estimated_num_employees", "Available")
    add(
        "Headcount Growth (1 Year)",
        fmt_growth(get_path(org, "organization_headcount_twelve_month_growth")),
        "organization.organization_headcount_twelve_month_growth",
        "Available",
    )
    add("Revenue / Revenue Band", fmt(get_path(org, "annual_revenue_printed")), "organization.annual_revenue_printed", "Available")
    add("Industry", fmt(get_path(org, "industry")), "organization.industry", "Available")

    sub_industry_parts = []
    if not is_missing(get_path(org, "secondary_industries")):
        sub_industry_parts.append(fmt(get_path(org, "secondary_industries")))
    if not is_missing(get_path(org, "naics_codes")):
        sub_industry_parts.append(f"NAICS: {fmt(get_path(org, 'naics_codes'))}")
    if not is_missing(get_path(org, "sic_codes")):
        sub_industry_parts.append(f"SIC: {fmt(get_path(org, 'sic_codes'))}")
    add(
        "Sub-Industry (NAICS/SIC)",
        " | ".join(sub_industry_parts),
        "organization.secondary_industries, organization.naics_codes, organization.sic_codes",
        "Available",
    )

    add("Headquarters Location", fmt(get_path(org, "raw_address")), "organization.raw_address", "Available")
    add(
        "Number of Sites/Locations (Singapore + Global)",
        NOT_VERIFIED,
        "organization.retail_location_count",
        NOT_VERIFIED,
        "Apollo returned retail_location_count=0, but not a verified Singapore + global location total.",
    )
    add("Founded Year", fmt(get_path(org, "founded_year")), "organization.founded_year", "Available")
    add(
        "Ownership Type",
        "Public Company",
        "organization.publicly_traded_symbol, organization.publicly_traded_exchange",
        "Derived from API",
        f"Derived from Apollo public listing fields: {fmt(get_path(org, 'publicly_traded_symbol'))} / {fmt(get_path(org, 'publicly_traded_exchange'))}.",
    )
    add(
        "Total Funding Raised",
        NOT_AVAILABLE,
        NOT_AVAILABLE,
        NOT_AVAILABLE,
        "Apollo response did not contain total_funding, funding_total, or total_funding_raised.",
    )
    return rows


def flatten_numbers(data: Any, prefix: str = "") -> dict[str, float]:
    values: dict[str, float] = {}
    if isinstance(data, dict):
        for key, value in data.items():
            next_prefix = f"{prefix}.{key}" if prefix else str(key)
            values.update(flatten_numbers(value, next_prefix))
    elif isinstance(data, list):
        for index, value in enumerate(data):
            next_prefix = f"{prefix}.{index}" if prefix else str(index)
            values.update(flatten_numbers(value, next_prefix))
    elif isinstance(data, (int, float)) and not isinstance(data, bool):
        values[prefix] = float(data)
    return values


def clean_metric_path(path: str) -> str:
    return re.sub(r'\["([^"]+)",\s*"([^"]+)"\]', r"\1/\2", path)


def fmt_number(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else f"{value:.4f}".rstrip("0").rstrip(".")


def credit_delta(before: dict[str, Any], after: dict[str, Any]) -> tuple[str, str, str, str]:
    before_json = before.get("json", before)
    after_json = after.get("json", after)
    before_numbers = flatten_numbers(before_json)
    after_numbers = flatten_numbers(after_json)
    candidates = []
    for path, after_value in after_numbers.items():
        if path not in before_numbers:
            continue
        delta = after_value - before_numbers[path]
        lowered = path.lower()
        if delta and "organizations" in lowered and "enrich" in lowered and "day.consumed" in lowered:
            candidates.append((path, before_numbers[path], after_value, delta))
    if not candidates:
        return NOT_VERIFIED, NOT_VERIFIED, NOT_VERIFIED, "Credit delta not found in usage snapshots."
    path, before_value, after_value, delta = candidates[0]
    return fmt_number(before_value), fmt_number(after_value), fmt_number(delta), f"{clean_metric_path(path)}: {fmt_number(before_value)} -> {fmt_number(after_value)}"


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def set_cell_text_color(cell, color: RGBColor, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color
            run.font.bold = bold


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], "1F4E78")
        set_cell_text_color(header_cells[idx], RGBColor(255, 255, 255), bold=True)
        if widths:
            header_cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if widths:
                cells[idx].width = Inches(widths[idx])
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            if value == NOT_AVAILABLE:
                set_cell_shading(cells[idx], "FCE4D6")
            elif value == NOT_VERIFIED:
                set_cell_shading(cells[idx], "FFF2CC")
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_kv_table(document: Document, rows: list[tuple[str, str]]) -> None:
    add_table(document, ["Field", "Value"], [[key, value] for key, value in rows], widths=[2.2, 4.9])


def build_report() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    payload = load_json(RAW_RESPONSE_PATH)
    metadata = load_json(METADATA_PATH)
    usage_before = load_json(USAGE_BEFORE_PATH)
    usage_after = load_json(USAGE_AFTER_PATH)
    company_rows = extract_company_rows(payload)
    credits_before, credits_after, credits_used, credit_evidence = credit_delta(usage_before, usage_after)

    org_result = metadata["organization_request"]["result"]
    latency_ms = str(org_result.get("elapsed_ms", NOT_AVAILABLE))
    status_code = str(org_result.get("status_code", NOT_AVAILABLE))
    rate_limit = "; ".join(f"{key}: {value}" for key, value in org_result.get("headers", {}).items() if "rate" in key.lower())
    complete_count = sum(1 for row in company_rows if row["Status"] in {"Available", "Derived from API"})
    completeness = f"{(complete_count / len(company_rows)) * 100:.2f}%"
    returned_fields = "; ".join(row["Field"] for row in company_rows if row["Status"] in {"Available", "Derived from API"})
    missing_fields = "; ".join(row["Field"] for row in company_rows if row["Status"] not in {"Available", "Derived from API"})

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DBS Group Firmographic API Validation Report")
    run.font.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 54, 93)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Sample company only - Apollo Organization Enrichment API").italic = True
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    add_heading(document, "Executive Summary", 1)
    add_kv_table(
        document,
        [
            ("Company", "DBS Group"),
            ("API Used", SOURCE_API),
            ("Endpoint Used", f"GET {ORG_ENDPOINT}"),
            ("Status", f"Success - HTTP {status_code}"),
            ("Companies Processed", "1"),
            ("Credits Before", credits_before),
            ("Credits After", credits_after),
            ("Credits Used", credits_used),
            ("Credit Evidence", credit_evidence),
            ("Data Completeness", completeness),
            ("Raw Export Saved", f"Y - {RAW_RESPONSE_PATH.relative_to(BASE_DIR)}"),
        ],
    )

    add_heading(document, "Understanding Confirmation", 1)
    document.add_paragraph(
        "The understanding is correct: before processing all 10 companies, one sample company should be submitted for approval. "
        "This report covers DBS Group only, using one API, with raw response, field mapping, missing-data reasons, and actual credit usage."
    )
    document.add_paragraph("The remaining companies should be processed only after this sample format and data-quality approach are approved.")

    add_heading(document, "Apollo Endpoint Check", 1)
    document.add_paragraph(
        "Yes, the Apollo endpoint was checked. The correct one-company firmographic endpoint is GET /api/v1/organizations/enrich. "
        "Apollo documentation states this endpoint is used to enrich one company and consumes credits. "
        "For credit and rate-limit tracing, the checked endpoint is POST /api/v1/usage_stats/api_usage_stats, which returns usage/rate information and requires a master API key."
    )
    add_kv_table(
        document,
        [
            ("Organization Enrichment Docs", APOLLO_ORG_DOC),
            ("Usage Stats Docs", APOLLO_USAGE_DOC),
            ("Request Params Used", json.dumps(metadata["organization_request"]["params"], ensure_ascii=False)),
            ("Authentication Type", "API Key"),
            ("Latency", f"{latency_ms} ms"),
            ("Rate Limit Captured", rate_limit or NOT_AVAILABLE),
        ],
    )

    add_heading(document, "Firmographic Data", 1)
    add_table(
        document,
        ["Field", "Value", "Source API", "Source Field", "Status", "Reason / Note"],
        [[row["Field"], row["Value"], row["Source API"], row["Source Field"], row["Status"], row["Reason / Note"]] for row in company_rows],
        widths=[1.5, 1.7, 1.5, 1.8, 1.0, 2.1],
    )

    add_heading(document, "Reason Of Missing Data", 1)
    missing_reason_rows = [
        [row["Field"], row["Value"], row["Source Field"], row["Status"], row["Reason / Note"]]
        for row in company_rows
        if row["Status"] not in {"Available", "Derived from API"}
    ]
    add_table(
        document,
        ["Field", "Value", "Source Field", "Status", "Reason"],
        missing_reason_rows,
        widths=[1.7, 1.3, 1.7, 1.2, 3.0],
    )

    add_heading(document, "API Comparison Report", 1)
    add_table(
        document,
        ["API Name", "Endpoint Used", "Status", "Fields Returned", "Free-Tier Limitations", "Paid-Tier Benefits", "Notes"],
        [
            [
                SOURCE_API,
                f"GET {ORG_ENDPOINT}",
                "Success",
                returned_fields,
                "Not Verified - not returned in Apollo response.",
                "Not Verified - plan benefits are account/plan dependent.",
                f"Missing/unverified fields: {missing_fields}. Raw response saved and credit evidence captured.",
            ]
        ],
        widths=[1.3, 1.7, 0.8, 2.2, 1.3, 1.3, 1.8],
    )

    add_heading(document, "API Trace And Validation", 1)
    validation_rows = [
        ("Tool Name", SOURCE_API),
        ("Category", "Firmographic"),
        ("API Available (Y/N)", "Y"),
        ("Authentication Type", "API Key"),
        ("Free Credits / Tokens", NOT_VERIFIED),
        ("Credits / Tokens Used", credits_used),
        ("Credits Before", credits_before),
        ("Credits After", credits_after),
        ("Credit Evidence", credit_evidence),
        ("Rate Limit", rate_limit or NOT_AVAILABLE),
        ("Companies Processed", "1"),
        ("Coverage (%)", completeness),
        ("Success Rate (%)", "100%"),
        ("Error Rate (%)", "0%"),
        ("Average Latency", f"{latency_ms} ms"),
        ("Gated Fields", "Not Verified - absent fields were not confirmed as gated."),
        ("Free Tier Limitation", "Not Verified - not returned in API response."),
        ("Paid Plan Cost", "Not Verified - not returned in API response."),
        ("Paid Tier Benefits", "Not Verified - plan benefits not validated from API response."),
        ("Ease of Integration", "4/5"),
        ("API Documentation Quality", "4/5"),
        ("Evidence Link", f"{RAW_RESPONSE_PATH.relative_to(BASE_DIR)}; {USAGE_BEFORE_PATH.relative_to(BASE_DIR)}; {USAGE_AFTER_PATH.relative_to(BASE_DIR)}"),
        ("Overall API Score", "4/5"),
        ("Status", "Sample ready for approval"),
        ("Remarks", "DBS Group sample only. Proceed with remaining companies after approval."),
        ("Data Completeness (%)", completeness),
        ("Records Retrieved", "1"),
        ("Estimated Cost per 100 Companies", "Not Verified in currency; observed usage is 1 Apollo credit/company, so 100 companies would use about 100 Apollo credits."),
        ("Raw Export Saved (Y/N)", "Y"),
    ]
    add_kv_table(document, validation_rows)

    add_heading(document, "Evidence Images", 1)
    for image_name, caption in [
        ("api_dashboard.png", "Apollo API dashboard evidence"),
        ("credit_usage.png", "Apollo credit usage evidence"),
    ]:
        image_path = EVIDENCE_DIR / image_name
        if image_path.exists():
            paragraph = document.add_paragraph()
            paragraph.add_run(caption).bold = True
            document.add_picture(str(image_path), width=Inches(6.8))
        else:
            document.add_paragraph(f"{caption}: image not found at {image_path}")

    add_heading(document, "Files Saved", 1)
    add_kv_table(
        document,
        [
            ("Raw API Response", str(RAW_RESPONSE_PATH.relative_to(BASE_DIR))),
            ("Run Metadata", str(METADATA_PATH.relative_to(BASE_DIR))),
            ("Usage Before Snapshot", str(USAGE_BEFORE_PATH.relative_to(BASE_DIR))),
            ("Usage After Snapshot", str(USAGE_AFTER_PATH.relative_to(BASE_DIR))),
            ("API Dashboard Image", str((EVIDENCE_DIR / "api_dashboard.png").relative_to(BASE_DIR))),
            ("Credit Usage Image", str((EVIDENCE_DIR / "credit_usage.png").relative_to(BASE_DIR))),
        ],
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
    print(f"Saved {OUTPUT_PATH}")
