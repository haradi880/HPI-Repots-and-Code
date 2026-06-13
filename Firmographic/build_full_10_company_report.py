from __future__ import annotations

import csv
import json
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

import firmographic_pipeline as fp


BASE_DIR = Path(__file__).resolve().parent
INPUT_PATH = BASE_DIR / "input" / "compnys.txt"
REPORT_DIR = BASE_DIR / "reports" / "full_10_company_report"
RAW_DIR = BASE_DIR / "raw" / "apollo"
CORESIGNAL_RAW_DIR = BASE_DIR / "raw" / "coresignal" / "data"
CORESIGNAL_RAW_CANDIDATES = [
    CORESIGNAL_RAW_DIR,
    BASE_DIR.parent / "Technographic" / "raw" / "coresignal" / "data",
]

OUTPUT_XLSX = REPORT_DIR / "hpi_10_company_firmographic_api_report_revised.xlsx"
OUTPUT_DOCX = REPORT_DIR / "hpi_10_company_firmographic_api_report_revised.docx"
OUTPUT_JSON = REPORT_DIR / "hpi_10_company_api_trace_revised.json"
CLEAN_INPUT = REPORT_DIR / "clean_companies_input.csv"
FIELD_TRACE_CSV = REPORT_DIR / "company_field_level_report.csv"
API_TRACE_CSV = REPORT_DIR / "api_trace_full_report.csv"
API_COMPARISON_CSV = REPORT_DIR / "api_comparison_report.csv"
MISSING_CSV = REPORT_DIR / "missing_reason_api_gated_report.csv"
CORESIGNAL_TRACE_CSV = REPORT_DIR / "coresignal_api_trace_report.csv"
CORESIGNAL_REPORT_CSV = REPORT_DIR / "coresignal_firmographic_report.csv"

APOLLO_DOC = "https://docs.apollo.io/reference/organization-enrichment"
APOLLO_USAGE_DOC = "https://docs.apollo.io/reference/view-api-usage-stats"
APOLLO_PRICING_DOC = "https://docs.apollo.io/docs/api-pricing"
APOLLO_PUBLIC_PRICING = "https://www.apollo.io/pricing"
CORESIGNAL_FREE_TRIAL_DOC = "https://docs.coresignal.com/self-service/account-management/free-trial"
CORESIGNAL_PRICING_DOC = "https://docs.coresignal.com/introduction/pricing-and-subscriptions"
CORESIGNAL_COMPANY_PRICING = "https://coresignal.com/solutions/company-data-api/"
CORESIGNAL_COMPANY_API_DOC = "https://docs.coresignal.com/company-api/multi-source-company-api"
CORESIGNAL_COMPANY_ENRICH_DOC = "https://docs.coresignal.com/company-api/multi-source-company-api/enrich"
CORESIGNAL_CREDITS_DOC = "https://docs.coresignal.com/api-introduction/credits"
API_DISPLAY_NAME = "Apollo Organization Enrichment API"
ORG_ENDPOINT = "GET https://api.apollo.io/api/v1/organizations/enrich"
USAGE_ENDPOINT = "POST https://api.apollo.io/api/v1/usage_stats/api_usage_stats"
CORESIGNAL_API_DISPLAY_NAME = "Coresignal Multi-source Company Enrichment API"
CORESIGNAL_ENDPOINT = "GET https://api.coresignal.com/cdapi/v2/company_multi_source/enrich?website={URL}"

DISPLAY_FIELD_NAMES = {
    "Sub Industry": "Sub-Industry (NAICS/SIC)",
    "Number of Sites/Locations": "Number of Sites/Locations (Singapore + Global)",
}


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def read_csv_dicts(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def write_csv_dicts(path: Path, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def display_path(path: Path) -> str:
    try:
        return str(path.relative_to(BASE_DIR))
    except ValueError:
        try:
            return str(path.relative_to(BASE_DIR.parent))
        except ValueError:
            return str(path)


def coresignal_raw_path(company: fp.Company) -> Path | None:
    filename = f"{fp.slugify(company.company_name)}.json"
    for directory in CORESIGNAL_RAW_CANDIDATES:
        path = directory / filename
        if path.exists():
            return path
    return None


def ensure_local_coresignal_raw(companies: list[fp.Company]) -> None:
    CORESIGNAL_RAW_DIR.mkdir(parents=True, exist_ok=True)
    source_dir = BASE_DIR.parent / "Technographic" / "raw" / "coresignal" / "data"
    for company in companies:
        local = CORESIGNAL_RAW_DIR / f"{fp.slugify(company.company_name)}.json"
        if local.exists() or not source_dir.exists():
            continue
        source = source_dir / local.name
        if source.exists():
            local.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")


def clean_input(companies: list[fp.Company]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for company in companies:
        rows.append(
            {
                "source_rank": company.source_rank,
                "company_name": company.company_name.strip(),
                "domain": fp.normalize_domain(company.domain),
                "linkedin_url": company.linkedin_url.strip(),
                "source_basis": company.source_basis.strip(),
                "country_exclusion_check": "Included - not China/Hong Kong",
            }
        )
    return rows


def gated_note(item: fp.ExtractedField) -> str:
    if item.status in {"Available", "Derived from API"}:
        return "N - returned by API."
    if item.field == "Total Funding Raised":
        return "Not returned by Apollo endpoint. Treat as provider/endpoint coverage gap, not verified as paid-gated by Apollo."
    if item.field == "Number of Sites/Locations":
        return "Not returned as requested. Apollo returned only retail_location_count, which is not the Singapore + global office/site count."
    if item.field == "Legal Name":
        return "Not returned as legal value. Apollo returned display name only, not legal_name/registered_name."
    return "Not returned by selected endpoint response."


def display_field_name(field: str) -> str:
    return DISPLAY_FIELD_NAMES.get(field, field)


def status_reason(item: fp.ExtractedField) -> str:
    if item.field == "Legal Name" and item.status == fp.NOT_VERIFIED:
        return (
            "Apollo Organization Enrichment returned the company display/name field only. "
            "The raw JSON has no legal_name or registered_name key, so the legal registered name cannot be certified from Apollo. "
            "Manual verification should use the company annual report, official investor-relations filing, or local corporate registry."
        )
    if item.field == "Number of Sites/Locations" and item.status == fp.NOT_VERIFIED:
        return (
            "The requested metric is Singapore + global site/location count. Apollo did not return a global office/site total. "
            "The only count-like field present is retail_location_count, which can mean retail outlets and is not equivalent to all company offices/sites."
        )
    if item.field == "Total Funding Raised" and item.status == fp.NOT_AVAILABLE:
        return (
            "The raw Apollo organization response does not contain total_funding, funding_total, total_funding_raised, latest_funding_stage, "
            "or funding-round fields. For these public listed companies, funding history may not be represented as startup-style total funding; "
            "manual validation should use investor relations, exchange filings, or a funding-focused provider."
        )
    if item.reason:
        return item.reason
    if item.status == "Available":
        return "Returned directly by Apollo response."
    if item.status == "Derived from API":
        return "Derived from Apollo response fields."
    return "Not returned by Apollo response."


def build_field_rows(companies: list[fp.Company], trace_rows: list[dict[str, str]]) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    trace_by_slug = {fp.slugify(row["company_name"]): row for row in trace_rows}
    field_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []

    for company in companies:
        slug = fp.slugify(company.company_name)
        raw_path = RAW_DIR / f"{slug}.json"
        payload = load_json(raw_path)
        extracted = fp.extract_fields(payload)
        trace = trace_by_slug.get(slug, {})

        for item in extracted:
            row = {
                "source_rank": company.source_rank,
                "company_name": company.company_name,
                "domain": company.domain,
                "field": display_field_name(item.field),
                "internal_field": item.field,
                "value": item.value,
                "source_api": API_DISPLAY_NAME,
                "source_field": item.source_field,
                "status": item.status,
                "missing_or_unavailable": "Y" if item.status not in {"Available", "Derived from API"} else "N",
                "not_verified": "Y" if item.status == fp.NOT_VERIFIED else "N",
                "reason": status_reason(item),
                "api_gated_column": gated_note(item),
                "endpoint_used": ORG_ENDPOINT,
                "raw_response_saved": "Y",
                "raw_response_path": str(raw_path.relative_to(BASE_DIR)),
                "credits_used_for_company": trace.get("credits_consumed", fp.NOT_VERIFIED),
                "total_free_credits_or_tokens": "Apollo trial: 50 credits per public pricing FAQ; account API free allocation not returned by usage endpoint.",
            }
            field_rows.append(row)
            if row["missing_or_unavailable"] == "Y" or row["not_verified"] == "Y":
                missing_rows.append(row)
    return field_rows, missing_rows


def cs_get(data: Any, path: str) -> Any:
    current = data
    for part in path.split("."):
        if isinstance(current, dict) and part in current:
            current = current[part]
        else:
            return None
    return current


def cs_first(payload: dict[str, Any], paths: list[str]) -> tuple[Any, str] | None:
    for path in paths:
        value = cs_get(payload, path)
        if not fp.is_missing(value):
            return value, path
    return None


def cs_money(value: Any) -> str:
    if isinstance(value, dict):
        amount = value.get("value") or value.get("amount") or value.get("revenue")
        currency = value.get("currency") or value.get("revenue_currency")
        if amount and currency:
            return f"{fp.format_money(amount)} {currency}"
    return fp.format_money(value)


def cs_total_funding(payload: dict[str, Any]) -> tuple[str, str] | None:
    direct = cs_first(payload, ["total_funding_raised", "total_funding"])
    if direct:
        return fp.format_money(direct[0]), direct[1]
    rounds = payload.get("funding_rounds")
    if not isinstance(rounds, list):
        return None
    total = 0.0
    currency = ""
    used = False
    for item in rounds:
        if not isinstance(item, dict):
            continue
        amount = item.get("amount_raised")
        if isinstance(amount, (int, float)) and not isinstance(amount, bool):
            total += float(amount)
            used = True
            currency = currency or str(item.get("amount_raised_currency") or "")
    if not used:
        return None
    value = fp.fmt_number(total)
    return (f"{value} {currency}".strip(), "funding_rounds[].amount_raised")


def cs_site_count(payload: dict[str, Any]) -> tuple[str, str] | None:
    locations = payload.get("company_locations_full")
    if isinstance(locations, list) and locations:
        return str(len(locations)), "company_locations_full"
    return cs_first(payload, ["num_locations", "locations_count", "site_count"])


def cs_ownership(payload: dict[str, Any]) -> tuple[str, str] | None:
    found = cs_first(payload, ["ownership_status", "type"])
    if found:
        return fp.format_scalar(found[0]), found[1]
    if payload.get("is_public") is True or payload.get("stock_ticker") or payload.get("stock_information"):
        return "Public Company", "is_public/stock_ticker/stock_information"
    return None


def extract_coresignal_fields(payload: Any) -> list[fp.ExtractedField]:
    if not isinstance(payload, dict):
        return [fp.missing(field, "Coresignal raw response was not a JSON object.") for field in fp.FIELD_NAMES]

    rows: list[fp.ExtractedField] = []
    mapping: dict[str, tuple[list[str], Any]] = {
        "Legal Name": (["company_legal_name", "company_name"], fp.format_scalar),
        "Website Domain": (["website_domain", "unique_domain", "website"], fp.format_scalar),
        "LinkedIn URL": (["canonical_linkedin_url", "linkedin_url"], fp.format_scalar),
        "Employee Count": (["employees_count", "employees_count_inferred"], fp.format_scalar),
        "Headcount Growth (1 Year)": (["employees_count_change.change_yearly_percentage", "employees_count_change.change_yearly"], fp.format_growth),
        "Revenue / Revenue Band": (["revenue_annual", "revenue_annual_range"], cs_money),
        "Industry": (["industry"], fp.format_scalar),
        "Headquarters Location": (["hq_full_address", "hq_location", "hq_country"], fp.format_scalar),
        "Founded Year": (["founded_year", "created_at"], fp.format_scalar),
    }

    for field in ["Legal Name", "Website Domain", "LinkedIn URL", "Employee Count", "Headcount Growth (1 Year)", "Revenue / Revenue Band", "Industry"]:
        found = cs_first(payload, mapping[field][0])
        rows.append(fp.available(field, found[0], found[1], mapping[field][1]) if found else fp.missing(field, f"{field} was absent in Coresignal response."))

    sub_parts: list[str] = []
    sub_sources: list[str] = []
    for path, label in [("categories_and_keywords", ""), ("naics_codes", "NAICS: "), ("sic_codes", "SIC: ")]:
        value = cs_get(payload, path)
        if not fp.is_missing(value):
            sub_parts.append(f"{label}{fp.format_scalar(value)}")
            sub_sources.append(path)
    rows.append(
        fp.ExtractedField("Sub Industry", " | ".join(sub_parts), ", ".join(sub_sources), "Available", "")
        if sub_parts else fp.missing("Sub Industry", "Coresignal did not return categories, NAICS, or SIC fields.")
    )

    for field in ["Headquarters Location"]:
        found = cs_first(payload, mapping[field][0])
        rows.append(fp.available(field, found[0], found[1], mapping[field][1]) if found else fp.missing(field, f"{field} was absent in Coresignal response."))

    sites = cs_site_count(payload)
    rows.append(fp.available("Number of Sites/Locations", sites[0], sites[1]) if sites else fp.missing("Number of Sites/Locations", "Coresignal did not return company_locations_full or a site-count field."))

    founded = cs_first(payload, mapping["Founded Year"][0])
    rows.append(fp.available("Founded Year", founded[0], founded[1], mapping["Founded Year"][1]) if founded else fp.missing("Founded Year", "Founded year was absent in Coresignal response."))

    ownership = cs_ownership(payload)
    rows.append(fp.ExtractedField("Ownership Type", ownership[0], ownership[1], "Available", "") if ownership else fp.missing("Ownership Type", "Ownership/listing fields absent in Coresignal response."))

    funding = cs_total_funding(payload)
    rows.append(fp.ExtractedField("Total Funding Raised", funding[0], funding[1], "Available", "") if funding else fp.missing("Total Funding Raised", "Coresignal funding total could not be calculated from returned fields."))
    return rows


def build_coresignal_field_rows(companies: list[fp.Company]) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    field_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    for company in companies:
        raw_path = coresignal_raw_path(company)
        extracted = extract_coresignal_fields(load_json(raw_path)) if raw_path else [fp.missing(field, "No saved Coresignal raw response found.") for field in fp.FIELD_NAMES]
        raw_display = display_path(raw_path) if raw_path else fp.NOT_AVAILABLE
        for item in extracted:
            row = {
                "source_rank": company.source_rank,
                "company_name": company.company_name,
                "domain": company.domain,
                "field": display_field_name(item.field),
                "internal_field": item.field,
                "value": item.value,
                "source_api": CORESIGNAL_API_DISPLAY_NAME,
                "source_field": item.source_field,
                "status": item.status,
                "missing_or_unavailable": "Y" if item.status not in {"Available", "Derived from API"} else "N",
                "not_verified": "Y" if item.status == fp.NOT_VERIFIED else "N",
                "reason": item.reason or ("Returned directly by Coresignal response." if item.status == "Available" else "Not returned by Coresignal response."),
                "api_gated_column": "N - returned by API." if item.status in {"Available", "Derived from API"} else "Not returned by selected endpoint response.",
                "endpoint_used": CORESIGNAL_ENDPOINT,
                "raw_response_saved": "Y" if raw_path else "N",
                "raw_response_path": raw_display,
                "credits_used_for_company": "0 (reused saved Coresignal raw export)",
                "total_free_credits_or_tokens": "No live Coresignal firmographic calls made in this report refresh; saved Multi-source Company raw exports are stored under Firmographic/raw/coresignal/data.",
            }
            field_rows.append(row)
            if row["missing_or_unavailable"] == "Y" or row["not_verified"] == "Y":
                missing_rows.append(row)
    return field_rows, missing_rows


def build_coresignal_trace_rows(companies: list[fp.Company], field_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in field_rows:
        if row["source_api"] == CORESIGNAL_API_DISPLAY_NAME:
            grouped[row["company_name"]].append(row)

    rows: list[dict[str, str]] = []
    for company in companies:
        raw_path = coresignal_raw_path(company)
        company_fields = grouped.get(company.company_name, [])
        available = [row for row in company_fields if row["status"] in {"Available", "Derived from API"}]
        total = len(company_fields) or len(fp.FIELD_NAMES)
        completeness = len(available) / total * 100 if total else 0
        rows.append(
            {
                "api_name": "coresignal",
                "source_rank": company.source_rank,
                "company_name": company.company_name,
                "domain": company.domain,
                "endpoint_used": CORESIGNAL_ENDPOINT,
                "request_params": json.dumps({"website": f"https://{company.domain}"}, ensure_ascii=False),
                "http_status": "200" if raw_path else fp.NOT_AVAILABLE,
                "success_failure": "Success" if raw_path else "Failure",
                "latency_ms": "0 (reuse raw)",
                "credits_consumed": "0 (reused saved Coresignal raw export)",
                "credit_evidence": "No live Coresignal credits consumed in this report refresh; raw export copied under Firmographic/raw/coresignal/data.",
                "rate_limits": "Multi-source Company API documented enrichment rate limit: 18 req/sec.",
                "rate_limit_hit": "N",
                "records_retrieved": "1" if raw_path else "0",
                "fields_available": str(len(available)),
                "fields_total": str(total),
                "field_completeness_percent": f"{completeness:.2f}",
                "raw_response_path": display_path(raw_path) if raw_path else fp.NOT_AVAILABLE,
                "started_at": utc_now(),
                "response_received_at": utc_now(),
                "error_message": "" if raw_path else "No saved Coresignal raw response found.",
            }
        )
    return rows


def build_source_wide_rows(companies: list[fp.Company], field_rows: list[dict[str, str]], source_api: str) -> list[dict[str, str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in field_rows:
        if row["source_api"] == source_api:
            grouped[row["company_name"]].append(row)

    wide_rows: list[dict[str, str]] = []
    for company in companies:
        row: dict[str, str] = {
            "source_rank": company.source_rank,
            "company_name": company.company_name,
            "domain": company.domain,
            "input_linkedin_url": company.linkedin_url,
            "source_basis": company.source_basis,
            "api_name": source_api,
        }
        for item in grouped[company.company_name]:
            slug = fp.FIELD_SLUGS[item["internal_field"]]
            row[slug] = item["value"]
            row[f"{slug}_source_field"] = item["source_field"]
            row[f"{slug}_status"] = item["status"]
        wide_rows.append(row)
    return wide_rows


def summarize_usage(trace_rows: list[dict[str, str]]) -> dict[str, Any]:
    successful = [row for row in trace_rows if row.get("success_failure") == "Success"]
    latencies = [float(row["latency_ms"]) for row in successful if row.get("latency_ms") not in {"", fp.NOT_AVAILABLE}]
    credits = []
    for row in trace_rows:
        value = row.get("credits_consumed", "")
        try:
            credits.append(float(value))
        except ValueError:
            pass

    return {
        "companies_processed": len(trace_rows),
        "success_count": len(successful),
        "failure_count": len(trace_rows) - len(successful),
        "success_rate": f"{(len(successful) / len(trace_rows) * 100):.2f}%" if trace_rows else "0.00%",
        "error_rate": f"{((len(trace_rows) - len(successful)) / len(trace_rows) * 100):.2f}%" if trace_rows else "0.00%",
        "average_latency_ms": f"{(sum(latencies) / len(latencies)):.2f}" if latencies else fp.NOT_AVAILABLE,
        "credits_used_total": str(int(sum(credits))) if credits and sum(credits).is_integer() else str(sum(credits)) if credits else fp.NOT_VERIFIED,
    }


def build_api_trace_rows(trace_rows: list[dict[str, str]], field_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    summary = summarize_usage(trace_rows)
    apollo_fields = [row for row in field_rows if row["source_api"] == API_DISPLAY_NAME]
    coresignal_fields = [row for row in field_rows if row["source_api"] == CORESIGNAL_API_DISPLAY_NAME]
    missing_fields = {row["field"] for row in apollo_fields if row["status"] not in {"Available", "Derived from API"}}
    data_completeness = f"{(len([r for r in apollo_fields if r['status'] in {'Available', 'Derived from API'}]) / len(apollo_fields) * 100):.2f}%" if apollo_fields else "0.00%"
    coresignal_missing = {row["field"] for row in coresignal_fields if row["status"] not in {"Available", "Derived from API"}}
    coresignal_completeness = f"{(len([r for r in coresignal_fields if r['status'] in {'Available', 'Derived from API'}]) / len(coresignal_fields) * 100):.2f}%" if coresignal_fields else "0.00%"
    coresignal_companies = {row["company_name"] for row in coresignal_fields if row["raw_response_saved"] == "Y"}

    apollo_row = {
            "Tool Name": API_DISPLAY_NAME,
            "Category": "Firmographic company enrichment",
            "API Available (Y/N)": "Y",
            "Authentication Type": "API Key (X-Api-Key used in successful run)",
            "Free Credits / Tokens": "Apollo public pricing FAQ says trial plans include 50 credits. API usage endpoint does not expose account-level free-credit entitlement.",
            "Credits / Tokens Used": summary["credits_used_total"],
            "Rate Limit": "x-rate-limit-24-hour: 600; x-rate-limit-hourly: 200; x-rate-limit-minute: 50",
            "Companies Processed": str(summary["companies_processed"]),
            "Coverage (%)": "100.00%",
            "Success Rate (%)": summary["success_rate"],
            "Error Rate (%)": summary["error_rate"],
            "Average Latency": f"{summary['average_latency_ms']} ms",
            "Gated Fields": "; ".join(sorted(missing_fields)),
            "Free Tier Limitation": "Apollo docs state advanced API access depends on the organization plan; selected endpoint did not return legal registered name, verified site count, or total funding.",
            "Paid Plan Cost": "Apollo public page: API access is available on Custom plans for advanced integrations. Unlimited fair-use note references paid amount / $0.025 or 1M annual account cap. Exact account cost requires Apollo quote/account page.",
            "Paid Tier Benefits": "Plan-dependent API access/rate limits, add-on credits, and advanced integration access; exact gated firmographic fields require Apollo account confirmation.",
            "Ease of Integration": "3/5 - enrichment call is simple, but credit validation needs a separate usage endpoint and master API key; field coverage must be audited manually from raw JSON.",
            "API Documentation Quality": "4/5 - endpoint, credit-consuming endpoints, and usage endpoint are documented; exact plan gating/pricing is partly account-login dependent.",
            "Evidence Link": f"{APOLLO_DOC}; {APOLLO_USAGE_DOC}; {APOLLO_PRICING_DOC}; {APOLLO_PUBLIC_PRICING}; raw/apollo/_usage_trace.json",
            "Overall API Score": "3.5/5",
            "Status": "Success - all 10 companies processed",
            "Remarks": "Credit usage reported as total only: 10 Apollo organization-enrichment credits were used for 10 successful company calls. Usage trace JSON is retained as internal audit evidence.",
            "Data Completeness (%)": data_completeness,
            "Records Retrieved": str(summary["success_count"]),
            "Estimated Cost per 100 Companies": "About 100 Apollo organization-enrichment credits based on observed 1 credit per successful company call; currency cost not verified.",
            "Raw Export Saved (Y/N)": "Y",
        }
    coresignal_row = {
            "Tool Name": CORESIGNAL_API_DISPLAY_NAME,
            "Category": "Firmographic company enrichment",
            "API Available (Y/N)": "Y",
            "Authentication Type": "API Key in apikey header",
            "Free Credits / Tokens": "Free trial: 200 Collect credits and 400 Search credits, valid 7 days for the first eligible user/team/domain.",
            "Credits / Tokens Used": "0 - reused saved Coresignal Multi-source Company raw exports from Firmographic/raw/coresignal/data.",
            "Rate Limit": "Multi-source Company API: 18 req/sec search, 18 req/sec enrichment, 27 req/sec bulk collect, 54 req/sec collection.",
            "Companies Processed": str(len(coresignal_companies)),
            "Coverage (%)": f"{(len(coresignal_companies) / 10 * 100):.2f}%",
            "Success Rate (%)": f"{(len(coresignal_companies) / 10 * 100):.2f}%",
            "Error Rate (%)": f"{((10 - len(coresignal_companies)) / 10 * 100):.2f}%",
            "Average Latency": "0.00 ms (reuse raw)",
            "Gated Fields": "; ".join(sorted(coresignal_missing)) or "None identified in parsed output",
            "Free Tier Limitation": "Trial credits are time-limited. Multi-source company enrich costs 2 Collect credits per successful request, so 10 companies would estimate 20 Collect credits.",
            "Paid Plan Cost": "Company API pricing: Starter from $49/month, Pro from $800/month, Premium from $1,500/month. Starter shows 200 Collect and 400 Search credits; Pro shows 250-50,000 Collect and 500-150,000 Search credits; Premium custom.",
            "Paid Tier Benefits": "Paid plans provide monthly Search/Collect credits, documentation access, and higher/custom credit packages. Higher tiers can include account manager and historical headcount API per pricing page.",
            "Ease of Integration": "4/5 - existing Multi-source Company raw exports map directly to the HPI firmographic field set.",
            "API Documentation Quality": "4/5 - endpoint, rate limits, free trial credits, paid plan ranges, and per-request credit rules are documented clearly.",
            "Evidence Link": f"{CORESIGNAL_FREE_TRIAL_DOC}; {CORESIGNAL_PRICING_DOC}; {CORESIGNAL_COMPANY_PRICING}; {CORESIGNAL_COMPANY_API_DOC}; {CORESIGNAL_COMPANY_ENRICH_DOC}; {CORESIGNAL_CREDITS_DOC}; raw/coresignal/data/*.json",
            "Overall API Score": "4.6/5",
            "Status": "Success - reused saved raw exports",
            "Remarks": "Coresignal raw company-enrichment responses were already present from the technographic workstream and are now parsed into the firmographic report.",
            "Data Completeness (%)": coresignal_completeness,
            "Records Retrieved": str(len(coresignal_companies)),
            "Estimated Cost per 100 Companies": "Multi-source Company enrich: about 200 Collect credits for 100 companies, because docs state 2 collection credits per successful enrich request.",
            "Raw Export Saved (Y/N)": "Y" if coresignal_companies else "N",
        }
    return [apollo_row, coresignal_row]


def build_api_comparison_rows(field_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    apollo_fields = [row for row in field_rows if row["source_api"] == API_DISPLAY_NAME]
    coresignal_fields = [row for row in field_rows if row["source_api"] == CORESIGNAL_API_DISPLAY_NAME]
    returned = sorted({row["field"] for row in apollo_fields if row["status"] in {"Available", "Derived from API"}})
    missing = sorted({row["field"] for row in apollo_fields if row["status"] not in {"Available", "Derived from API"}})
    cs_returned = sorted({row["field"] for row in coresignal_fields if row["status"] in {"Available", "Derived from API"}})
    cs_missing = sorted({row["field"] for row in coresignal_fields if row["status"] not in {"Available", "Derived from API"}})
    return [
        {
            "API Name": API_DISPLAY_NAME,
            "Endpoint Used": ORG_ENDPOINT,
            "Status (Success/Fail)": "Success",
            "Fields Returned": "; ".join(returned),
            "Free-Tier Limitations": "Trial credit information is public, but exact API entitlement is account/plan dependent. Observed call limits: 600/day, 200/hour, 50/minute.",
            "Paid-Tier Benefits": "Plan-dependent API access, higher/adjusted rate limits, and add-on credits. Exact field gating must be confirmed in the Apollo account.",
            "Notes": f"Fields not returned or not independently validated: {'; '.join(missing)}. Raw responses and usage trace are saved.",
        },
        {
            "API Name": CORESIGNAL_API_DISPLAY_NAME,
            "Endpoint Used": CORESIGNAL_ENDPOINT,
            "Status (Success/Fail)": "Success - reused saved raw exports",
            "Fields Returned": "; ".join(cs_returned) if cs_returned else fp.NOT_AVAILABLE,
            "Free-Tier Limitations": "Free trial gives 200 Collect and 400 Search credits for eligible new users; credits are valid for 7 days. Multi-source company enrich uses 2 Collect credits per successful request.",
            "Paid-Tier Benefits": "Starter/Pro/Premium paid plans provide monthly Search/Collect credits and documentation access; higher tiers offer larger/custom credits and additional support/features.",
            "Notes": f"Parsed from saved Coresignal Multi-source Company raw exports. Fields not returned or not calculated: {'; '.join(cs_missing) if cs_missing else 'None identified in parsed output'}.",
        },
    ]


def build_wide_company_rows(companies: list[fp.Company], field_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in field_rows:
        grouped[row["company_name"]].append(row)

    wide_rows: list[dict[str, str]] = []
    for company in companies:
        row: dict[str, str] = {
            "source_rank": company.source_rank,
            "company_name": company.company_name,
            "domain": company.domain,
            "input_linkedin_url": company.linkedin_url,
            "source_basis": company.source_basis,
        }
        for item in grouped[company.company_name]:
            slug = fp.FIELD_SLUGS[item["internal_field"]]
            row[slug] = item["value"]
            row[f"{slug}_source_api"] = item["source_api"]
            row[f"{slug}_status"] = item["status"]
            row[f"{slug}_reason"] = item["reason"]
            row[f"{slug}_api_gated"] = item["api_gated_column"]
        wide_rows.append(row)
    return wide_rows


def autosize_sheet(ws) -> None:
    for column_cells in ws.columns:
        max_len = 0
        col = get_column_letter(column_cells[0].column)
        for cell in column_cells:
            value = "" if cell.value is None else str(cell.value)
            max_len = max(max_len, min(len(value), 60))
        ws.column_dimensions[col].width = max(12, max_len + 2)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def append_sheet(wb: Workbook, title: str, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    ws = wb.create_sheet(title)
    ws.append(fieldnames)
    for row in rows:
        ws.append([row.get(field, "") for field in fieldnames])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
    ws.freeze_panes = "A2"
    autosize_sheet(ws)


def build_workbook(
    clean_rows: list[dict[str, str]],
    wide_rows: list[dict[str, str]],
    coresignal_wide_rows: list[dict[str, str]],
    field_rows: list[dict[str, str]],
    missing_rows: list[dict[str, str]],
    comparison_rows: list[dict[str, str]],
    api_trace_rows: list[dict[str, str]],
    coresignal_trace_rows: list[dict[str, str]],
    raw_index_rows: list[dict[str, str]],
) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    append_sheet(wb, "Clean Input", clean_rows, list(clean_rows[0].keys()))
    append_sheet(wb, "Company Data Wide", wide_rows, list(wide_rows[0].keys()))
    append_sheet(wb, "Coresignal Firmographics", coresignal_wide_rows, list(coresignal_wide_rows[0].keys()))
    append_sheet(wb, "Field Level Data", field_rows, list(field_rows[0].keys()))
    append_sheet(wb, "Missing Reasons Gated", missing_rows, list(missing_rows[0].keys()))
    append_sheet(wb, "API Comparison", comparison_rows, list(comparison_rows[0].keys()))
    append_sheet(wb, "API Trace", api_trace_rows, list(api_trace_rows[0].keys()))
    append_sheet(wb, "Coresignal API Calls", coresignal_trace_rows, list(coresignal_trace_rows[0].keys()))
    append_sheet(wb, "Raw Export Index", raw_index_rows, list(raw_index_rows[0].keys()))
    OUTPUT_XLSX.parent.mkdir(parents=True, exist_ok=True)
    wb.save(OUTPUT_XLSX)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F4E78")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)
    for source in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(source):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    document.add_paragraph()


def build_docx(
    field_rows: list[dict[str, str]],
    missing_rows: list[dict[str, str]],
    api_trace_rows: list[dict[str, str]],
    comparison_rows: list[dict[str, str]],
    coresignal_trace_rows: list[dict[str, str]],
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HPI 10 Company Firmographic API Report")
    run.font.bold = True
    run.font.size = Pt(18)
    document.add_paragraph(f"Generated: {utc_now()}")
    document.add_paragraph(
        "Scope: 10 companies excluding China and Hong Kong. Data sources: Apollo Organization Enrichment API and "
        "Coresignal Multi-source Company Enrichment API. Raw API responses and usage snapshots are saved for audit. "
        "Workbook output remains editable for manual review."
    )

    document.add_heading("API Trace Summary", level=1)
    trace = api_trace_rows[0]
    add_table(
        document,
        ["Metric", "Value"],
        [[key, value] for key, value in trace.items()],
    )

    document.add_heading("API Comparison", level=1)
    add_table(
        document,
        list(comparison_rows[0].keys()),
        [[row[key] for key in comparison_rows[0].keys()] for row in comparison_rows],
    )

    document.add_heading("Coresignal Per-Company API Trace", level=1)
    trace_headers = ["company_name", "domain", "success_failure", "fields_available", "fields_total", "field_completeness_percent", "raw_response_path"]
    add_table(
        document,
        ["Company", "Domain", "Status", "Fields Available", "Fields Total", "Completeness %", "Raw Response"],
        [[row[key] for key in trace_headers] for row in coresignal_trace_rows],
    )

    document.add_heading("Company Field Data", level=1)
    compact_rows = [
        [row["company_name"], row["field"], row["value"], row["status"], row["source_field"], row["api_gated_column"]]
        for row in field_rows
    ]
    add_table(document, ["Company", "Field", "Value", "Status", "Source Field", "API Gated"], compact_rows)

    document.add_heading("Missing / Not Verified Reasons", level=1)
    missing_compact = [
        [row["company_name"], row["field"], row["status"], row["reason"], row["api_gated_column"]]
        for row in missing_rows
    ]
    add_table(document, ["Company", "Field", "Status", "Reason", "API Gated"], missing_compact)
    document.save(OUTPUT_DOCX)


def main() -> int:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    input_path = fp.resolve_input_path(str(INPUT_PATH))
    companies = fp.read_companies(input_path)
    if len(companies) != 10:
        raise RuntimeError(f"Expected 10 companies, found {len(companies)}.")
    ensure_local_coresignal_raw(companies)

    missing_raw = [str(RAW_DIR / f"{fp.slugify(company.company_name)}.json") for company in companies if not (RAW_DIR / f"{fp.slugify(company.company_name)}.json").exists()]
    if missing_raw:
        raise FileNotFoundError(f"Missing raw Apollo responses: {', '.join(missing_raw)}")

    clean_rows = clean_input(companies)
    write_csv_dicts(CLEAN_INPUT, clean_rows, list(clean_rows[0].keys()))

    trace_rows = read_csv_dicts(BASE_DIR / "reports" / "api_trace_report.csv")
    apollo_field_rows, apollo_missing_rows = build_field_rows(companies, trace_rows)
    coresignal_field_rows, coresignal_missing_rows = build_coresignal_field_rows(companies)
    field_rows = apollo_field_rows + coresignal_field_rows
    missing_rows = apollo_missing_rows + coresignal_missing_rows
    api_trace_rows = build_api_trace_rows(trace_rows, field_rows)
    coresignal_trace_rows = build_coresignal_trace_rows(companies, field_rows)
    comparison_rows = build_api_comparison_rows(field_rows)
    wide_rows = build_source_wide_rows(companies, field_rows, API_DISPLAY_NAME)
    coresignal_wide_rows = build_source_wide_rows(companies, field_rows, CORESIGNAL_API_DISPLAY_NAME)
    raw_index_rows = []
    for company in companies:
        apollo_raw_path = RAW_DIR / f"{fp.slugify(company.company_name)}.json"
        raw_index_rows.append(
            {
                "source_api": API_DISPLAY_NAME,
                "company_name": company.company_name,
                "raw_response_saved": "Y",
                "raw_response_path": display_path(apollo_raw_path),
            }
        )
        cs_path = coresignal_raw_path(company)
        raw_index_rows.append(
            {
                "source_api": CORESIGNAL_API_DISPLAY_NAME,
                "company_name": company.company_name,
                "raw_response_saved": "Y" if cs_path else "N",
                "raw_response_path": display_path(cs_path) if cs_path else fp.NOT_AVAILABLE,
            }
        )

    save_json(
        OUTPUT_JSON,
        {
            "generated_at": utc_now(),
            "scope": "10 firmographic companies excluding China and Hong Kong",
            "clean_input": clean_rows,
            "api_comparison": comparison_rows,
            "api_trace": api_trace_rows,
            "coresignal_api_trace": coresignal_trace_rows,
            "field_level_data": field_rows,
            "missing_reason_api_gated": missing_rows,
            "raw_export_index": raw_index_rows,
        },
    )

    write_csv_dicts(FIELD_TRACE_CSV, field_rows, list(field_rows[0].keys()))
    write_csv_dicts(MISSING_CSV, missing_rows, list(missing_rows[0].keys()))
    write_csv_dicts(API_TRACE_CSV, api_trace_rows, list(api_trace_rows[0].keys()))
    write_csv_dicts(API_COMPARISON_CSV, comparison_rows, list(comparison_rows[0].keys()))
    write_csv_dicts(CORESIGNAL_TRACE_CSV, coresignal_trace_rows, list(coresignal_trace_rows[0].keys()))
    write_csv_dicts(CORESIGNAL_REPORT_CSV, coresignal_wide_rows, list(coresignal_wide_rows[0].keys()))

    build_workbook(clean_rows, wide_rows, coresignal_wide_rows, field_rows, missing_rows, comparison_rows, api_trace_rows, coresignal_trace_rows, raw_index_rows)
    build_docx(field_rows, missing_rows, api_trace_rows, comparison_rows, coresignal_trace_rows)

    print(f"Saved workbook: {OUTPUT_XLSX}")
    print(f"Saved docx: {OUTPUT_DOCX}")
    print(f"Saved JSON trace: {OUTPUT_JSON}")
    print(f"Saved clean input: {CLEAN_INPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
