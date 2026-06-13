from __future__ import annotations

import argparse
import csv
import json
import os
import re
import shutil
import sys
import time
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
INPUT_FILE = BASE_DIR / "input" / "compnys.txt"
REPORTS_DIR = BASE_DIR / "reports"
RAW_ROOT = BASE_DIR / "raw"
APILOG_ROOT = BASE_DIR / "apilogs"

APOLLO_DOC = "https://docs.apollo.io/reference/organization-jobs-postings"
THEIRSTACK_DOC = "https://theirstack.com/en/docs/api-reference/jobs/search_jobs_v1"
CORESIGNAL_DOC = "https://docs.coresignal.com/jobs-api/multi-source-jobs-api"
PREDICTLEADS_DOC = "https://docs.predictleads.com/"
LINKUP_DOC = "https://www.linkup.com/developers/"
AURA_DOC = "https://www.aura.com/"

APOLLO_JOB_ENDPOINT = "GET https://api.apollo.io/api/v1/organizations/{organization_id}/job_postings"
THEIRSTACK_JOB_ENDPOINT = "POST https://api.theirstack.com/v1/jobs/search"
CORESIGNAL_SOURCE_ENDPOINT = "Saved Coresignal Multi-source Company raw export with active_job_postings fields"

NOT_AVAILABLE = "Not Available"
NOT_TESTED = "Not Tested"

TARGET_FIELDS = [
    "Active job count (total)",
    "Active job count (SG)",
    "Roles by function (eng/design/IT/sales/ops)",
    "Hiring velocity (postings last 90 days)",
    "Job locations",
    "Posting first-seen date",
    "Source URL",
]

COMPANY_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "active_job_count_total",
    "active_job_count_sg",
    "roles_by_function",
    "hiring_velocity_90_days",
    "job_locations",
    "posting_first_seen_date",
    "source_urls",
    "sample_job_count",
    "data_completeness_percent",
    "raw_response_path",
]

JOB_DETAIL_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "job_title",
    "function",
    "location",
    "country",
    "posted_at",
    "first_seen",
    "source_url",
    "source_field",
    "raw_response_path",
]

API_CALL_FIELDS = [
    "timestamp",
    "api_name",
    "company_name",
    "domain",
    "endpoint_used",
    "http_status",
    "status",
    "latency_ms",
    "credits_used",
    "rate_limit",
    "records_retrieved",
    "error_message",
    "raw_response_path",
    "api_log_path",
]

API_TRACE_FIELDS = [
    "Tool Name",
    "Category",
    "API Available (Y/N)",
    "Authentication Type",
    "Free Credits / Tokens",
    "Credits / Tokens Used",
    "Rate Limit",
    "Companies Processed",
    "Coverage (%)",
    "Success Rate (%)",
    "Error Rate (%)",
    "Average Latency",
    "Gated Fields",
    "Free Tier Limitation",
    "Paid Plan Cost",
    "Paid Tier Benefits",
    "Ease of Integration",
    "API Documentation Quality",
    "Evidence Link",
    "Overall API Score",
    "Status",
    "Remarks",
    "Data Completeness (%)",
    "Records Retrieved",
    "Estimated Cost per 100 Companies",
    "Raw Export Saved (Y/N)",
]

API_COMPARISON_FIELDS = [
    "API Name",
    "Endpoint Used",
    "Status (Success/Fail)",
    "Fields Returned",
    "Free-Tier Limitations",
    "Paid-Tier Benefits",
    "Notes",
]

MISSING_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "field",
    "status",
    "reason",
    "raw_response_path",
]

ROLE_PATTERNS = {
    "engineering": re.compile(r"\b(engineer|developer|software|mechanical|electrical|civil|manufacturing|r&d|research|cad|cam|architect)\b", re.I),
    "design": re.compile(r"\b(design|designer|ux|ui|creative|product design|industrial design|graphics?)\b", re.I),
    "it": re.compile(r"\b(it|information technology|systems?|network|security|cyber|cloud|infrastructure|devops|database|support|helpdesk)\b", re.I),
    "sales": re.compile(r"\b(sales|account executive|business development|customer success|presales|inside sales|commercial)\b", re.I),
    "operations": re.compile(r"\b(operations?|ops|supply chain|logistics|warehouse|procurement|fulfillment|manufacturing operator)\b", re.I),
}


@dataclass
class Company:
    source_rank: str
    company_name: str
    domain: str
    linkedin_url: str
    source_basis: str


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_") or "company"


def normalize_domain(domain: str) -> str:
    domain = re.sub(r"^https?://", "", (domain or "").strip(), flags=re.I).split("/")[0]
    return domain[4:] if domain.lower().startswith("www.") else domain


def load_dotenv() -> None:
    for env_path in [BASE_DIR / ".env", ROOT_DIR / "Firmographic" / ".env", ROOT_DIR / "Technographic" / ".env"]:
        if not env_path.exists():
            continue
        for raw in env_path.read_text(encoding="utf-8").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def ensure_dirs() -> None:
    REPORTS_DIR.mkdir(exist_ok=True)
    for name in ["apollo", "theirstack", "coresignal", "predictleads"]:
        (RAW_ROOT / name / "data").mkdir(parents=True, exist_ok=True)
        (APILOG_ROOT / name).mkdir(parents=True, exist_ok=True)


def read_companies(path: Path, limit: int | None) -> list[Company]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        companies: list[Company] = []
        for row in reader:
            if not (row.get("company_name") or "").strip():
                continue
            companies.append(
                Company(
                    source_rank=(row.get("source_rank") or "").strip(),
                    company_name=(row.get("company_name") or "").strip(),
                    domain=normalize_domain(row.get("domain") or ""),
                    linkedin_url=(row.get("linkedin_url") or "").strip(),
                    source_basis=(row.get("source_basis") or "").strip(),
                )
            )
            if limit and len(companies) >= limit:
                break
        return companies


def safe_rel(path: Path) -> str:
    try:
        return str(path.relative_to(BASE_DIR)).replace("\\", "/")
    except ValueError:
        try:
            return str(path.relative_to(ROOT_DIR)).replace("\\", "/")
        except ValueError:
            return str(path)


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def append_jsonl(path: Path, event: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(event, ensure_ascii=False) + "\n")


def redact_headers(headers: dict[str, str]) -> dict[str, str]:
    safe = dict(headers)
    for key in list(safe):
        if key.lower() in {"authorization", "x-api-key", "apikey", "api-key"}:
            safe[key] = "[REDACTED]"
    return safe


def clean_headers(headers: requests.structures.CaseInsensitiveDict[str]) -> dict[str, str]:
    out: dict[str, str] = {}
    for key, value in headers.items():
        if key.lower() in {"date", "content-type"} or "rate" in key.lower() or "retry-after" in key.lower():
            out[key] = value
    return out


def http_request(method: str, url: str, *, headers: dict[str, str], params: dict[str, Any] | None = None, body: dict[str, Any] | None = None, timeout: int = 60) -> dict[str, Any]:
    started = time.perf_counter()
    started_at = now_iso()
    try:
        response = requests.request(method, url, headers=headers, params=params, json=body, timeout=timeout)
        latency = round((time.perf_counter() - started) * 1000, 2)
        try:
            payload = response.json()
        except ValueError:
            payload = {"text": response.text}
        return {
            "ok": 200 <= response.status_code < 300,
            "status_code": response.status_code,
            "latency_ms": latency,
            "payload": payload,
            "headers": clean_headers(response.headers),
            "error": "",
            "started_at": started_at,
            "finished_at": now_iso(),
        }
    except requests.RequestException as exc:
        return {
            "ok": False,
            "status_code": "",
            "latency_ms": round((time.perf_counter() - started) * 1000, 2),
            "payload": {"error": str(exc)},
            "headers": {},
            "error": str(exc),
            "started_at": started_at,
            "finished_at": now_iso(),
        }


def classify_function(title: str) -> str:
    for label, pattern in ROLE_PATTERNS.items():
        if pattern.search(title or ""):
            return label
    return "other"


def parse_date(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    match = re.search(r"\d{4}-\d{2}-\d{2}", value)
    return match.group(0) if match else ""


def first_value(obj: dict[str, Any], keys: list[str]) -> Any:
    for key in keys:
        value = obj.get(key)
        if value not in (None, "", [], {}):
            return value
    return ""


def find_jobs_list(payload: Any) -> list[dict[str, Any]]:
    if not isinstance(payload, dict):
        return []
    for key in ["jobs", "job_postings", "organization_job_postings", "postings", "data", "results", "active_job_postings"]:
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
    for value in payload.values():
        if isinstance(value, dict):
            found = find_jobs_list(value)
            if found:
                return found
    return []


def total_from_payload(payload: Any, fallback: int) -> str:
    if not isinstance(payload, dict):
        return str(fallback)
    for key in ["total_results", "total", "count", "total_count", "active_job_postings_count"]:
        value = payload.get(key)
        if isinstance(value, int):
            return str(value)
    metadata = payload.get("metadata")
    if isinstance(metadata, dict):
        for key in ["total_results", "total", "count", "total_count"]:
            value = metadata.get(key)
            if isinstance(value, int):
                return str(value)
    pagination = payload.get("pagination")
    if isinstance(pagination, dict):
        for key in ["total_entries", "total", "total_count"]:
            value = pagination.get(key)
            if isinstance(value, int):
                return str(value)
    return str(fallback)


def normalize_job(item: dict[str, Any], api_name: str, source_field: str) -> dict[str, str]:
    title = str(first_value(item, ["title", "job_title", "job_posting_title", "name"]) or "")
    location_value = first_value(item, ["location", "short_location", "long_location", "formatted_location", "job_location", "locations", "city"])
    location = format_location(location_value)
    country = normalize_country(first_value(item, ["country_code", "job_country_code", "country", "job_country", "country_codes", "countries"]) or "")
    location = enrich_location_from_parts(item, location)
    if not location and country:
        location = country
    posted_at = parse_date(first_value(item, ["posted_at", "date_posted", "published_at", "created_at", "last_seen_at"]))
    first_seen = parse_date(first_value(item, ["first_seen_at", "first_seen", "discovered_at", "posted_at", "created_at"]))
    source_url = str(first_value(item, ["url", "job_url", "source_url", "final_url", "apply_url", "redirect_url"]) or "")
    return {
        "job_title": title or NOT_AVAILABLE,
        "function": classify_function(title),
        "location": location or NOT_AVAILABLE,
        "country": country or infer_country(location),
        "posted_at": posted_at or NOT_AVAILABLE,
        "first_seen": first_seen or posted_at or NOT_AVAILABLE,
        "source_url": source_url or NOT_AVAILABLE,
        "source_field": source_field,
    }


def infer_country(location: str) -> str:
    text = (location or "").lower()
    if "singapore" in text or re.search(r"\bsg\b", text):
        return "SG"
    return ""


def format_location(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        for key in ["name", "formatted", "formatted_location", "long_location", "short_location", "city"]:
            if value.get(key):
                return str(value[key]).strip()
        parts = [value.get("city"), value.get("region"), value.get("country_code") or value.get("country")]
        return ", ".join(str(part).strip() for part in parts if part)
    if isinstance(value, list):
        locations = [format_location(item) for item in value[:3]]
        return "; ".join(location for location in locations if location)
    return ""


def enrich_location_from_parts(item: dict[str, Any], location: str) -> str:
    city = str(first_value(item, ["city", "job_city"]) or "").strip()
    state = str(first_value(item, ["state", "region", "job_state"]) or "").strip()
    country = normalize_country(first_value(item, ["country", "job_country", "country_code", "job_country_code"]) or "")
    if not city:
        return location
    if location and location.strip().lower() != city.lower():
        return location
    parts = [city, state, country]
    return ", ".join(part for part in parts if part)


def normalize_country(value: Any) -> str:
    if isinstance(value, list):
        value = first_value({"value": value}, ["value"])
        if isinstance(value, list):
            value = value[0] if value else ""
    if isinstance(value, dict):
        value = first_value(value, ["country_code", "code", "name", "country"])
    text = str(value or "").strip()
    if text.lower() in {"singapore", "sg", "sgp"}:
        return "SG"
    return text


def summarize_jobs(company: Company, api_name: str, payload: Any, raw_path: str, sg_total: str = "") -> tuple[dict[str, str], list[dict[str, str]], list[dict[str, str]]]:
    jobs = find_jobs_list(payload)
    details = []
    for idx, item in enumerate(jobs):
        row = normalize_job(item, api_name, f"jobs.{idx}")
        details.append(
            {
                "source_rank": company.source_rank,
                "company_name": company.company_name,
                "domain": company.domain,
                "api_name": api_name,
                "raw_response_path": raw_path,
                **row,
            }
        )

    total = total_from_payload(payload, len(jobs))
    sg_count = sg_total or str(sum(1 for row in details if row.get("country") == "SG" or "singapore" in row.get("location", "").lower()))
    role_counts = Counter(row["function"] for row in details)
    roles = "; ".join(f"{key}: {role_counts.get(key, 0)}" for key in ["engineering", "design", "it", "sales", "operations", "other"])
    locations = sorted({row["location"] for row in details if row["location"] != NOT_AVAILABLE})
    first_seen_dates = sorted({row["first_seen"] for row in details if row["first_seen"] != NOT_AVAILABLE})
    urls = [row["source_url"] for row in details if row["source_url"] != NOT_AVAILABLE]

    velocity = str(len(jobs))
    if isinstance(payload, dict):
        by_month = payload.get("active_job_postings_count_by_month")
        if isinstance(by_month, list):
            recent = []
            for entry in by_month[:4]:
                if isinstance(entry, dict) and isinstance(entry.get("active_job_postings_count"), int):
                    recent.append(entry["active_job_postings_count"])
            if recent:
                velocity = str(recent[0] - recent[-1]) if len(recent) > 1 else str(recent[0])

    populated = [
        total not in {"0", NOT_AVAILABLE, ""},
        sg_count not in {"", NOT_AVAILABLE},
        bool(details),
        velocity not in {"", NOT_AVAILABLE},
        bool(locations),
        bool(first_seen_dates),
        bool(urls),
    ]
    completeness = round(sum(populated) / len(TARGET_FIELDS) * 100, 2)
    summary = {
        "source_rank": company.source_rank,
        "company_name": company.company_name,
        "domain": company.domain,
        "api_name": api_name,
        "active_job_count_total": total,
        "active_job_count_sg": sg_count or NOT_AVAILABLE,
        "roles_by_function": roles if details else NOT_AVAILABLE,
        "hiring_velocity_90_days": velocity,
        "job_locations": "; ".join(locations[:25]) if locations else NOT_AVAILABLE,
        "posting_first_seen_date": first_seen_dates[0] if first_seen_dates else NOT_AVAILABLE,
        "source_urls": "; ".join(urls[:10]) if urls else NOT_AVAILABLE,
        "sample_job_count": str(len(details)),
        "data_completeness_percent": f"{completeness:.2f}",
        "raw_response_path": raw_path,
    }
    reasons = [
        (TARGET_FIELDS[0], total not in {"0", NOT_AVAILABLE, ""}, "No active total job count returned or parsed."),
        (TARGET_FIELDS[1], sg_count not in {"", NOT_AVAILABLE}, "No Singapore-specific count returned or inferred."),
        (TARGET_FIELDS[2], bool(details), "No job titles available for role/function classification."),
        (TARGET_FIELDS[3], velocity not in {"", NOT_AVAILABLE}, "No 90-day posting velocity returned or inferred."),
        (TARGET_FIELDS[4], bool(locations), "No job locations returned in sampled postings."),
        (TARGET_FIELDS[5], bool(first_seen_dates), "No posting first-seen/posted date returned in sampled postings."),
        (TARGET_FIELDS[6], bool(urls), "No source URLs returned in sampled postings."),
    ]
    missing = [
        {
            "source_rank": company.source_rank,
            "company_name": company.company_name,
            "domain": company.domain,
            "api_name": api_name,
            "field": field,
            "status": NOT_AVAILABLE,
            "reason": reason,
            "raw_response_path": raw_path,
        }
        for field, ok, reason in reasons
        if not ok
    ]
    return summary, details, missing


def copy_input() -> None:
    source = ROOT_DIR / "Firmographic" / "input" / "compnys.txt"
    if not INPUT_FILE.exists() and source.exists():
        shutil.copy2(source, INPUT_FILE)


def copy_coresignal_raw(companies: list[Company]) -> None:
    source_dir = ROOT_DIR / "Firmographic" / "raw" / "coresignal" / "data"
    target_dir = RAW_ROOT / "coresignal" / "data"
    for company in companies:
        src = source_dir / f"{slugify(company.company_name)}.json"
        dst = target_dir / src.name
        if src.exists() and not dst.exists():
            shutil.copy2(src, dst)


def apollo_headers(api_key: str) -> list[dict[str, str]]:
    preferred = os.getenv("APOLLO_AUTH_MODE", "").strip().lower()
    modes = [preferred] if preferred in {"bearer", "x-api-key"} else ["bearer", "x-api-key"]
    headers = []
    for mode in modes:
        base = {"Accept": "application/json", "Content-Type": "application/json"}
        if mode == "x-api-key":
            base["X-Api-Key"] = api_key
        else:
            base["Authorization"] = f"Bearer {api_key}"
        headers.append(base)
    return headers


def load_apollo_org_id(company: Company) -> str:
    raw_path = ROOT_DIR / "Firmographic" / "raw" / "apollo" / f"{slugify(company.company_name)}.json"
    if not raw_path.exists():
        return ""
    payload = load_json(raw_path)
    org = payload.get("organization") if isinstance(payload, dict) else {}
    return str(org.get("id") or "") if isinstance(org, dict) else ""


def collect_apollo(companies: list[Company], reuse_raw: bool, timeout: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_key = os.getenv("APOLLO_API_KEY") or os.getenv("APOLLO_MASTER_API_KEY") or ""
    api_name = "Apollo Organization Job Postings API"
    company_rows: list[dict[str, str]] = []
    detail_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    for company in companies:
        raw_path = RAW_ROOT / "apollo" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "apollo" / f"{slugify(company.company_name)}.jsonl"
        org_id = load_apollo_org_id(company)
        endpoint = f"https://api.apollo.io/api/v1/organizations/{org_id}/job_postings" if org_id else ""
        if reuse_raw and raw_path.exists():
            result = {"ok": True, "status_code": 200, "latency_ms": 0, "payload": load_json(raw_path), "headers": {}, "error": ""}
            used = "0 (reuse-raw)"
        elif not api_key or not org_id:
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": {"error": "Missing Apollo API key or organization ID."}, "headers": {}, "error": "Missing Apollo API key or organization ID."}
            save_json(raw_path, result["payload"])
            used = NOT_AVAILABLE
        else:
            result = {}
            for headers in apollo_headers(api_key):
                result = http_request("GET", endpoint, headers=headers, params={"page": 1, "per_page": 25}, timeout=timeout)
                if result["status_code"] != 401:
                    break
            save_json(raw_path, result["payload"])
            append_jsonl(log_path, {"logged_at": now_iso(), "request": {"method": "GET", "url": endpoint, "headers": redact_headers(headers), "params": {"page": 1, "per_page": 25}}, "response": {k: result[k] for k in ["status_code", "ok", "latency_ms", "headers", "error"]}, "raw_response_path": safe_rel(raw_path)})
            used = "Estimated 1"

        raw_rel = safe_rel(raw_path)
        summary, details, missing = summarize_jobs(company, api_name, result["payload"], raw_rel)
        company_rows.append(summary)
        detail_rows.extend(details)
        missing_rows.extend(missing)
        records = len(details)
        status = "Success" if result["ok"] and (records or summary["active_job_count_total"] != "0") else "Fail"
        call_rows.append(api_call_row(api_name, company, APOLLO_JOB_ENDPOINT, result, status, used, records, raw_rel, safe_rel(log_path)))
        print(f"apollo: {company.company_name} | {status} | HTTP {result['status_code']} | jobs={records}")
    return company_rows, detail_rows, missing_rows, call_rows


def collect_theirstack(companies: list[Company], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_key = os.getenv("THEIRSTACK_API_KEY") or ""
    api_name = "TheirStack Job Search API"
    company_rows: list[dict[str, str]] = []
    detail_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    for company in companies:
        raw_path = RAW_ROOT / "theirstack" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "theirstack" / f"{slugify(company.company_name)}.jsonl"
        payload = {"company_domain_or": [company.domain], "limit": limit, "page": 0, "include_total_results": True}
        sg_payload = {"company_domain_or": [company.domain], "job_country_code_or": ["SG"], "limit": 1, "page": 0, "include_total_results": True}
        if reuse_raw and raw_path.exists():
            result = {"ok": True, "status_code": 200, "latency_ms": 0, "payload": load_json(raw_path), "headers": {}, "error": ""}
            used = "0 (reuse-raw)"
        elif not api_key:
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": {"error": "Missing TheirStack API key."}, "headers": {}, "error": "Missing TheirStack API key."}
            save_json(raw_path, result["payload"])
            used = NOT_AVAILABLE
        else:
            headers = {"Accept": "application/json", "Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
            result = http_request("POST", "https://api.theirstack.com/v1/jobs/search", headers=headers, body=payload, timeout=timeout)
            sg_result = http_request("POST", "https://api.theirstack.com/v1/jobs/search", headers=headers, body=sg_payload, timeout=timeout)
            result["payload"] = {"all_jobs_response": result["payload"], "sg_jobs_response": sg_result["payload"]}
            result["ok"] = bool(result["ok"])
            result["latency_ms"] = round(float(result["latency_ms"]) + float(sg_result["latency_ms"]), 2)
            save_json(raw_path, result["payload"])
            append_jsonl(log_path, {"logged_at": now_iso(), "request": {"method": "POST", "url": "https://api.theirstack.com/v1/jobs/search", "headers": redact_headers(headers), "json": payload, "sg_json": sg_payload}, "response": {"status_code": result["status_code"], "ok": result["ok"], "latency_ms": result["latency_ms"], "headers": result["headers"], "error": result["error"]}, "raw_response_path": safe_rel(raw_path)})
            used = f"Estimated {limit + 1} max returned-job credits"

        raw_rel = safe_rel(raw_path)
        response_payload = result["payload"].get("all_jobs_response", result["payload"]) if isinstance(result["payload"], dict) else result["payload"]
        sg_payload_response = result["payload"].get("sg_jobs_response", {}) if isinstance(result["payload"], dict) else {}
        sg_total = total_from_payload(sg_payload_response, 0)
        summary, details, missing = summarize_jobs(company, api_name, response_payload, raw_rel, sg_total=sg_total)
        company_rows.append(summary)
        detail_rows.extend(details)
        missing_rows.extend(missing)
        records = len(details)
        status = "Success" if result["ok"] and (records or summary["active_job_count_total"] != "0") else "Fail"
        call_rows.append(api_call_row(api_name, company, THEIRSTACK_JOB_ENDPOINT, result, status, used, records, raw_rel, safe_rel(log_path)))
        print(f"theirstack: {company.company_name} | {status} | HTTP {result['status_code']} | jobs={records}")
    return company_rows, detail_rows, missing_rows, call_rows


def collect_coresignal(companies: list[Company]) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Coresignal Multi-source Company Jobs Signals"
    company_rows: list[dict[str, str]] = []
    detail_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    copy_coresignal_raw(companies)
    for company in companies:
        raw_path = RAW_ROOT / "coresignal" / "data" / f"{slugify(company.company_name)}.json"
        payload = load_json(raw_path) if raw_path.exists() else {"error": "Missing saved Coresignal raw export."}
        raw_rel = safe_rel(raw_path)
        log_path = APILOG_ROOT / "coresignal" / f"{slugify(company.company_name)}.jsonl"
        if log_path.exists():
            log_path.unlink()
        summary, details, missing = summarize_jobs(company, api_name, payload, raw_rel)
        company_rows.append(summary)
        detail_rows.extend(details)
        missing_rows.extend(missing)
        result = {"ok": raw_path.exists(), "status_code": 200 if raw_path.exists() else "", "latency_ms": 0, "headers": {}, "error": "" if raw_path.exists() else "Missing saved Coresignal raw export."}
        append_jsonl(
            log_path,
            {
                "logged_at": now_iso(),
                "request": {
                    "method": "REUSE_RAW",
                    "source": CORESIGNAL_SOURCE_ENDPOINT,
                    "raw_response_path": raw_rel,
                },
                "response": {
                    "status_code": result["status_code"],
                    "ok": result["ok"],
                    "latency_ms": result["latency_ms"],
                    "headers": result["headers"],
                    "error": result["error"],
                    "records_retrieved": len(details),
                },
            },
        )
        call_rows.append(api_call_row(api_name, company, CORESIGNAL_SOURCE_ENDPOINT, result, "Success" if raw_path.exists() else "Fail", "0 (reuse raw)", len(details), raw_rel, safe_rel(log_path)))
        print(f"coresignal: {company.company_name} | {'Success' if raw_path.exists() else 'Fail'} | jobs={len(details)}")
    return company_rows, detail_rows, missing_rows, call_rows


def collect_blocked(companies: list[Company], api_name: str, endpoint: str, reason: str) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    company_rows: list[dict[str, str]] = []
    missing_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    for company in companies:
        raw_path = RAW_ROOT / slugify(api_name) / "data" / f"{slugify(company.company_name)}.json"
        save_json(raw_path, {"status": "not_executed", "reason": reason})
        raw_rel = safe_rel(raw_path)
        summary = {
            "source_rank": company.source_rank,
            "company_name": company.company_name,
            "domain": company.domain,
            "api_name": api_name,
            "active_job_count_total": NOT_TESTED,
            "active_job_count_sg": NOT_TESTED,
            "roles_by_function": NOT_TESTED,
            "hiring_velocity_90_days": NOT_TESTED,
            "job_locations": NOT_TESTED,
            "posting_first_seen_date": NOT_TESTED,
            "source_urls": NOT_TESTED,
            "sample_job_count": "0",
            "data_completeness_percent": "0.00",
            "raw_response_path": raw_rel,
        }
        company_rows.append(summary)
        for field in TARGET_FIELDS:
            missing_rows.append({"source_rank": company.source_rank, "company_name": company.company_name, "domain": company.domain, "api_name": api_name, "field": field, "status": NOT_TESTED, "reason": reason, "raw_response_path": raw_rel})
        result = {"ok": False, "status_code": NOT_TESTED, "latency_ms": 0, "headers": {}, "error": reason}
        call_rows.append(api_call_row(api_name, company, endpoint, result, NOT_TESTED, "0", 0, raw_rel, ""))
    return company_rows, [], missing_rows, call_rows


def api_call_row(api_name: str, company: Company, endpoint: str, result: dict[str, Any], status: str, credits_used: str, records: int, raw_rel: str, log_rel: str) -> dict[str, str]:
    return {
        "timestamp": now_iso(),
        "api_name": api_name,
        "company_name": company.company_name,
        "domain": company.domain,
        "endpoint_used": endpoint,
        "http_status": str(result.get("status_code", "")),
        "status": status,
        "latency_ms": str(result.get("latency_ms", "")),
        "credits_used": credits_used,
        "rate_limit": "; ".join(f"{k}: {v}" for k, v in (result.get("headers") or {}).items() if "rate" in k.lower()),
        "records_retrieved": str(records),
        "error_message": str(result.get("error") or ""),
        "raw_response_path": raw_rel,
        "api_log_path": log_rel,
    }


def fields_returned(rows: list[dict[str, str]], api_name: str) -> list[str]:
    api_rows = [row for row in rows if row["api_name"] == api_name]
    returned = []
    checks = {
        "active_job_count_total": "Active job count (total)",
        "active_job_count_sg": "Active job count (SG)",
        "roles_by_function": "Roles by function",
        "hiring_velocity_90_days": "Hiring velocity",
        "job_locations": "Job locations",
        "posting_first_seen_date": "Posting first-seen date",
        "source_urls": "Source URL",
    }
    for key, label in checks.items():
        if any(row.get(key) not in {"", NOT_AVAILABLE, NOT_TESTED, "0"} for row in api_rows):
            returned.append(label)
    return returned


def build_api_comparison(company_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    configs = [
        ("PredictLeads Job Openings API", "PredictLeads", "Invalid API credential observed; not executed for this run.", "Free: 100 API requests/month", "Job openings dataset, history, categories, and URLs if valid credentials are available."),
        ("TheirStack Job Search API", THEIRSTACK_JOB_ENDPOINT, "Live job search executed where raw was absent.", "Free: 200 credits/month; job search consumes 1 credit per returned job.", "Large job-posting database with company, title, location, country, and date filters."),
        ("Coresignal Multi-source Company Jobs Signals", CORESIGNAL_SOURCE_ENDPOINT, "Reused saved Coresignal company raw exports.", "Trial credits are account-specific; Multi-source jobs API uses Search/Collect credits.", "Can use company jobs signals and dedicated Multi-source Jobs API for deeper job records."),
        ("Apollo Organization Job Postings API", APOLLO_JOB_ENDPOINT, "Live Apollo org job postings endpoint executed where raw was absent.", "Existing paid Apollo plan; job postings endpoint consumes credits.", "Convenient if Apollo org IDs already exist from firmographic enrichment."),
        ("LinkUp / Aura", "Sample/demo on request", "Not executed - sample/demo required.", "Demo/sample on request.", "Raw de-duplicated postings at scale; evaluate if free/available APIs fall short."),
    ]
    rows = []
    for api_name, endpoint, notes, free, paid in configs:
        returned = fields_returned(company_rows, api_name)
        status = "Success" if returned else "Not Tested" if api_name in {"PredictLeads Job Openings API", "LinkUp / Aura"} else "Fail"
        rows.append({"API Name": api_name, "Endpoint Used": endpoint, "Status (Success/Fail)": status, "Fields Returned": "; ".join(returned) if returned else NOT_AVAILABLE, "Free-Tier Limitations": free, "Paid-Tier Benefits": paid, "Notes": notes})
    return rows


def build_api_trace(company_rows: list[dict[str, str]], call_rows: list[dict[str, str]], missing_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    configs = {
        "PredictLeads Job Openings API": ("API key query/header", PREDICTLEADS_DOC, "Invalid credential in inventory; no valid live test completed.", "Free 100 requests/month", "Vendor paid dataset/API access", "100 requests per month free tier", "Not executed - invalid API key"),
        "TheirStack Job Search API": ("Bearer API key", THEIRSTACK_DOC, "Job endpoint docs are clear; use company_domain_or and include_total_results.", "200 credits/month; job search costs 1 credit per returned job", "Higher volume job search and filters", "Vendor/account rate limits in response headers", ""),
        "Coresignal Multi-source Company Jobs Signals": ("API key header", CORESIGNAL_DOC, "Reused saved company enrichment raw exports; dedicated Jobs API available for future deeper collection.", "Trial Search/Collect credits", "Multi-source Jobs Search/Collect endpoints and bulk delivery", "18 req/sec documented across CoreSignal enrichment endpoints; Jobs API docs list search/collect endpoints", ""),
        "Apollo Organization Job Postings API": ("Bearer or X-Api-Key", APOLLO_DOC, "Requires Apollo organization_id from company enrichment.", "Existing paid plan; credits are plan dependent", "Plan-dependent job postings access and rate limits", "Apollo rate limits are account/plan dependent", ""),
        "LinkUp / Aura": ("Demo/sample request", f"{LINKUP_DOC}; {AURA_DOC}", "Not executed because access is sample/demo on request.", "Sample on request", "Enterprise access to job postings feeds", "Vendor specific", "Not executed - sample/demo on request"),
    }
    rows = []
    for api_name, (auth, evidence, remarks, free, paid, rate, forced_status) in configs.items():
        calls = [row for row in call_rows if row["api_name"] == api_name]
        summaries = [row for row in company_rows if row["api_name"] == api_name]
        successes = [row for row in calls if row["status"] == "Success"]
        latencies = [float(row["latency_ms"]) for row in calls if str(row["latency_ms"]).replace(".", "", 1).isdigit()]
        completeness = [float(row["data_completeness_percent"]) for row in summaries if row.get("data_completeness_percent")]
        records = sum(int(row["records_retrieved"]) for row in calls if row["records_retrieved"].isdigit())
        missing = sorted({row["field"] for row in missing_rows if row["api_name"] == api_name and row["status"] != NOT_TESTED})
        processed = len(calls)
        success_rate = len(successes) / processed * 100 if processed else 0
        status = forced_status or ("Success" if successes else "Fail")
        rows.append({
            "Tool Name": api_name,
            "Category": "Jobs / Hiring",
            "API Available (Y/N)": "Y" if api_name not in {"LinkUp / Aura"} else "Sample on request",
            "Authentication Type": auth,
            "Free Credits / Tokens": free,
            "Credits / Tokens Used": "; ".join(row["credits_used"] for row in calls if row["credits_used"]) or NOT_AVAILABLE,
            "Rate Limit": rate,
            "Companies Processed": str(processed),
            "Coverage (%)": f"{success_rate:.2f}",
            "Success Rate (%)": f"{success_rate:.2f}",
            "Error Rate (%)": f"{(100 - success_rate if processed else 0):.2f}",
            "Average Latency": f"{(sum(latencies) / len(latencies)):.2f} ms" if latencies else NOT_AVAILABLE,
            "Gated Fields": "; ".join(missing) if missing else "None identified in parsed output",
            "Free Tier Limitation": free,
            "Paid Plan Cost": "Vendor/account specific",
            "Paid Tier Benefits": paid,
            "Ease of Integration": "Good - standard REST/JSON" if api_name not in {"PredictLeads Job Openings API", "LinkUp / Aura"} else "Blocked by access/credential",
            "API Documentation Quality": "Good - endpoint documented" if api_name != "LinkUp / Aura" else "Not assessed - sample/demo required",
            "Evidence Link": evidence,
            "Overall API Score": score_api(success_rate, sum(completeness) / len(completeness) if completeness else 0, records),
            "Status": status,
            "Remarks": remarks,
            "Data Completeness (%)": f"{(sum(completeness) / len(completeness)):.2f}" if completeness else "0.00",
            "Records Retrieved": str(records),
            "Estimated Cost per 100 Companies": estimated_cost(api_name),
            "Raw Export Saved (Y/N)": "Y" if any(row["raw_response_path"] for row in calls) else "N",
        })
    return rows


def estimated_cost(api_name: str) -> str:
    if api_name == "TheirStack Job Search API":
        return "About returned-job credits only; this run capped samples to control spend."
    if api_name == "Apollo Organization Job Postings API":
        return "About 100 Apollo job-postings calls for 100 companies; exact credit cost is plan specific."
    if api_name == "Coresignal Multi-source Company Jobs Signals":
        return "Reuse from company enrichment is 0 additional calls; dedicated Jobs API pricing depends on Search/Collect credits."
    if api_name == "PredictLeads Job Openings API":
        return "100 free requests/month if valid; paid plan required beyond free tier."
    return "Sample/demo required."


def score_api(success_rate: float, completeness: float, records: int) -> str:
    score = 1.0 + min(success_rate / 100 * 1.5, 1.5) + min(completeness / 100 * 1.5, 1.5) + (1.0 if records else 0.0)
    return f"{min(score, 5.0):.2f}/5"


def write_workbook(path: Path, sheets: dict[str, tuple[list[str], list[dict[str, Any]]]]) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    for title, (fields, rows) in sheets.items():
        ws = wb.create_sheet(title)
        ws.append(fields)
        for row in rows:
            ws.append([row.get(field, "") for field in fields])
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = Font(bold=True, color="FFFFFF")
        ws.freeze_panes = "A2"
        for column in ws.columns:
            col = get_column_letter(column[0].column)
            width = max(12, min(max(len(str(cell.value or "")) for cell in column) + 2, 60))
            ws.column_dimensions[col].width = width
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(path)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F4E78")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    document.add_paragraph()


def write_docx(path: Path, company_rows: list[dict[str, str]], api_trace_rows: list[dict[str, str]], comparison_rows: list[dict[str, str]], missing_rows: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HPI Jobs / Hiring API Evaluation Report")
    run.font.bold = True
    run.font.size = Pt(18)
    document.add_paragraph(f"Generated: {now_iso()}")
    document.add_paragraph("Scope: Jobs/Hiring provider evaluation for the 10-company HPI pilot set. Target fields follow Section 7 of the HPI Data Provider Evaluation brief.")
    document.add_heading("API Trace Summary", level=1)
    add_table(document, API_TRACE_FIELDS, [[row.get(field, "") for field in API_TRACE_FIELDS] for row in api_trace_rows])
    document.add_heading("API Comparison", level=1)
    add_table(document, API_COMPARISON_FIELDS, [[row.get(field, "") for field in API_COMPARISON_FIELDS] for row in comparison_rows])
    document.add_heading("Company Jobs Summary", level=1)
    compact_fields = ["company_name", "api_name", "active_job_count_total", "active_job_count_sg", "roles_by_function", "hiring_velocity_90_days", "data_completeness_percent"]
    add_table(document, ["Company", "API", "Total Jobs", "SG Jobs", "Roles", "90-Day Velocity", "Completeness %"], [[row.get(field, "") for field in compact_fields] for row in company_rows])
    document.add_heading("Missing / Not Tested Fields", level=1)
    add_table(document, MISSING_FIELDS, [[row.get(field, "") for field in MISSING_FIELDS] for row in missing_rows])
    document.save(path)


def run(args: argparse.Namespace) -> int:
    load_dotenv()
    ensure_dirs()
    copy_input()
    companies = read_companies(Path(args.input) if args.input else INPUT_FILE, args.limit)
    if not companies:
        raise ValueError("No companies found for Jobs/Hiring pipeline.")

    all_company_rows: list[dict[str, str]] = []
    all_detail_rows: list[dict[str, str]] = []
    all_missing_rows: list[dict[str, str]] = []
    all_call_rows: list[dict[str, str]] = []

    collectors = [
        collect_coresignal(companies),
        collect_theirstack(companies, args.reuse_raw, args.timeout, args.theirstack_limit),
        collect_apollo(companies, args.reuse_raw, args.timeout),
        collect_blocked(companies, "PredictLeads Job Openings API", "PredictLeads job openings endpoint", "PredictLeads API credential was reported invalid; leave as blocked until a valid key is provided."),
        collect_blocked(companies, "LinkUp / Aura", "Sample/demo on request", "LinkUp/Aura require sample/demo access; not executed for this free-tier run."),
    ]
    for company_rows, detail_rows, missing_rows, call_rows in collectors:
        all_company_rows.extend(company_rows)
        all_detail_rows.extend(detail_rows)
        all_missing_rows.extend(missing_rows)
        all_call_rows.extend(call_rows)

    comparison_rows = build_api_comparison(all_company_rows)
    trace_rows = build_api_trace(all_company_rows, all_call_rows, all_missing_rows)

    write_csv(REPORTS_DIR / "company_jobs_hiring.csv", all_company_rows, COMPANY_FIELDS)
    write_csv(REPORTS_DIR / "job_detail.csv", all_detail_rows, JOB_DETAIL_FIELDS)
    write_csv(REPORTS_DIR / "missing_fields_report.csv", all_missing_rows, MISSING_FIELDS)
    write_csv(REPORTS_DIR / "api_call_log.csv", all_call_rows, API_CALL_FIELDS)
    write_csv(REPORTS_DIR / "api_comparison_report.csv", comparison_rows, API_COMPARISON_FIELDS)
    write_csv(REPORTS_DIR / "api_tracing_report.csv", trace_rows, API_TRACE_FIELDS)
    save_json(REPORTS_DIR / "company_jobs_hiring.json", all_company_rows)
    save_json(REPORTS_DIR / "run_manifest.json", {"generated_at": now_iso(), "companies_processed": len(companies), "providers": ["coresignal", "theirstack", "apollo", "predictleads", "linkup/aura"], "reuse_raw": args.reuse_raw})

    xlsx_path = REPORTS_DIR / f"hpi_jobs_hiring_api_evaluation_{stamp()}.xlsx"
    docx_path = REPORTS_DIR / f"hpi_jobs_hiring_api_evaluation_{stamp()}.docx"
    write_workbook(
        xlsx_path,
        {
            "API Comparison": (API_COMPARISON_FIELDS, comparison_rows),
            "API Trace": (API_TRACE_FIELDS, trace_rows),
            "API Call Log": (API_CALL_FIELDS, all_call_rows),
            "Company Jobs Summary": (COMPANY_FIELDS, all_company_rows),
            "Job Detail": (JOB_DETAIL_FIELDS, all_detail_rows),
            "Missing Fields": (MISSING_FIELDS, all_missing_rows),
        },
    )
    write_docx(docx_path, all_company_rows, trace_rows, comparison_rows, all_missing_rows)
    print(f"Saved workbook: {safe_rel(xlsx_path)}")
    print(f"Saved docx: {safe_rel(docx_path)}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Run HPI Jobs/Hiring API evaluation.")
    parser.add_argument("--input", help="Input CSV path. Defaults to JobsHiring/input/compnys.txt.")
    parser.add_argument("--limit", type=int, default=10)
    parser.add_argument("--reuse-raw", action="store_true", help="Reuse saved raw responses.")
    parser.add_argument("--timeout", type=int, default=60)
    parser.add_argument("--theirstack-limit", type=int, default=10, help="Sample jobs returned per company for TheirStack.")
    args = parser.parse_args()
    try:
        return run(args)
    except Exception as exc:
        print(f"Jobs/Hiring pipeline failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
