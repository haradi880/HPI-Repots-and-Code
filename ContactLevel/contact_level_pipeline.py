from __future__ import annotations

import argparse
import csv
import json
import os
import re
import shutil
import sys
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
INPUT_FILE = BASE_DIR / "input" / "compnys.txt"
REPORTS_DIR = BASE_DIR / "reports"
RAW_ROOT = BASE_DIR / "raw"
APILOG_ROOT = BASE_DIR / "apilogs"

NOT_AVAILABLE = "Not Available"
NOT_TESTED = "Not Tested"

APOLLO_DOC = "https://docs.apollo.io/reference/people-api-search"
FULLENRICH_DOC = "https://docs.fullenrich.com/api/v2/contact/enrich/bulk/post"
PROSPEO_DOC = "https://docs.prospeo.io/reference/person-enrichment"
SIGNALHIRE_DOC = "https://www.signalhire.com/api"
PDL_DOC = "https://docs.peopledatalabs.com/"

APOLLO_ENDPOINT = "POST https://api.apollo.io/api/v1/mixed_people/api_search"
APOLLO_ENRICHMENT_ENDPOINT = "POST https://api.apollo.io/api/v1/people/match"
FULLENRICH_BULK_ENDPOINT = "POST https://app.fullenrich.com/api/v2/contact/enrich/bulk"
FULLENRICH_RESULT_ENDPOINT = "GET https://app.fullenrich.com/api/v1/contact/enrich/bulk/{enrichment_id}"
PROSPEO_ENDPOINT = "POST https://api.prospeo.io/enrich-person"
SIGNALHIRE_ENDPOINT = "POST https://www.signalhire.com/api/v1/candidate/search"

TARGET_FIELDS = [
    "Verified work email + confidence",
    "Direct/mobile phone + confidence",
    "Title, seniority, department",
    "Reports-to / org-chart hierarchy",
    "LinkedIn URL",
    "Match rate per requested contact",
]

CONTACT_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "requested_contact",
    "matched_contact",
    "match_status",
    "match_confidence",
    "verified_work_email",
    "email_confidence",
    "direct_mobile_phone",
    "phone_confidence",
    "title",
    "seniority",
    "department",
    "reports_to",
    "linkedin_url",
    "raw_response_path",
]

COMPANY_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "requested_contacts",
    "matched_contacts",
    "match_rate_percent",
    "verified_work_email_count",
    "direct_mobile_phone_count",
    "linkedin_url_count",
    "title_count",
    "reports_to_count",
    "data_completeness_percent",
    "raw_response_path",
]

API_CALL_FIELDS = [
    "timestamp",
    "api_name",
    "company_name",
    "domain",
    "endpoint_used",
    "http_status",
    "status",
    "latency_ms",
    "credits_used",
    "rate_limit",
    "records_retrieved",
    "error_message",
    "raw_response_path",
    "api_log_path",
]

API_TRACE_FIELDS = [
    "Tool Name",
    "Category",
    "API Available (Y/N)",
    "Authentication Type",
    "Free Credits / Tokens",
    "Credits / Tokens Used",
    "Rate Limit",
    "Companies Processed",
    "Coverage (%)",
    "Success Rate (%)",
    "Error Rate (%)",
    "Average Latency",
    "Gated Fields",
    "Free Tier Limitation",
    "Paid Plan Cost",
    "Paid Tier Benefits",
    "Ease of Integration",
    "API Documentation Quality",
    "Evidence Link",
    "Overall API Score",
    "Status",
    "Remarks",
    "Data Completeness (%)",
    "Records Retrieved",
    "Estimated Cost per 100 Companies",
    "Raw Export Saved (Y/N)",
]

API_COMPARISON_FIELDS = [
    "API Name",
    "Endpoint Used",
    "Status (Success/Fail)",
    "Fields Returned",
    "Free-Tier Limitations",
    "Paid-Tier Benefits",
    "Notes",
]

MISSING_FIELDS = [
    "source_rank",
    "company_name",
    "domain",
    "api_name",
    "requested_contact",
    "field",
    "status",
    "reason",
    "raw_response_path",
]


@dataclass
class Company:
    source_rank: str
    company_name: str
    domain: str
    linkedin_url: str
    source_basis: str


@dataclass
class SeedContact:
    company: Company
    first_name: str
    last_name: str
    full_name: str
    title: str
    seniority: str
    department: str
    linkedin_url: str
    source_api: str


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_") or "company"


def normalize_domain(domain: str) -> str:
    domain = re.sub(r"^https?://", "", (domain or "").strip(), flags=re.I).split("/")[0]
    return domain[4:] if domain.lower().startswith("www.") else domain


def load_dotenv() -> None:
    for env_path in [BASE_DIR / ".env", ROOT_DIR / "Firmographic" / ".env", ROOT_DIR / "Technographic" / ".env", ROOT_DIR / "JobsHiring" / ".env"]:
        if not env_path.exists():
            continue
        for raw in env_path.read_text(encoding="utf-8").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def ensure_dirs() -> None:
    REPORTS_DIR.mkdir(exist_ok=True)
    (BASE_DIR / "input").mkdir(exist_ok=True)
    for provider in ["apollo", "apollo_enrichment", "fullenrich", "fullenrich_search", "prospeo", "signalhire", "people_data_labs"]:
        (RAW_ROOT / provider / "data").mkdir(parents=True, exist_ok=True)
        (APILOG_ROOT / provider).mkdir(parents=True, exist_ok=True)


def copy_input() -> None:
    source = ROOT_DIR / "Firmographic" / "input" / "compnys.txt"
    if not INPUT_FILE.exists() and source.exists():
        shutil.copy2(source, INPUT_FILE)


def read_companies(path: Path, limit: int | None) -> list[Company]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        companies: list[Company] = []
        for row in reader:
            if not (row.get("company_name") or "").strip():
                continue
            companies.append(
                Company(
                    source_rank=(row.get("source_rank") or "").strip(),
                    company_name=(row.get("company_name") or "").strip(),
                    domain=normalize_domain(row.get("domain") or ""),
                    linkedin_url=(row.get("linkedin_url") or "").strip(),
                    source_basis=(row.get("source_basis") or "").strip(),
                )
            )
            if limit and len(companies) >= limit:
                break
        return companies


def safe_rel(path: Path) -> str:
    try:
        return str(path.relative_to(BASE_DIR)).replace("\\", "/")
    except ValueError:
        try:
            return str(path.relative_to(ROOT_DIR)).replace("\\", "/")
        except ValueError:
            return str(path)


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def payload_has_error(payload: Any) -> bool:
    if isinstance(payload, dict):
        if payload.get("error") or payload.get("error_code"):
            return True
        for key in ["start_response", "result_response"]:
            value = payload.get(key)
            if isinstance(value, dict) and payload_has_error(value):
                return True
    return False


def result_from_saved_payload(payload: Any) -> dict[str, Any]:
    ok = not payload_has_error(payload)
    return {"ok": ok, "status_code": 200 if ok else "Saved Error", "latency_ms": 0, "payload": payload, "headers": {}, "error": "" if ok else summarize_error(payload)}


def write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def append_jsonl(path: Path, event: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(event, ensure_ascii=False) + "\n")


def redact_headers(headers: dict[str, str]) -> dict[str, str]:
    safe = dict(headers)
    for key in list(safe):
        if key.lower() in {"authorization", "x-api-key", "x-key", "apikey", "api-key"}:
            safe[key] = "[REDACTED]"
    return safe


def clean_headers(headers: requests.structures.CaseInsensitiveDict[str]) -> dict[str, str]:
    out: dict[str, str] = {}
    for key, value in headers.items():
        low = key.lower()
        if low in {"date", "content-type", "x-credits-left"} or "rate" in low or "retry-after" in low:
            out[key] = value
    return out


def http_request(method: str, url: str, *, headers: dict[str, str], params: dict[str, Any] | None = None, body: dict[str, Any] | None = None, timeout: int = 60) -> dict[str, Any]:
    started = time.perf_counter()
    started_at = now_iso()
    try:
        response = requests.request(method, url, headers=headers, params=params, json=body, timeout=timeout)
        latency = round((time.perf_counter() - started) * 1000, 2)
        try:
            payload = response.json()
        except ValueError:
            payload = {"text": response.text}
        return {
            "ok": 200 <= response.status_code < 300,
            "status_code": response.status_code,
            "latency_ms": latency,
            "payload": payload,
            "headers": clean_headers(response.headers),
            "error": "" if 200 <= response.status_code < 300 else summarize_error(payload),
            "started_at": started_at,
            "finished_at": now_iso(),
        }
    except requests.RequestException as exc:
        return {
            "ok": False,
            "status_code": "",
            "latency_ms": round((time.perf_counter() - started) * 1000, 2),
            "payload": {"error": str(exc)},
            "headers": {},
            "error": str(exc),
            "started_at": started_at,
            "finished_at": now_iso(),
        }


def summarize_error(payload: Any) -> str:
    if isinstance(payload, dict):
        for key in ["error", "message", "detail"]:
            if payload.get(key):
                return str(payload[key])[:300]
    return ""


def api_call_row(api_name: str, company: Company, endpoint: str, result: dict[str, Any], status: str, credits_used: str, records: int, raw_rel: str, log_rel: str) -> dict[str, str]:
    return {
        "timestamp": now_iso(),
        "api_name": api_name,
        "company_name": company.company_name,
        "domain": company.domain,
        "endpoint_used": endpoint,
        "http_status": str(result.get("status_code", "")),
        "status": status,
        "latency_ms": str(result.get("latency_ms", "")),
        "credits_used": credits_used,
        "rate_limit": "; ".join(f"{k}: {v}" for k, v in (result.get("headers") or {}).items() if "rate" in k.lower() or k.lower() == "x-credits-left"),
        "records_retrieved": str(records),
        "error_message": str(result.get("error") or ""),
        "raw_response_path": raw_rel,
        "api_log_path": log_rel,
    }


def apollo_headers(api_key: str) -> list[dict[str, str]]:
    preferred = os.getenv("APOLLO_AUTH_MODE", "").strip().lower()
    modes = [preferred] if preferred in {"bearer", "x-api-key"} else ["bearer", "x-api-key"]
    headers = []
    for mode in modes:
        base = {"Accept": "application/json", "Content-Type": "application/json"}
        if mode == "x-api-key":
            base["X-Api-Key"] = api_key
        else:
            base["Authorization"] = f"Bearer {api_key}"
        headers.append(base)
    return headers


def first_value(obj: dict[str, Any], keys: list[str]) -> Any:
    for key in keys:
        value = obj.get(key)
        if value not in (None, "", [], {}):
            return value
    return ""


def find_people_list(payload: Any) -> list[dict[str, Any]]:
    if not isinstance(payload, dict):
        return []
    for key in ["people", "contacts", "data", "results", "profiles"]:
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
    for value in payload.values():
        if isinstance(value, dict):
            found = find_people_list(value)
            if found:
                return found
    return []


def split_name(name: str) -> tuple[str, str]:
    parts = [part for part in re.split(r"\s+", (name or "").strip()) if part]
    if not parts:
        return "", ""
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], parts[-1]


def normalize_seed(company: Company, person: dict[str, Any], source_api: str) -> SeedContact:
    full_name = str(first_value(person, ["name", "full_name", "display_name"]) or "").strip()
    first_name = str(first_value(person, ["first_name", "firstName"]) or "").strip()
    last_name = str(first_value(person, ["last_name", "lastName"]) or "").strip()
    if not (first_name and last_name):
        inferred_first, inferred_last = split_name(full_name)
        first_name = first_name or inferred_first
        last_name = last_name or inferred_last
    full_name = full_name or " ".join(part for part in [first_name, last_name] if part)
    org = person.get("organization") if isinstance(person.get("organization"), dict) else {}
    employment = person.get("employment") if isinstance(person.get("employment"), dict) else {}
    current = employment.get("current") if isinstance(employment.get("current"), dict) else {}
    social = person.get("social_profiles") if isinstance(person.get("social_profiles"), dict) else {}
    professional_network = social.get("professional_network") if isinstance(social.get("professional_network"), dict) else {}
    title = str(first_value(person, ["title", "job_title", "headline"]) or first_value(current, ["title", "job_title"]) or "")
    return SeedContact(
        company=company,
        first_name=first_name,
        last_name=last_name,
        full_name=full_name,
        title=title,
        seniority=str(first_value(person, ["seniority", "seniority_level", "management_level"]) or first_value(current, ["seniority"]) or ""),
        department=str(first_value(person, ["department", "departments", "function"]) or first_value(current, ["department", "job_functions"]) or ""),
        linkedin_url=str(first_value(person, ["linkedin_url", "linkedin", "linkedin_profile_url"]) or first_value(professional_network, ["url"]) or first_value(org, ["linkedin_url"]) or ""),
        source_api=source_api,
    )


def seed_key(seed: SeedContact) -> str:
    return "|".join([slugify(seed.company.company_name), seed.full_name.lower(), seed.linkedin_url.lower()])


def contact_name(seed: SeedContact) -> str:
    return seed.full_name or " ".join(part for part in [seed.first_name, seed.last_name] if part) or NOT_AVAILABLE


def base_contact_row(seed: SeedContact, api_name: str, raw_rel: str) -> dict[str, str]:
    return {
        "source_rank": seed.company.source_rank,
        "company_name": seed.company.company_name,
        "domain": seed.company.domain,
        "api_name": api_name,
        "requested_contact": contact_name(seed),
        "matched_contact": NOT_AVAILABLE,
        "match_status": "No Match",
        "match_confidence": NOT_AVAILABLE,
        "verified_work_email": NOT_AVAILABLE,
        "email_confidence": NOT_AVAILABLE,
        "direct_mobile_phone": NOT_AVAILABLE,
        "phone_confidence": NOT_AVAILABLE,
        "title": seed.title or NOT_AVAILABLE,
        "seniority": seed.seniority or NOT_AVAILABLE,
        "department": seed.department or NOT_AVAILABLE,
        "reports_to": NOT_AVAILABLE,
        "linkedin_url": seed.linkedin_url or NOT_AVAILABLE,
        "raw_response_path": raw_rel,
    }


def normalize_email(value: Any) -> tuple[str, str]:
    if isinstance(value, list):
        for item in value:
            email, confidence = normalize_email(item)
            if email != NOT_AVAILABLE:
                return email, confidence
    if isinstance(value, dict):
        email = first_value(value, ["email", "value", "address", "email_address"])
        confidence = first_value(value, ["confidence", "status", "verification_status", "type"])
        return str(email or NOT_AVAILABLE), str(confidence or NOT_AVAILABLE)
    if isinstance(value, str) and "@" in value:
        return value, NOT_AVAILABLE
    return NOT_AVAILABLE, NOT_AVAILABLE


def normalize_phone(value: Any) -> tuple[str, str]:
    if isinstance(value, list):
        for item in value:
            phone, confidence = normalize_phone(item)
            if phone != NOT_AVAILABLE:
                return phone, confidence
    if isinstance(value, dict):
        phone = first_value(value, ["phone", "value", "number", "phone_number", "mobile_phone"])
        confidence = first_value(value, ["confidence", "status", "type"])
        return str(phone or NOT_AVAILABLE), str(confidence or NOT_AVAILABLE)
    if isinstance(value, str) and re.search(r"\d", value):
        return value, NOT_AVAILABLE
    return NOT_AVAILABLE, NOT_AVAILABLE


def summarize_contacts(companies: list[Company], api_name: str, rows: list[dict[str, str]], raw_by_company: dict[str, str]) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    summaries: list[dict[str, str]] = []
    missing: list[dict[str, str]] = []
    for company in companies:
        company_rows = [row for row in rows if row["company_name"] == company.company_name and row["api_name"] == api_name]
        requested = len(company_rows)
        matched = sum(1 for row in company_rows if row["match_status"] == "Matched")
        email_count = sum(1 for row in company_rows if row["verified_work_email"] != NOT_AVAILABLE)
        phone_count = sum(1 for row in company_rows if row["direct_mobile_phone"] != NOT_AVAILABLE)
        linkedin_count = sum(1 for row in company_rows if row["linkedin_url"] != NOT_AVAILABLE)
        title_count = sum(1 for row in company_rows if row["title"] != NOT_AVAILABLE)
        reports_to_count = sum(1 for row in company_rows if row["reports_to"] != NOT_AVAILABLE)
        field_checks = [email_count, phone_count, title_count, reports_to_count, linkedin_count, matched]
        completeness = round(sum(1 for count in field_checks if count > 0) / len(TARGET_FIELDS) * 100, 2)
        raw_rel = raw_by_company.get(company.company_name, "")
        summaries.append(
            {
                "source_rank": company.source_rank,
                "company_name": company.company_name,
                "domain": company.domain,
                "api_name": api_name,
                "requested_contacts": str(requested),
                "matched_contacts": str(matched),
                "match_rate_percent": f"{(matched / requested * 100 if requested else 0):.2f}",
                "verified_work_email_count": str(email_count),
                "direct_mobile_phone_count": str(phone_count),
                "linkedin_url_count": str(linkedin_count),
                "title_count": str(title_count),
                "reports_to_count": str(reports_to_count),
                "data_completeness_percent": f"{completeness:.2f}",
                "raw_response_path": raw_rel,
            }
        )
        reasons = [
            (TARGET_FIELDS[0], email_count, "No verified work email returned in parsed contacts."),
            (TARGET_FIELDS[1], phone_count, "No direct/mobile phone returned in parsed contacts or phone enrichment was gated."),
            (TARGET_FIELDS[2], title_count, "No title/seniority/department returned in parsed contacts."),
            (TARGET_FIELDS[3], reports_to_count, "No reports-to/org-chart hierarchy returned."),
            (TARGET_FIELDS[4], linkedin_count, "No LinkedIn URL returned in parsed contacts."),
            (TARGET_FIELDS[5], matched, "No matched contacts returned."),
        ]
        for field, count, reason in reasons:
            if count:
                continue
            missing.append(
                {
                    "source_rank": company.source_rank,
                    "company_name": company.company_name,
                    "domain": company.domain,
                    "api_name": api_name,
                    "requested_contact": "Company level",
                    "field": field,
                    "status": NOT_AVAILABLE if requested else NOT_TESTED,
                    "reason": reason,
                    "raw_response_path": raw_rel,
                }
            )
    return summaries, missing


def collect_apollo(companies: list[Company], reuse_raw: bool, timeout: int, per_company: int) -> tuple[list[SeedContact], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Apollo People Search API"
    api_key = os.getenv("APOLLO_API_KEY") or os.getenv("APOLLO_MASTER_API_KEY") or ""
    contact_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    seeds: list[SeedContact] = []
    for company in companies:
        raw_path = RAW_ROOT / "apollo" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "apollo" / f"{slugify(company.company_name)}.jsonl"
        result: dict[str, Any]
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = result_from_saved_payload(payload)
        elif not api_key:
            payload = {"error": "Missing APOLLO_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            body = {
                "q_organization_domains_list": [company.domain],
                "per_page": per_company,
                "page": 1,
                "person_seniorities": ["owner", "founder", "c_suite", "vp", "director", "manager", "head"],
            }
            result = {}
            for headers in apollo_headers(api_key):
                result = http_request("POST", "https://api.apollo.io/api/v1/mixed_people/api_search", headers=headers, body=body, timeout=timeout)
                append_jsonl(log_path, {"request": {"method": "POST", "url": "https://api.apollo.io/api/v1/mixed_people/api_search", "headers": redact_headers(headers), "body": body}, "response": result})
                if result["ok"] or result["status_code"] not in {401, 403}:
                    break
            save_json(raw_path, result.get("payload"))
        payload = result.get("payload", {})
        raw_rel = safe_rel(raw_path)
        raw_by_company[company.company_name] = raw_rel
        people = find_people_list(payload)[:per_company]
        for person in people:
            seed = normalize_seed(company, person, api_name)
            seeds.append(seed)
            row = base_contact_row(seed, api_name, raw_rel)
            row.update(
                {
                    "matched_contact": contact_name(seed),
                    "match_status": "Matched",
                    "match_confidence": "Apollo people search result",
                    "verified_work_email": NOT_AVAILABLE,
                    "email_confidence": "Not returned by Apollo people search endpoint",
                    "direct_mobile_phone": NOT_AVAILABLE,
                    "phone_confidence": "Not returned by Apollo people search endpoint",
                }
            )
            contact_rows.append(row)
        status = "Success" if result.get("ok") else "Fail"
        call_rows.append(api_call_row(api_name, company, APOLLO_ENDPOINT, result, status, "0 (people search does not reveal email/phone)", len(people), raw_rel, safe_rel(log_path)))
        print(f"apollo: {company.company_name} | {status} | seeds={len(people)}")
    company_rows, missing_rows = summarize_contacts(companies, api_name, contact_rows, raw_by_company)
    return seeds, company_rows, contact_rows, missing_rows, call_rows


def collect_fullenrich_search_seeds(companies: list[Company], reuse_raw: bool, timeout: int, per_company: int) -> tuple[list[SeedContact], list[dict[str, str]]]:
    api_name = "FullEnrich People Search API"
    api_key = os.getenv("FULLENRICH_API_KEY") or os.getenv("FULLENRICH_API_TOKEN") or ""
    seeds: list[SeedContact] = []
    call_rows: list[dict[str, str]] = []
    for company in companies:
        raw_path = RAW_ROOT / "fullenrich_search" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "fullenrich_search" / f"{slugify(company.company_name)}.jsonl"
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = result_from_saved_payload(payload)
        elif not api_key:
            payload = {"error": "Missing FULLENRICH_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Accept": "application/json", "Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
            body = {
                "offset": 0,
                "limit": per_company,
                "current_company_domains": [{"value": company.domain, "exact_match": True, "exclude": False}],
                "current_position_seniority_level": [
                    {"value": "C-level", "exact_match": False, "exclude": False},
                    {"value": "VP", "exact_match": False, "exclude": False},
                    {"value": "Director", "exact_match": False, "exclude": False},
                    {"value": "Manager", "exact_match": False, "exclude": False},
                ],
            }
            result = http_request("POST", "https://app.fullenrich.com/api/v2/people/search", headers=headers, body=body, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://app.fullenrich.com/api/v2/people/search", "headers": redact_headers(headers), "body": body}, "response": result})
            save_json(raw_path, result.get("payload"))
        people = find_people_list(result.get("payload", {}))[:per_company]
        for person in people:
            seeds.append(normalize_seed(company, person, api_name))
        status = "Success" if result.get("ok") else "Fail"
        call_rows.append(api_call_row(api_name, company, "POST https://app.fullenrich.com/api/v2/people/search", result, status, "People search credits are account specific", len(people), safe_rel(raw_path), safe_rel(log_path)))
        print(f"fullenrich search: {company.company_name} | {status} | seeds={len(people)}")
    return seeds, call_rows


def load_saved_seed_contacts(companies: list[Company], per_company: int) -> list[SeedContact]:
    seeds: list[SeedContact] = []
    for company in companies:
        raw_path = RAW_ROOT / "fullenrich_search" / "data" / f"{slugify(company.company_name)}.json"
        if not raw_path.exists():
            continue
        payload = load_json(raw_path)
        for person in find_people_list(payload)[:per_company]:
            seeds.append(normalize_seed(company, person, "Saved FullEnrich People Search seed"))
    return seeds


def collect_apollo_enrichment(companies: list[Company], seeds: list[SeedContact], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Apollo People Enrichment API"
    api_key = os.getenv("APOLLO_API_KEY") or os.getenv("APOLLO_MASTER_API_KEY") or ""
    selected = limited_seeds(seeds, limit)
    contact_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for seed in selected:
        company = seed.company
        raw_path = RAW_ROOT / "apollo_enrichment" / "data" / f"{slugify(company.company_name)}__{slugify(contact_name(seed))}.json"
        log_path = APILOG_ROOT / "apollo_enrichment" / f"{slugify(company.company_name)}.jsonl"
        raw_by_company[company.company_name] = safe_rel(raw_path)
        params = {
            "name": contact_name(seed),
            "domain": company.domain,
            "reveal_personal_emails": "false",
            "reveal_phone_number": "false",
        }
        if seed.first_name:
            params["first_name"] = seed.first_name
        if seed.last_name:
            params["last_name"] = seed.last_name
        if seed.linkedin_url:
            params["linkedin_url"] = seed.linkedin_url
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = result_from_saved_payload(payload)
        elif not api_key:
            payload = {"error": "Missing APOLLO_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Accept": "application/json", "Content-Type": "application/json", "Cache-Control": "no-cache", "X-Api-Key": api_key}
            result = http_request("POST", "https://api.apollo.io/api/v1/people/match", headers=headers, params=params, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://api.apollo.io/api/v1/people/match", "headers": redact_headers(headers), "params": params}, "response": result})
            save_json(raw_path, result.get("payload"))
        row = parse_generic_enrichment(seed, api_name, result.get("payload", {}), safe_rel(raw_path))
        contact_rows.append(row)
        status = "Success" if result.get("ok") else "Fail"
        credits = "Apollo enrichment credit if matched; phone reveal disabled"
        call_rows.append(api_call_row(api_name, company, APOLLO_ENRICHMENT_ENDPOINT, result, status, credits, 1 if row["match_status"] == "Matched" else 0, safe_rel(raw_path), safe_rel(log_path)))
        print(f"apollo enrichment: {company.company_name} / {contact_name(seed)} | {status}")
    add_not_requested_rows(companies, selected, contact_rows, api_name)
    company_rows, missing_rows = summarize_contacts(companies, api_name, contact_rows, raw_by_company)
    return company_rows, contact_rows, missing_rows, call_rows


def limited_seeds(seeds: list[SeedContact], limit: int) -> list[SeedContact]:
    if limit <= 0:
        return []
    unique: dict[str, SeedContact] = {}
    for seed in seeds:
        if not seed.full_name and not seed.linkedin_url:
            continue
        unique.setdefault(seed_key(seed), seed)
        if len(unique) >= limit:
            break
    return list(unique.values())


def collect_prospeo(companies: list[Company], seeds: list[SeedContact], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "Prospeo Person Enrichment API"
    api_key = os.getenv("PROSPEO_API_KEY") or ""
    contact_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    selected = limited_seeds(seeds, limit)
    for seed in selected:
        company = seed.company
        raw_path = RAW_ROOT / "prospeo" / "data" / f"{slugify(company.company_name)}__{slugify(contact_name(seed))}.json"
        log_path = APILOG_ROOT / "prospeo" / f"{slugify(company.company_name)}.jsonl"
        raw_by_company[company.company_name] = safe_rel(raw_path)
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = result_from_saved_payload(payload)
        elif not api_key:
            payload = {"error": "Missing PROSPEO_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Accept": "application/json", "Content-Type": "application/json", "X-KEY": api_key}
            body = {
                "only_verified_email": True,
                "enrich_mobile": True,
                "data": {
                    "first_name": seed.first_name,
                    "last_name": seed.last_name,
                    "full_name": seed.full_name,
                    "company": company.company_name,
                    "company_name": company.company_name,
                    "company_website": company.domain,
                    "linkedin_url": seed.linkedin_url,
                },
            }
            result = http_request("POST", "https://api.prospeo.io/enrich-person", headers=headers, body=body, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://api.prospeo.io/enrich-person", "headers": redact_headers(headers), "body": body}, "response": result})
            save_json(raw_path, result.get("payload"))
        payload = result.get("payload", {})
        row = parse_generic_enrichment(seed, api_name, payload, safe_rel(raw_path))
        contact_rows.append(row)
        status = "Success" if result.get("ok") else "Fail"
        credits = "1 email credit if email found; +10 mobile credits if mobile found"
        call_rows.append(api_call_row(api_name, company, PROSPEO_ENDPOINT, result, status, credits, 1 if row["match_status"] == "Matched" else 0, safe_rel(raw_path), safe_rel(log_path)))
        print(f"prospeo: {company.company_name} / {contact_name(seed)} | {status}")
    add_not_requested_rows(companies, selected, contact_rows, api_name)
    company_rows, missing_rows = summarize_contacts(companies, api_name, contact_rows, raw_by_company)
    return company_rows, contact_rows, missing_rows, call_rows


def collect_fullenrich(companies: list[Company], seeds: list[SeedContact], reuse_raw: bool, timeout: int, limit: int, poll_attempts: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "FullEnrich Contact Bulk Enrichment API"
    api_key = os.getenv("FULLENRICH_API_KEY") or os.getenv("FULLENRICH_API_TOKEN") or ""
    selected = limited_seeds(seeds, limit)
    contact_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    grouped: dict[str, list[SeedContact]] = {}
    for seed in selected:
        grouped.setdefault(seed.company.company_name, []).append(seed)
    for company in companies:
        company_seeds = grouped.get(company.company_name, [])
        raw_path = RAW_ROOT / "fullenrich" / "data" / f"{slugify(company.company_name)}.json"
        log_path = APILOG_ROOT / "fullenrich" / f"{slugify(company.company_name)}.jsonl"
        raw_by_company[company.company_name] = safe_rel(raw_path)
        if not company_seeds:
            continue
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = result_from_saved_payload(payload)
        elif not api_key:
            payload = {"error": "Missing FULLENRICH_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Accept": "application/json", "Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
            data = []
            for seed in company_seeds:
                data.append(
                    {
                        "first_name": seed.first_name,
                        "last_name": seed.last_name,
                        "domain": company.domain,
                        "company_name": company.company_name,
                        "linkedin_url": seed.linkedin_url,
                        "enrich_fields": ["contact.work_emails", "contact.phones"],
                    }
                )
            body = {"name": f"HPI contact-level {company.company_name} {stamp()}", "data": data}
            result = http_request("POST", "https://app.fullenrich.com/api/v2/contact/enrich/bulk", headers=headers, body=body, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://app.fullenrich.com/api/v2/contact/enrich/bulk", "headers": redact_headers(headers), "body": body}, "response": result})
            payload = result.get("payload", {})
            enrichment_id = str(first_value(payload if isinstance(payload, dict) else {}, ["enrichment_id", "id", "bulk_id"]) or "")
            final_payload: Any = payload
            if result.get("ok") and enrichment_id:
                for _ in range(max(poll_attempts, 0)):
                    time.sleep(10)
                    poll = http_request("GET", f"https://app.fullenrich.com/api/v1/contact/enrich/bulk/{enrichment_id}", headers=headers, timeout=timeout)
                    append_jsonl(log_path, {"request": {"method": "GET", "url": f"https://app.fullenrich.com/api/v1/contact/enrich/bulk/{enrichment_id}", "headers": redact_headers(headers)}, "response": poll})
                    final_payload = {"start_response": payload, "result_response": poll.get("payload")}
                    if poll.get("ok") and str(first_value(poll.get("payload") if isinstance(poll.get("payload"), dict) else {}, ["status", "state"]) or "").lower() in {"finished", "completed", "done", "success"}:
                        break
            result = {**result, "payload": final_payload}
            save_json(raw_path, final_payload)
        payload = result.get("payload", {})
        for seed in company_seeds:
            row = parse_generic_enrichment(seed, api_name, payload, safe_rel(raw_path))
            contact_rows.append(row)
        status = "Success" if result.get("ok") else "Fail"
        call_rows.append(api_call_row(api_name, company, FULLENRICH_BULK_ENDPOINT, result, status, "Bulk enrichment credits; email/phone waterfall cost is account specific", len(company_seeds), safe_rel(raw_path), safe_rel(log_path)))
        print(f"fullenrich: {company.company_name} | {status} | contacts={len(company_seeds)}")
    add_not_requested_rows(companies, selected, contact_rows, api_name)
    company_rows, missing_rows = summarize_contacts(companies, api_name, contact_rows, raw_by_company)
    return company_rows, contact_rows, missing_rows, call_rows


def collect_signalhire(companies: list[Company], seeds: list[SeedContact], reuse_raw: bool, timeout: int, limit: int) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "SignalHire Candidate Search API"
    api_key = os.getenv("SIGNALHIRE_API_KEY") or ""
    selected = [seed for seed in limited_seeds(seeds, limit) if seed.linkedin_url]
    contact_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for seed in selected:
        company = seed.company
        raw_path = RAW_ROOT / "signalhire" / "data" / f"{slugify(company.company_name)}__{slugify(contact_name(seed))}.json"
        log_path = APILOG_ROOT / "signalhire" / f"{slugify(company.company_name)}.jsonl"
        raw_by_company[company.company_name] = safe_rel(raw_path)
        if reuse_raw and raw_path.exists():
            payload = load_json(raw_path)
            result = result_from_saved_payload(payload)
        elif not api_key:
            payload = {"error": "Missing SIGNALHIRE_API_KEY."}
            result = {"ok": False, "status_code": "", "latency_ms": 0, "payload": payload, "headers": {}, "error": payload["error"]}
            save_json(raw_path, payload)
        else:
            headers = {"Accept": "application/json", "Content-Type": "application/json", "apikey": api_key}
            body = {"items": [seed.linkedin_url], "withoutWaterfall": True}
            result = http_request("POST", "https://www.signalhire.com/api/v1/candidate/search", headers=headers, body=body, timeout=timeout)
            append_jsonl(log_path, {"request": {"method": "POST", "url": "https://www.signalhire.com/api/v1/candidate/search", "headers": redact_headers(headers), "body": body}, "response": result})
            save_json(raw_path, result.get("payload"))
        row = parse_generic_enrichment(seed, api_name, result.get("payload", {}), safe_rel(raw_path))
        contact_rows.append(row)
        status = "Success" if result.get("ok") else "Fail"
        credits_left = (result.get("headers") or {}).get("X-Credits-Left", "")
        credits = f"1 standard credit when matched; credits left: {credits_left}" if credits_left else "1 standard credit when matched; capped at 5 lookups"
        call_rows.append(api_call_row(api_name, company, SIGNALHIRE_ENDPOINT, result, status, credits, 1 if row["match_status"] == "Matched" else 0, safe_rel(raw_path), safe_rel(log_path)))
        print(f"signalhire: {company.company_name} / {contact_name(seed)} | {status}")
    add_not_requested_rows(companies, selected, contact_rows, api_name)
    company_rows, missing_rows = summarize_contacts(companies, api_name, contact_rows, raw_by_company)
    return company_rows, contact_rows, missing_rows, call_rows


def parse_generic_enrichment(seed: SeedContact, api_name: str, payload: Any, raw_rel: str) -> dict[str, str]:
    row = base_contact_row(seed, api_name, raw_rel)
    if isinstance(payload, dict) and (payload.get("error") or payload.get("error_code")):
        row["match_confidence"] = str(payload.get("error_code") or payload.get("error"))
        return row
    candidates = flatten_dicts(payload)
    best = choose_best_candidate(seed, candidates)
    if not best:
        return row
    email, email_conf = find_first_email(best)
    phone, phone_conf = find_first_phone(best)
    employment = best.get("employment") if isinstance(best.get("employment"), dict) else {}
    current = employment.get("current") if isinstance(employment.get("current"), dict) else {}
    social = best.get("social_profiles") if isinstance(best.get("social_profiles"), dict) else {}
    professional_network = social.get("professional_network") if isinstance(social.get("professional_network"), dict) else {}
    title = str(first_value(best, ["title", "job_title", "position", "headline"]) or first_value(current, ["title"]) or seed.title or NOT_AVAILABLE)
    linkedin = str(first_value(best, ["linkedin_url", "linkedin", "linkedin_profile_url", "url"]) or first_value(professional_network, ["url"]) or seed.linkedin_url or NOT_AVAILABLE)
    full_name = str(first_value(best, ["name", "full_name", "display_name"]) or contact_name(seed))
    reports_to = find_reports_to(best)
    row.update(
        {
            "matched_contact": full_name,
            "match_status": "Matched",
            "match_confidence": str(first_value(best, ["confidence", "match_confidence", "score"]) or "Provider match"),
            "verified_work_email": email,
            "email_confidence": email_conf,
            "direct_mobile_phone": phone,
            "phone_confidence": phone_conf,
            "title": title,
            "seniority": str(first_value(best, ["seniority", "seniority_level", "management_level"]) or first_value(current, ["seniority"]) or seed.seniority or NOT_AVAILABLE),
            "department": str(first_value(best, ["department", "departments", "function"]) or first_value(current, ["job_functions"]) or seed.department or NOT_AVAILABLE),
            "reports_to": reports_to,
            "linkedin_url": linkedin,
        }
    )
    return row


def flatten_dicts(payload: Any) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    if isinstance(payload, dict):
        out.append(payload)
        for value in payload.values():
            if isinstance(value, (dict, list)):
                out.extend(flatten_dicts(value))
    elif isinstance(payload, list):
        for item in payload:
            out.extend(flatten_dicts(item))
    return out


def choose_best_candidate(seed: SeedContact, candidates: list[dict[str, Any]]) -> dict[str, Any]:
    target_name = contact_name(seed).lower()
    for candidate in candidates:
        linkedin = str(first_value(candidate, ["linkedin_url", "linkedin", "linkedin_profile_url", "url"]) or "").lower()
        if seed.linkedin_url and linkedin and seed.linkedin_url.lower().rstrip("/") in linkedin.rstrip("/"):
            return candidate
    for candidate in candidates:
        name = str(first_value(candidate, ["name", "full_name", "display_name"]) or "").lower()
        if target_name != NOT_AVAILABLE.lower() and name and (target_name in name or name in target_name):
            return candidate
    for candidate in candidates:
        email, _ = find_first_email(candidate)
        phone, _ = find_first_phone(candidate)
        if email != NOT_AVAILABLE or phone != NOT_AVAILABLE:
            return candidate
    return candidates[0] if candidates else {}


def find_first_email(obj: dict[str, Any]) -> tuple[str, str]:
    for key in ["work_email", "verified_email", "professional_email", "email", "email_address", "emails", "work_emails"]:
        if key in obj:
            email, confidence = normalize_email(obj.get(key))
            if email != NOT_AVAILABLE:
                return email, confidence
    for value in obj.values():
        if isinstance(value, (dict, list)):
            for nested in flatten_dicts(value):
                for key in ["work_email", "verified_email", "professional_email", "email", "email_address", "emails", "work_emails"]:
                    if key in nested:
                        email, confidence = normalize_email(nested.get(key))
                        if email != NOT_AVAILABLE:
                            return email, confidence
    return NOT_AVAILABLE, NOT_AVAILABLE


def find_first_phone(obj: dict[str, Any]) -> tuple[str, str]:
    for key in ["mobile_phone", "phone", "phone_number", "direct_phone", "phones", "mobile_phones"]:
        if key in obj:
            phone, confidence = normalize_phone(obj.get(key))
            if phone != NOT_AVAILABLE:
                return phone, confidence
    for value in obj.values():
        if isinstance(value, (dict, list)):
            for nested in flatten_dicts(value):
                for key in ["mobile_phone", "phone", "phone_number", "direct_phone", "phones", "mobile_phones"]:
                    if key in nested:
                        phone, confidence = normalize_phone(nested.get(key))
                        if phone != NOT_AVAILABLE:
                            return phone, confidence
    return NOT_AVAILABLE, NOT_AVAILABLE


def find_reports_to(obj: dict[str, Any]) -> str:
    for key in ["reports_to", "manager", "manager_name", "boss", "org_chart", "hierarchy"]:
        value = obj.get(key)
        if not value:
            continue
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return str(first_value(value, ["name", "full_name", "title"]) or json.dumps(value, ensure_ascii=False)[:200])
        return json.dumps(value, ensure_ascii=False)[:200]
    return NOT_AVAILABLE


def add_not_requested_rows(companies: list[Company], selected: list[SeedContact], contact_rows: list[dict[str, str]], api_name: str) -> None:
    selected_companies = {seed.company.company_name for seed in selected}
    for company in companies:
        if company.company_name in selected_companies:
            continue
        raw_rel = safe_rel(RAW_ROOT / slugify(api_name) / "data" / f"{slugify(company.company_name)}.json")
        seed = SeedContact(company, "", "", "Not requested - credit cap", "", "", "", "", "credit cap")
        row = base_contact_row(seed, api_name, raw_rel)
        row["match_status"] = NOT_TESTED
        contact_rows.append(row)


def collect_blocked(companies: list[Company]) -> tuple[list[dict[str, str]], list[dict[str, str]], list[dict[str, str]], list[dict[str, str]]]:
    api_name = "People Data Labs Person Enrichment API"
    reason = "Marked server_down/not required for this run; API was not called."
    contact_rows: list[dict[str, str]] = []
    call_rows: list[dict[str, str]] = []
    raw_by_company: dict[str, str] = {}
    for company in companies:
        raw_path = RAW_ROOT / "people_data_labs" / "data" / f"{slugify(company.company_name)}.json"
        save_json(raw_path, {"status": "not_executed", "reason": reason})
        raw_rel = safe_rel(raw_path)
        raw_by_company[company.company_name] = raw_rel
        seed = SeedContact(company, "", "", "Not executed", "", "", "", "", "blocked")
        row = base_contact_row(seed, api_name, raw_rel)
        row["match_status"] = NOT_TESTED
        contact_rows.append(row)
        result = {"ok": False, "status_code": NOT_TESTED, "latency_ms": 0, "headers": {}, "error": reason}
        call_rows.append(api_call_row(api_name, company, "Server down / not called", result, NOT_TESTED, "0", 0, raw_rel, ""))
    company_rows, missing_rows = summarize_contacts(companies, api_name, contact_rows, raw_by_company)
    for row in missing_rows:
        row["reason"] = reason
    return company_rows, contact_rows, missing_rows, call_rows


def fields_returned(rows: list[dict[str, str]], api_name: str) -> list[str]:
    api_rows = [row for row in rows if row["api_name"] == api_name]
    checks = {
        "verified_work_email": "Verified work email",
        "direct_mobile_phone": "Direct/mobile phone",
        "title": "Title/seniority/department",
        "reports_to": "Reports-to hierarchy",
        "linkedin_url": "LinkedIn URL",
    }
    returned = []
    for key, label in checks.items():
        if any(row.get(key) not in {"", NOT_AVAILABLE, NOT_TESTED} for row in api_rows):
            returned.append(label)
    if any(row.get("match_status") == "Matched" for row in api_rows):
        returned.append("Match rate")
    return returned


def build_api_comparison(contact_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    configs = [
        ("Apollo People Search API", APOLLO_ENDPOINT, "Used as seed source for named contacts; email/phone not returned by people search.", "People search endpoint does not expose work email or phone.", "Apollo enrichment endpoints can add gated email/phone if additional credits are approved."),
        ("Apollo People Enrichment API", APOLLO_ENRICHMENT_ENDPOINT, "Enriched saved seed contacts by name, domain, and LinkedIn URL with phone reveal disabled.", "Consumes Apollo enrichment credits when matched; phone reveal was disabled for this run.", "Can reveal work email and optionally phone/waterfall data when plan and credits allow it."),
        ("FullEnrich Contact Bulk Enrichment API", FULLENRICH_BULK_ENDPOINT, "Bulk waterfall contact enrichment for capped seed contacts.", "Async job and enrichment credits are account specific.", "Waterfall enrichment across providers for emails and phones."),
        ("Prospeo Person Enrichment API", PROSPEO_ENDPOINT, "Person enrichment for capped seed contacts.", "Email credit used only if email found; mobile is higher cost.", "Verified email plus mobile enrichment when credits are available."),
        ("SignalHire Candidate Search API", SIGNALHIRE_ENDPOINT, "Capped at 5 lookups because user reported very low credits.", "Very low credit balance; run capped and credit headers logged.", "Can return phones/emails/social profiles for matched candidates."),
        ("People Data Labs Person Enrichment API", "Server down / not called", "Not executed per user status.", "Server_down/not required for this evaluation.", "Large contact/person enrichment if access stabilizes."),
    ]
    rows = []
    for api_name, endpoint, notes, free, paid in configs:
        returned = fields_returned(contact_rows, api_name)
        status = "Success" if returned else "Not Tested" if api_name == "People Data Labs Person Enrichment API" else "Fail"
        rows.append({"API Name": api_name, "Endpoint Used": endpoint, "Status (Success/Fail)": status, "Fields Returned": "; ".join(returned) if returned else NOT_AVAILABLE, "Free-Tier Limitations": free, "Paid-Tier Benefits": paid, "Notes": notes})
    return rows


def build_api_trace(company_rows: list[dict[str, str]], contact_rows: list[dict[str, str]], call_rows: list[dict[str, str]], missing_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    configs = {
        "Apollo People Search API": ("Bearer or X-Api-Key", APOLLO_DOC, "Existing Apollo access", "Plan-specific", "People search and paid enrichment endpoints", "Plan-specific", "Useful seed source, but email/phone requires enrichment."),
        "Apollo People Enrichment API": ("X-Api-Key header", "https://docs.apollo.io/reference/people-enrichment", "Apollo enrichment credits", "Plan-specific", "Person enrichment and optional email/phone reveal", "Plan-specific", "Uses saved seed contacts; phone reveal disabled to avoid extra credit exposure."),
        "FullEnrich Contact Bulk Enrichment API": ("Bearer API key", FULLENRICH_DOC, "Account-specific waterfall credits", "Account-specific", "Waterfall enrichment across contact providers", "Account-specific", "Async workflow; raw start/result responses are saved."),
        "Prospeo Person Enrichment API": ("X-KEY header", PROSPEO_DOC, "Email credit if found; mobile costs more", "Email/mobile credits", "Verified emails and mobile enrichment", "Account-specific", "Straightforward synchronous REST endpoint."),
        "SignalHire Candidate Search API": ("apikey header", SIGNALHIRE_DOC, "Only 5 credits reported by user", "X-Credits-Left response header captured where returned", "Candidate profile enrichment", "Account-specific", "Run capped at 5 lookups to preserve credits."),
        "People Data Labs Person Enrichment API": ("API key", PDL_DOC, "Server_down/not used", "Not assessed", "Person enrichment at scale", "Account-specific", "Not executed per user status."),
    }
    rows = []
    for api_name, (auth, evidence, free, rate, paid, cost, remarks) in configs.items():
        calls = [row for row in call_rows if row["api_name"] == api_name]
        summaries = [row for row in company_rows if row["api_name"] == api_name]
        successes = [row for row in calls if row["status"] == "Success"]
        latencies = [float(row["latency_ms"]) for row in calls if str(row["latency_ms"]).replace(".", "", 1).isdigit()]
        completeness = [float(row["data_completeness_percent"]) for row in summaries if row.get("data_completeness_percent")]
        records = sum(int(row["records_retrieved"]) for row in calls if row["records_retrieved"].isdigit())
        missing = sorted({row["field"] for row in missing_rows if row["api_name"] == api_name and row["status"] != NOT_TESTED})
        processed = len(calls)
        success_rate = len(successes) / processed * 100 if processed else 0
        status = "Not Tested" if api_name == "People Data Labs Person Enrichment API" else ("Success" if successes else "Fail")
        avg_completeness = sum(completeness) / len(completeness) if completeness else 0
        rows.append(
            {
                "Tool Name": api_name,
                "Category": "Contact-Level",
                "API Available (Y/N)": "N" if api_name == "People Data Labs Person Enrichment API" else "Y",
                "Authentication Type": auth,
                "Free Credits / Tokens": free,
                "Credits / Tokens Used": "; ".join(row["credits_used"] for row in calls if row["credits_used"]) or NOT_AVAILABLE,
                "Rate Limit": rate,
                "Companies Processed": str(processed),
                "Coverage (%)": f"{success_rate:.2f}",
                "Success Rate (%)": f"{success_rate:.2f}",
                "Error Rate (%)": f"{(100 - success_rate if processed else 0):.2f}",
                "Average Latency": f"{(sum(latencies) / len(latencies)):.2f} ms" if latencies else NOT_AVAILABLE,
                "Gated Fields": "; ".join(missing) if missing else "None identified in parsed output",
                "Free Tier Limitation": free,
                "Paid Plan Cost": cost,
                "Paid Tier Benefits": paid,
                "Ease of Integration": "Good - standard REST/JSON" if api_name != "People Data Labs Person Enrichment API" else "Not assessed",
                "API Documentation Quality": "Good - endpoint documented" if api_name != "People Data Labs Person Enrichment API" else "Not assessed",
                "Evidence Link": evidence,
                "Overall API Score": score_api(success_rate, avg_completeness, records),
                "Status": status,
                "Remarks": remarks,
                "Data Completeness (%)": f"{avg_completeness:.2f}",
                "Records Retrieved": str(records),
                "Estimated Cost per 100 Companies": estimate_cost(api_name),
                "Raw Export Saved (Y/N)": "Y" if any(row["raw_response_path"] for row in calls) else "N",
            }
        )
    return rows


def score_api(success_rate: float, completeness: float, records: int) -> str:
    score = 1.0 + min(success_rate / 100 * 1.5, 1.5) + min(completeness / 100 * 1.5, 1.5) + (1.0 if records else 0.0)
    return f"{min(score, 5.0):.2f}/5"


def estimate_cost(api_name: str) -> str:
    if api_name == "Apollo People Search API":
        return "About 100 people-search calls for 100 companies; email/phone enrichment would add plan-specific credits."
    if api_name == "Apollo People Enrichment API":
        return "One enrichment request per requested contact; phone reveal/waterfall would add plan-specific credits."
    if api_name == "FullEnrich Contact Bulk Enrichment API":
        return "Depends on selected enrichment waterfall fields and matched contacts."
    if api_name == "Prospeo Person Enrichment API":
        return "At least one email credit per found contact; mobile enrichment costs more."
    if api_name == "SignalHire Candidate Search API":
        return "One standard candidate lookup credit per matched item; run capped at 5."
    return "Not assessed due server_down/not executed."


def write_workbook(path: Path, sheets: dict[str, tuple[list[str], list[dict[str, Any]]]]) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    for title, (fields, rows) in sheets.items():
        ws = wb.create_sheet(title[:31])
        ws.append(fields)
        for row in rows:
            ws.append([row.get(field, "") for field in fields])
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = Font(bold=True, color="FFFFFF")
        ws.freeze_panes = "A2"
        for column in ws.columns:
            col = get_column_letter(column[0].column)
            width = max(12, min(max(len(str(cell.value or "")) for cell in column) + 2, 60))
            ws.column_dimensions[col].width = width
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(path)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F4E78")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    document.add_paragraph()


def write_docx(path: Path, company_rows: list[dict[str, str]], trace_rows: list[dict[str, str]], comparison_rows: list[dict[str, str]], missing_rows: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HPI Contact-Level API Evaluation Report")
    run.font.bold = True
    run.font.size = Pt(18)
    document.add_paragraph(f"Generated: {now_iso()}")
    document.add_paragraph("Scope: Contact-level provider evaluation for the 10-company HPI pilot set. SignalHire was capped at 5 lookups because the account has very low credits.")
    document.add_heading("API Trace Summary", level=1)
    add_table(document, API_TRACE_FIELDS, [[row.get(field, "") for field in API_TRACE_FIELDS] for row in trace_rows])
    document.add_heading("API Comparison", level=1)
    add_table(document, API_COMPARISON_FIELDS, [[row.get(field, "") for field in API_COMPARISON_FIELDS] for row in comparison_rows])
    document.add_heading("Company Contact Summary", level=1)
    compact_fields = ["company_name", "api_name", "requested_contacts", "matched_contacts", "match_rate_percent", "verified_work_email_count", "direct_mobile_phone_count", "data_completeness_percent"]
    add_table(document, ["Company", "API", "Requested", "Matched", "Match %", "Emails", "Phones", "Completeness %"], [[row.get(field, "") for field in compact_fields] for row in company_rows])
    document.add_heading("Missing / Not Tested Fields", level=1)
    add_table(document, MISSING_FIELDS, [[row.get(field, "") for field in MISSING_FIELDS] for row in missing_rows])
    document.save(path)


def run(args: argparse.Namespace) -> int:
    load_dotenv()
    ensure_dirs()
    copy_input()
    companies = read_companies(Path(args.input) if args.input else INPUT_FILE, args.limit)
    if not companies:
        raise ValueError("No companies found for Contact-Level pipeline.")

    all_company_rows: list[dict[str, str]] = []
    all_contact_rows: list[dict[str, str]] = []
    all_missing_rows: list[dict[str, str]] = []
    all_call_rows: list[dict[str, str]] = []

    api_filter = {item.strip().lower() for item in args.apis.split(",") if item.strip()}
    run_all = "all" in api_filter

    seeds: list[SeedContact] = []
    if run_all or "apollo" in api_filter:
        seeds, company_rows, contact_rows, missing_rows, call_rows = collect_apollo(companies, args.reuse_raw, args.timeout, args.apollo_contacts_per_company)
        all_company_rows.extend(company_rows)
        all_contact_rows.extend(contact_rows)
        all_missing_rows.extend(missing_rows)
        all_call_rows.extend(call_rows)

    if run_all or "apollo_enrichment" in api_filter:
        enrichment_seeds = seeds or load_saved_seed_contacts(companies, args.apollo_contacts_per_company)
        company_rows, contact_rows, missing_rows, call_rows = collect_apollo_enrichment(companies, enrichment_seeds, args.reuse_raw, args.timeout, args.apollo_enrichment_limit)
        all_company_rows.extend(company_rows)
        all_contact_rows.extend(contact_rows)
        all_missing_rows.extend(missing_rows)
        all_call_rows.extend(call_rows)

    if not seeds and (run_all or "fullenrich_search" in api_filter):
        fallback_seeds, fallback_call_rows = collect_fullenrich_search_seeds(companies, args.reuse_raw, args.timeout, args.apollo_contacts_per_company)
        seeds.extend(fallback_seeds)
        all_call_rows.extend(fallback_call_rows)

    collectors = []
    if run_all or "fullenrich" in api_filter:
        collectors.append(collect_fullenrich(companies, seeds, args.reuse_raw, args.timeout, args.enrichment_limit, args.fullenrich_poll_attempts))
    if run_all or "prospeo" in api_filter:
        collectors.append(collect_prospeo(companies, seeds, args.reuse_raw, args.timeout, args.enrichment_limit))
    if run_all or "signalhire" in api_filter:
        collectors.append(collect_signalhire(companies, seeds, args.reuse_raw, args.timeout, args.signalhire_limit))
    if run_all or "people_data_labs" in api_filter:
        collectors.append(collect_blocked(companies))
    for company_rows, contact_rows, missing_rows, call_rows in collectors:
        all_company_rows.extend(company_rows)
        all_contact_rows.extend(contact_rows)
        all_missing_rows.extend(missing_rows)
        all_call_rows.extend(call_rows)

    comparison_rows = build_api_comparison(all_contact_rows)
    trace_rows = build_api_trace(all_company_rows, all_contact_rows, all_call_rows, all_missing_rows)

    write_csv(REPORTS_DIR / "company_contact_level.csv", all_company_rows, COMPANY_FIELDS)
    write_csv(REPORTS_DIR / "contact_detail.csv", all_contact_rows, CONTACT_FIELDS)
    write_csv(REPORTS_DIR / "api_call_log.csv", all_call_rows, API_CALL_FIELDS)
    write_csv(REPORTS_DIR / "api_comparison_report.csv", comparison_rows, API_COMPARISON_FIELDS)
    write_csv(REPORTS_DIR / "api_tracing_report.csv", trace_rows, API_TRACE_FIELDS)
    write_csv(REPORTS_DIR / "missing_fields_report.csv", all_missing_rows, MISSING_FIELDS)
    save_json(REPORTS_DIR / "company_contact_level.json", all_company_rows)
    save_json(
        REPORTS_DIR / "run_manifest.json",
        {
            "generated_at": now_iso(),
            "companies_processed": len(companies),
            "providers": ["apollo", "apollo_enrichment", "fullenrich", "prospeo", "signalhire", "people_data_labs"],
            "apis": sorted(api_filter),
            "reuse_raw": args.reuse_raw,
            "apollo_contacts_per_company": args.apollo_contacts_per_company,
            "enrichment_limit": args.enrichment_limit,
            "signalhire_limit": args.signalhire_limit,
        },
    )

    suffix = stamp()
    xlsx_path = REPORTS_DIR / f"hpi_contact_level_api_evaluation_{suffix}.xlsx"
    docx_path = REPORTS_DIR / f"hpi_contact_level_api_evaluation_{suffix}.docx"
    write_workbook(
        xlsx_path,
        {
            "API Comparison": (API_COMPARISON_FIELDS, comparison_rows),
            "API Trace": (API_TRACE_FIELDS, trace_rows),
            "API Call Log": (API_CALL_FIELDS, all_call_rows),
            "Company Summary": (COMPANY_FIELDS, all_company_rows),
            "Contact Detail": (CONTACT_FIELDS, all_contact_rows),
            "Missing Fields": (MISSING_FIELDS, all_missing_rows),
        },
    )
    write_docx(docx_path, all_company_rows, trace_rows, comparison_rows, all_missing_rows)
    print(f"Saved workbook: {safe_rel(xlsx_path)}")
    print(f"Saved docx: {safe_rel(docx_path)}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Run HPI Contact-Level API evaluation.")
    parser.add_argument("--input", help="Input CSV path. Defaults to ContactLevel/input/compnys.txt.")
    parser.add_argument("--limit", type=int, default=10)
    parser.add_argument("--reuse-raw", action="store_true", help="Reuse saved raw responses.")
    parser.add_argument("--timeout", type=int, default=60)
    parser.add_argument("--apollo-contacts-per-company", type=int, default=2)
    parser.add_argument("--apollo-enrichment-limit", type=int, default=5, help="Maximum saved seed contacts sent to Apollo People Enrichment.")
    parser.add_argument("--enrichment-limit", type=int, default=5, help="Maximum seed contacts sent to FullEnrich and Prospeo.")
    parser.add_argument("--signalhire-limit", type=int, default=5, help="Maximum SignalHire candidate lookups. Keep at 5 unless more credits are approved.")
    parser.add_argument("--fullenrich-poll-attempts", type=int, default=2, help="Poll attempts for async FullEnrich bulk results.")
    parser.add_argument("--apis", default="all", help="Comma-separated APIs to run: all, apollo, apollo_enrichment, fullenrich_search, fullenrich, prospeo, signalhire, people_data_labs.")
    args = parser.parse_args()
    try:
        return run(args)
    except Exception as exc:
        print(f"Contact-Level pipeline failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
